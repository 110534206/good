from flask import Blueprint, request, jsonify, render_template, session, send_file, current_app
from config import get_db
from datetime import datetime
from werkzeug.utils import secure_filename
import os
import traceback
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from notification import create_notification
from semester import get_current_semester_code

company_bp = Blueprint("company_bp", __name__)

# =========================================================
# 📁 上傳設定
# =========================================================
UPLOAD_FOLDER = "uploads/company_docs"
ALLOWED_EXTENSIONS = {"docx", "doc"}

def ensure_upload_folder():
    project_root = os.path.dirname(current_app.root_path)
    upload_path = os.path.join(project_root, UPLOAD_FOLDER)
    os.makedirs(upload_path, exist_ok=True)
    return upload_path

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# =========================================================
# 📄 生成實習單位基本資料表 Word 檔
# =========================================================
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def generate_company_word_document(data):
    doc = Document()

    # --- 內部輔助：設定格式、字型、以及對齊方式 ---
    def set_cell_format(cell, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, distribute=False):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        
        # 實作分散對齊效果
        if distribute:
            p.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            # 標籤格微調左右邊距
            p.paragraph_format.left_indent = Pt(5)
            p.paragraph_format.right_indent = Pt(5)
        else:
            p.alignment = alignment
            p.paragraph_format.left_indent = Pt(6)
            
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        if text:
            run = p.add_run(str(text))
            run.font.name = '標楷體'
            run.font.size = Pt(12)
            run.bold = bold
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), '標楷體')
            rFonts.set(qn('w:ascii'), '標楷體')
            rFonts.set(qn('w:hAnsi'), '標楷體')

    def apply_table_style(table, col_widths, min_row_height=480):
        tblPr = table._tbl.tblPr
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        
        borders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            node = OxmlElement(f'w:{b}')
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), '4')
            node.set(qn('w:color'), '000000')
            borders.append(node)
        tblPr.append(borders)
        
        for i, w in enumerate(col_widths):
            if i < len(table.columns):
                table.columns[i].width = Inches(w)
            
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(min_row_height))
            trHeight.set(qn('w:hRule'), 'atLeast') 
            trPr.append(trHeight)

    # --- 頁面設定 ---
    section = doc.sections[0]
    section.left_margin = Inches(0.45) # 稍微縮減邊距以容納 7.6 吋表格
    section.right_margin = Inches(0.45)

    # --- 標題區 ---
    titles = [
        ('康寧學校財團法人康寧大學資訊管理科', 12, False),
        ('實習單位基本資料表', 16, True),
        ('實習期間：115年2月23日至115年6月26日止', 12, False)
    ]
    for text, size, is_bold in titles:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = '標楷體'
        run.font.size = Pt(size)
        run.bold = is_bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # 總寬度 7.6 吋分配：1.1 + 2.7 + 1.1 + 2.7 = 7.6
    STD_WIDTHS = [1.1, 2.7, 1.1, 2.7] 

    # --- 1. 基本資訊表格 ---
    table_rows = [
        ('編 號', data.get('serial_number', ''), '', '', True),
        ('單 位 名 稱', data.get('company_name', ''), '', '', True),
        ('負 責 人', data.get('owner', ''), '統一編號', data.get('tax_id', ''), False),
        ('聯 絡 人', data.get('contact_person', ''), '職 稱', data.get('contact_title', ''), False),
        ('聯 絡 電 話', data.get('contact_phone', ''), '傳 真', data.get('fax', ''), False),
        ('地 址', data.get('address', ''), '', '', True),
        ('交 通 說 明', data.get('traffic_guide', ''), '', '', True),
        ('E-mail', data.get('email', ''), '', '', True),
    ]

    basic_table = doc.add_table(rows=len(table_rows), cols=4)
    apply_table_style(basic_table, STD_WIDTHS)

    for i, (l_lab, l_val, r_lab, r_val, merge) in enumerate(table_rows):
        cells = basic_table.rows[i].cells
        set_cell_format(cells[0], l_lab, distribute=True) 
        
        if merge:
            cells[1].merge(cells[2]); cells[1].merge(cells[3])
            set_cell_format(cells[1], l_val, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        else:
            set_cell_format(cells[1], l_val, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_format(cells[2], r_lab, distribute=True)
            set_cell_format(cells[3], r_val, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # --- 2. 單位簡介、營業項目、企業規模 ---
    for lab, key in [('單 位 簡 介', 'company_intro'), ('營 業 項 目', 'business_scope'), ('企 業 規 模', 'company_scale')]:
        h = 1000 if '簡介' in lab else 480
        t = doc.add_table(rows=1, cols=4)
        apply_table_style(t, STD_WIDTHS, min_row_height=h)
        cells = t.rows[0].cells
        set_cell_format(cells[0], lab, distribute=True)
        cells[1].merge(cells[2]); cells[1].merge(cells[3])
        
        if '規模' in lab:
            opts = ['1000人以上', '500-999人', '100-499人', '10-99人', '10以下']
            val = data.get(key, '')
            text = ''.join([f'{"☑" if o == val else "☐"} {o}   ' for o in opts])
            set_cell_format(cells[1], text, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        else:
            set_cell_format(cells[1], data.get(key, ''), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # --- 3. 職缺明細 ---
    jobs = data.get('jobs', [])
    if jobs:
        jobs_table = doc.add_table(rows=len(jobs) + 1, cols=4)
        # 7.6 吋重新分配：1.1, 1.8, 3.6, 1.1
        apply_table_style(jobs_table, [1.1, 1.8, 3.6, 1.1])
        headers = ['工 作 編 號', '工 作 項 目', '需求條件/工作內容', '名 額']
        for i, h in enumerate(headers):
            set_cell_format(jobs_table.rows[0].cells[i], h, distribute=True)
            
        for idx, job in enumerate(jobs, 1):
            row = jobs_table.rows[idx].cells
            set_cell_format(row[0], str(idx), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_format(row[1], job.get('title', ''), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_format(row[2], job.get('description', ''), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_format(row[3], str(job.get('slots', 1)), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # --- 4. 待遇和來源 ---
    final_table = doc.add_table(rows=2, cols=4)
    apply_table_style(final_table, STD_WIDTHS)
    
    for row_idx, (lab, key_list, opts) in enumerate([
        ('待 遇', 'compensation', ['月薪', '時薪', '獎金(津貼)', '無']),
        ('來 源', 'source', ['廠商申請', '老師推薦', '學生申請', '其它'])
    ]):
        cells = final_table.rows[row_idx].cells
        set_cell_format(cells[0], lab, distribute=True)
        cells[1].merge(cells[2]); cells[1].merge(cells[3])
        
        selected = data.get(key_list, [])
        text = ''.join([f'{"☑" if o in selected else "☐"} {o}   ' for o in opts])
        if lab == '來 源' and '其它' in selected and data.get('source_other_text'):
            text += f'（{data.get("source_other_text")}）'
        set_cell_format(cells[1], text, alignment=WD_ALIGN_PARAGRAPH.LEFT)
 
    return doc
# =========================================================
# 📥 下載公司上傳範本
# =========================================================
@company_bp.route('/download_company_template', methods=['GET'])
def download_company_template():
    try:
        template_file_name = "114學年實習單位基本資料表.docx"
        backend_dir = current_app.root_path
        project_root = os.path.dirname(backend_dir)
        file_path = os.path.join(project_root, 'frontend', 'static', 'examples', template_file_name)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "message": "找不到範本檔案"}), 404

        return send_file(
            file_path,
            as_attachment=True,
            download_name=template_file_name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": "下載範本失敗"}), 500


# =========================================================
# 🔢 獲取下一個編號序號
# =========================================================
@company_bp.route('/api/get_next_serial_number', methods=['GET'])
def get_next_serial_number():
    """根據民國年份獲取下一個序號"""
    conn = None
    cursor = None
    try:
        year = request.args.get('year', '').strip()
        if not year or len(year) != 3:
            # 如果沒有提供年份，使用當前民國年份
            now = datetime.now()
            year = str(now.year - 1911).zfill(3)
        
        conn = get_db()
        cursor = conn.cursor()
        
        # 計算該年份的起始和結束日期（西元年）
        roc_year = int(year)
        gregorian_year_start = roc_year + 1911
        gregorian_year_end = gregorian_year_start + 1
        
        # 查詢該年份內提交的公司數量
        cursor.execute("""
            SELECT COUNT(*) 
            FROM internship_companies 
            WHERE submitted_at >= %s 
            AND submitted_at < %s
        """, (
            datetime(gregorian_year_start, 1, 1),
            datetime(gregorian_year_end, 1, 1)
        ))
        
        count = cursor.fetchone()[0]
        
        # 下一個序號 = 該年份的公司數量 + 1
        next_sequence = count + 1
        
        return jsonify({
            "success": True,
            "year": year,
            "next_sequence": next_sequence,
            "serial_number": year + str(next_sequence).zfill(3)
        })
        
    except Exception as e:
        traceback.print_exc()
        # 如果出錯，返回預設值 001
        now = datetime.now()
        year = str(now.year - 1911).zfill(3)
        return jsonify({
            "success": True,
            "year": year,
            "next_sequence": 1,
            "serial_number": year + "001"
        })
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# 📤 上傳公司資料（網頁填表，自動生成 Word 檔）
# =========================================================
@company_bp.route('/api/upload_company', methods=['POST'])
def upload_company():
    conn = None
    cursor = None
    file_path = None

    try:
        if 'user_id' not in session:
            return jsonify({"success": False, "message": "請先登入"}), 403

        role = session.get('role')
        if role not in ['teacher', 'director', 'ta', 'vendor']:
           return jsonify({"success": False, "message": "無權限操作此功能"}), 403

        user_id = session['user_id']
        upload_dir = ensure_upload_folder()

        # 判斷是 JSON 資料（新方式）還是表單資料（舊方式，保留向後兼容）
        if request.is_json:
            data = request.get_json()
            company_name = data.get("company_name", "").strip()
            jobs_data = data.get("jobs", [])
        else:
            # 舊方式：表單上傳（向後兼容）
            company_name = request.form.get("company_name", "").strip()
            jobs_data = []
            job_index = 0
            while True:
                job_title = request.form.get(f"job[{job_index}][title]", "").strip()
                slots_str = request.form.get(f"job[{job_index}][slots]", "0").strip()
                if not job_title:
                    break
                try:
                    slots = int(slots_str)
                    if slots <= 0:
                        raise ValueError
                except ValueError:
                    return jsonify({"success": False, "message": f"職缺 #{job_index+1} 名額必須是正整數"}), 400
                jobs_data.append({"title": job_title, "slots": slots})
                job_index += 1

        if not company_name:
            return jsonify({"success": False, "message": "公司名稱為必填欄位"}), 400

        if not jobs_data:
            return jsonify({"success": False, "message": "請至少新增一個職缺"}), 400

        # 如果是 JSON 資料，生成 Word 檔
        if request.is_json:
            # 生成 Word 檔
            doc = generate_company_word_document(data)
            
            # 儲存 Word 檔
            safe_name = secure_filename(f"{company_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx")
            save_path = os.path.join(upload_dir, safe_name)
            doc.save(save_path)
            
            # 驗證文件是否成功保存
            if not os.path.exists(save_path):
                return jsonify({"success": False, "message": "檔案保存失敗，請稍後再試"}), 500
            
            file_path = os.path.join(UPLOAD_FOLDER, safe_name)
        else:
            # 舊方式：處理上傳的 Word 檔案
            file = request.files.get("company_doc")
            if file and file.filename and allowed_file(file.filename):
                safe_name = secure_filename(f"{company_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{file.filename}")
                save_path = os.path.join(upload_dir, safe_name)
                file.save(save_path)
                file_path = os.path.join(UPLOAD_FOLDER, safe_name)
            else:
                return jsonify({"success": False, "message": "請上傳有效的 Word 檔案 (.doc 或 .docx)"}), 400

        # 寫入資料庫
        conn = get_db()
        cursor = conn.cursor()

        # 如果是科助，自動填入 advisor_user_id 和 reviewed_by_user_id，並設為已核准狀態
        if role == 'ta':
            advisor_user_id = user_id
            reviewed_by_user_id = user_id
            status = 'approved'
            reviewed_at = datetime.now()
        elif role == 'vendor':
            # 廠商上傳：根據廠商的 teacher_name 找到對應的指導老師
            cursor.execute("SELECT teacher_name FROM users WHERE id = %s", (user_id,))
            vendor_row = cursor.fetchone()
            advisor_user_id = None
            if vendor_row and vendor_row[0]:
                teacher_name = vendor_row[0].strip()
                if teacher_name:
                    cursor.execute("SELECT id FROM users WHERE name = %s AND role IN ('teacher', 'director')", (teacher_name,))
                    teacher_row = cursor.fetchone()
                    if teacher_row:
                        advisor_user_id = teacher_row[0]
            reviewed_by_user_id = None
            status = 'pending'
            reviewed_at = None
        else:
            # 如果是老師或主任，預設上傳教師為指導老師
            updated_vendor_username = None  # 初始化變數
            if role in ['teacher', 'director']:
                advisor_user_id = user_id
                
                # 取得上傳者的名字（用於更新廠商的 teacher_name）
                cursor.execute("SELECT name FROM users WHERE id = %s", (user_id,))
                teacher_info = cursor.fetchone()
                teacher_name = teacher_info[0] if teacher_info and teacher_info[0] else None
                
                # 檢查公司名稱是否匹配廠商對應的公司名稱
                # 如果匹配，則更新該廠商的 teacher_name 為該指導老師的名字
                vendor_company_map = {
                    'vendor': '人人人',
                    'vendora': '嘻嘻嘻'
                }
                
                # 檢查公司名稱是否在 vendor_company_map 的值中
                matched_vendor_username = None
                for vendor_username, mapped_company_name in vendor_company_map.items():
                    if company_name == mapped_company_name:
                        matched_vendor_username = vendor_username
                        break
                
                # 如果找到匹配的廠商，更新該廠商的 teacher_name 為該指導老師的名字
                if matched_vendor_username and teacher_name:
                    cursor.execute("""
                        UPDATE users 
                        SET teacher_name = %s 
                        WHERE username = %s AND role = 'vendor'
                    """, (teacher_name, matched_vendor_username))
                    # 記錄更新的廠商資訊（用於後續的成功訊息）
                    updated_vendor_username = matched_vendor_username
            else:
                advisor_user_id = None
            reviewed_by_user_id = None
            status = 'pending'
            reviewed_at = None

        # 準備公司資料
        if request.is_json:
            company_description = data.get("company_intro", "（詳見附檔）")
            company_location = data.get("address", "")
            contact_person = data.get("contact_person", "")
            contact_title = data.get("contact_title", "")
            contact_email = data.get("email", "")
            contact_phone = data.get("contact_phone", "")
        else:
            company_description = "（詳見附檔）"
            company_location = ""
            contact_person = ""
            contact_title = ""
            contact_email = ""
            contact_phone = ""
        
        cursor.execute("""
            INSERT INTO internship_companies 
            (company_name, uploaded_by_user_id, advisor_user_id, reviewed_by_user_id, status, submitted_at, reviewed_at, company_doc_path, 
             description, location, contact_person, contact_title, contact_email, contact_phone)
            VALUES (%s, %s, %s, %s, %s, NOW(), %s, %s, %s, %s, %s, %s, %s, %s)
        """, (company_name, user_id, advisor_user_id, reviewed_by_user_id, status, reviewed_at, file_path,
              company_description, company_location, contact_person, contact_title, contact_email, contact_phone))
        company_id = cursor.lastrowid

        # 插入職缺
        job_records = []
        for j in jobs_data:
            job_description = j.get("description", "（詳見附檔）")
            job_records.append((
                company_id,
                j.get("title", ""),
                j.get("slots", 1),
                job_description,
                "",
                "",
                "",
                True
            ))
        cursor.executemany("""
            INSERT INTO internship_jobs 
            (company_id, title, slots, description, period, work_time, remark, is_active)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
        """, job_records)

        conn.commit()

        job_count = len(jobs_data)
        
        # 根據角色顯示不同的成功訊息
        if role == 'ta':
            message = f"公司 '{company_name}' ({job_count} 個職缺) 上傳成功，已自動核准。"
        elif role == 'vendor':
            message = f"公司 '{company_name}' ({job_count} 個職缺) 上傳成功，資料已標記為「待科助開放」。"
        else:
            # 老師或主任上傳
            message = f"公司 '{company_name}' ({job_count} 個職缺) 上傳成功，等待審核。"
            # 如果匹配到廠商並更新了 teacher_name，在訊息中提示
            if updated_vendor_username:
                message += f" 已自動更新廠商 '{updated_vendor_username}' 的指導老師關聯。"

        response_data = {
            "success": True,
            "message": message,
            "company_id": company_id
        }
        
        # 如果是新方式（JSON），提供下載連結
        if request.is_json and file_path:
            response_data["download_url"] = f"/api/download_company_file/{company_id}"

        return jsonify(response_data)

    except Exception as e:
        traceback.print_exc()
        # 如果發生錯誤，刪除剛剛儲存的檔案
        if file_path:
            project_root = os.path.dirname(current_app.root_path)
            abs_path = os.path.join(project_root, file_path)
            if os.path.exists(abs_path):
                os.remove(abs_path)
        return jsonify({"success": False, "message": f"伺服器錯誤: {e}"}), 500

    finally:
        if cursor: cursor.close()
        if conn: conn.close()


# =========================================================
# 📜 查詢使用者上傳紀錄
# =========================================================
@company_bp.route('/api/get_my_companies', methods=['GET'])
def get_my_companies():
    conn = None
    cursor = None
    try:
        if 'user_id' not in session:
            return jsonify({"success": False, "message": "請先登入"}), 403

        user_id = session['user_id']
        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("""
            SELECT 
                ic.id,
                ic.company_name,
                ic.status,
                ic.company_doc_path AS filepath,
                ic.submitted_at AS upload_time,
                u.role AS uploader_role
            FROM internship_companies ic
            JOIN users u ON ic.uploaded_by_user_id = u.id
            WHERE ic.uploaded_by_user_id = %s
            ORDER BY ic.submitted_at DESC
        """, (user_id,))
        records = cursor.fetchall()

        # === 🕒 加上台灣時區轉換 ===
        from datetime import datetime, timezone, timedelta
        taiwan_tz = timezone(timedelta(hours=8))

        for r in records:
            if isinstance(r.get("upload_time"), datetime):
                # 將 UTC 轉為台灣時間
                r["upload_time"] = r["upload_time"].astimezone(taiwan_tz).strftime("%Y-%m-%d %H:%M")
            else:
                r["upload_time"] = "-"

            r["filename"] = os.path.basename(r["filepath"]) if r["filepath"] else None

        return jsonify({"success": True, "companies": records})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": "載入上傳紀錄失敗"}), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()


# =========================================================
# 📂 下載上傳的公司檔案
# =========================================================
@company_bp.route('/api/download_company_file/<int:file_id>', methods=['GET'])
def download_company_file(file_id):
    conn = None
    cursor = None
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT company_doc_path FROM internship_companies WHERE id=%s", (file_id,))
        record = cursor.fetchone()
        
        if not record:
            from flask import render_template_string
            return render_template_string('''
                <html><body>
                <h2>錯誤：找不到此公司紀錄</h2>
                <p>公司 ID: {{ file_id }}</p>
                <a href="javascript:history.back()">返回</a>
                </body></html>
            ''', file_id=file_id), 404
        
        if not record.get("company_doc_path"):
            from flask import render_template_string
            return render_template_string('''
                <html><body>
                <h2>錯誤：此公司沒有上傳檔案</h2>
                <p>公司 ID: {{ file_id }}</p>
                <a href="javascript:history.back()">返回</a>
                </body></html>
            ''', file_id=file_id), 404

        project_root = os.path.dirname(current_app.root_path)
        abs_path = os.path.join(project_root, record["company_doc_path"])
        
        if not os.path.exists(abs_path):
            from flask import render_template_string
            return render_template_string('''
                <html><body>
                <h2>錯誤：檔案不存在</h2>
                <p>檔案路徑: {{ file_path }}</p>
                <p>公司 ID: {{ file_id }}</p>
                <a href="javascript:history.back()">返回</a>
                </body></html>
            ''', file_path=record["company_doc_path"], file_id=file_id), 404

        filename = os.path.basename(abs_path)
        return send_file(abs_path, as_attachment=True, download_name=filename)
    except Exception as e:
        traceback.print_exc()
        from flask import render_template_string
        return render_template_string('''
            <html><body>
            <h2>錯誤：下載失敗</h2>
            <p>錯誤訊息: {{ error }}</p>
            <a href="javascript:history.back()">返回</a>
            </body></html>
        ''', error=str(e)), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()


# =========================================================
# 🗑️ 刪除公司上傳紀錄
# =========================================================
@company_bp.route('/api/delete_company/<int:company_id>', methods=['DELETE'])
def delete_company(company_id):
    conn = None
    cursor = None
    try:
        if 'user_id' not in session:
            return jsonify({"success": False, "message": "請先登入"}), 403

        user_id = session['user_id']
        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        # 先查資料，確認是否為本人上傳
        cursor.execute("""
            SELECT company_doc_path FROM internship_companies 
            WHERE id=%s AND uploaded_by_user_id=%s
        """, (company_id, user_id))
        record = cursor.fetchone()

        if not record:
            return jsonify({"success": False, "message": "找不到該公司資料或您無權限刪除"}), 404

        # 刪除檔案（如果存在）
        if record["company_doc_path"]:
            project_root = os.path.dirname(current_app.root_path)
            abs_path = os.path.join(project_root, record["company_doc_path"])
            if os.path.exists(abs_path):
                os.remove(abs_path)

        # 先刪除志願序中引用到該公司/職缺的資料以免觸發 FK
        cursor.execute("SELECT id FROM internship_jobs WHERE company_id=%s", (company_id,))
        job_rows = cursor.fetchall() or []
        job_ids = [row["id"] for row in job_rows]

        # 刪除指定公司下的志願序（包含未指定職缺與指定職缺）
        cursor.execute("DELETE FROM student_preferences WHERE company_id=%s", (company_id,))
        if job_ids:
            placeholders = ", ".join(["%s"] * len(job_ids))
            cursor.execute(f"DELETE FROM student_preferences WHERE job_id IN ({placeholders})", tuple(job_ids))

        # 刪除公司開放設定，避免 company_openings → internship_companies FK 擋住
        cursor.execute("DELETE FROM company_openings WHERE company_id=%s", (company_id,))

        # 刪除相關職缺資料
        cursor.execute("DELETE FROM internship_jobs WHERE company_id=%s", (company_id,))

        # 刪除公司主資料
        cursor.execute("DELETE FROM internship_companies WHERE id=%s", (company_id,))
        conn.commit()

        return jsonify({"success": True, "message": "公司資料已刪除。"})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"刪除失敗: {e}"}), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# API - 取得所有已開放職缺的公司清單
# =========================================================
@company_bp.route("/api/get_companies_for_resume_delivery", methods=["GET"])
def get_companies_for_resume_delivery():
    # 必須登入
    if "user_id" not in session or session.get("role") != "student":
        return jsonify({"success": False, "message": "未授權"}), 403

    student_id = session["user_id"]

    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 1️⃣ 檢查學生是否有填寫志願序
        cursor.execute("""
            SELECT DISTINCT company_id, job_id
            FROM student_preferences
            WHERE student_id = %s
        """, (student_id,))
        pref_records = cursor.fetchall()

        use_preferences = len(pref_records) > 0

        # 2️⃣ 根據是否有志願序，決定 SQL 條件
        if use_preferences:
            # 如果有志願序，只顯示學生在志願序中選擇的公司和職缺
            company_ids = list(set([r["company_id"] for r in pref_records if r["company_id"]]))
            job_ids = [r["job_id"] for r in pref_records if r["job_id"] is not None]
            
            format_strings = ",".join(["%s"] * len(company_ids))
            
            if job_ids:
                # 如果有明確的 job_id，只顯示這些職缺
                job_format_strings = ",".join(["%s"] * len(job_ids))
                cursor.execute(f"""
                    SELECT
                        c.id AS company_id,
                        c.company_name,
                        j.id AS job_id,
                        j.title AS job_title
                    FROM internship_companies c
                    JOIN internship_jobs j ON j.company_id = c.id
                    WHERE c.status = 'approved'
                      AND j.is_active = TRUE
                      AND c.id IN ({format_strings})
                      AND j.id IN ({job_format_strings})
                    ORDER BY c.company_name, j.id
                """, tuple(company_ids + job_ids))
            else:
                # 如果志願序中沒有 job_id（舊資料），顯示該公司下的所有職缺
                cursor.execute(f"""
                    SELECT
                        c.id AS company_id,
                        c.company_name,
                        j.id AS job_id,
                        j.title AS job_title
                    FROM internship_companies c
                    JOIN internship_jobs j ON j.company_id = c.id
                    WHERE c.status = 'approved'
                      AND j.is_active = TRUE
                      AND c.id IN ({format_strings})
                    ORDER BY c.company_name, j.id
                """, tuple(company_ids))
        else:
            # 尚未填志願序 → 顯示全部
            cursor.execute("""
                SELECT
                    c.id AS company_id,
                    c.company_name,
                    j.id AS job_id,
                    j.title AS job_title
                FROM internship_companies c
                JOIN internship_jobs j ON j.company_id = c.id
                WHERE c.status = 'approved'
                  AND j.is_active = TRUE
                ORDER BY c.company_name, j.id
            """)

        rows = cursor.fetchall()

        # 3️⃣ 整理成 company -> jobs
        companies = {}
        for r in rows:
            cid = r["company_id"]
            if cid not in companies:
                companies[cid] = {
                    "company_id": cid,
                    "company_name": r["company_name"],
                    "jobs": []
                }
            companies[cid]["jobs"].append({
                "job_id": r["job_id"],
                "job_title": r["job_title"]
            })

        return jsonify({
            "success": True,
            "use_preferences": use_preferences, 
            "companies": list(companies.values())
        })

    finally:
        cursor.close()
        conn.close()

# =========================================================
# API - 取得已審核公司（歷史紀錄）
# =========================================================
@company_bp.route("/api/get_reviewed_companies", methods=["GET"])
def api_get_reviewed_companies():
    conn = None
    cursor = None
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        # 取得當前學期代碼
        current_semester_code = get_current_semester_code(cursor)

        # 如果沒有設定當前學期，仍然可以顯示公司列表，但無法顯示開放狀態
        if current_semester_code:
            cursor.execute("""
                SELECT 
                    ic.id,
                    u.name AS upload_teacher_name,
                    COALESCE(advisor.name, 
                        CASE 
                            WHEN ic.advisor_user_id IS NULL AND u.role IN ('teacher', 'director') THEN u.name 
                            ELSE NULL 
                        END
                    ) AS advisor_teacher_name,
                    COALESCE(ic.advisor_user_id, 
                        CASE 
                            WHEN u.role IN ('teacher', 'director') THEN ic.uploaded_by_user_id 
                            ELSE NULL 
                        END
                    ) AS advisor_user_id,
                    ic.company_name, 
                    ic.status,
                    ic.submitted_at AS upload_time,
                    ic.reviewed_at,
                    COALESCE(co.is_open, FALSE) AS is_open_current_semester
                FROM internship_companies ic
                LEFT JOIN users u ON ic.uploaded_by_user_id = u.id
                LEFT JOIN users advisor ON ic.advisor_user_id = advisor.id
                LEFT JOIN company_openings co ON ic.id = co.company_id 
                    AND co.semester = %s
                WHERE ic.status = 'approved'
                ORDER BY 
                    CASE WHEN ic.reviewed_at IS NULL THEN 1 ELSE 0 END,
                    ic.reviewed_at DESC,
                    ic.submitted_at DESC
            """, (current_semester_code,))
        else:
            cursor.execute("""
                SELECT 
                    ic.id,
                    u.name AS upload_teacher_name,
                    COALESCE(advisor.name, 
                        CASE 
                            WHEN ic.advisor_user_id IS NULL AND u.role IN ('teacher', 'director') THEN u.name 
                            ELSE NULL 
                        END
                    ) AS advisor_teacher_name,
                    COALESCE(ic.advisor_user_id, 
                        CASE 
                            WHEN u.role IN ('teacher', 'director') THEN ic.uploaded_by_user_id 
                            ELSE NULL 
                        END
                    ) AS advisor_user_id,
                    ic.company_name, 
                    ic.status,
                    ic.submitted_at AS upload_time,
                    ic.reviewed_at,
                    FALSE AS is_open_current_semester
                FROM internship_companies ic
                LEFT JOIN users u ON ic.uploaded_by_user_id = u.id
                LEFT JOIN users advisor ON ic.advisor_user_id = advisor.id
                WHERE ic.status = 'approved'
                ORDER BY 
                    CASE WHEN ic.reviewed_at IS NULL THEN 1 ELSE 0 END,
                    ic.reviewed_at DESC,
                    ic.submitted_at DESC
            """)

        companies = cursor.fetchall()

        # 取得各公司的職缺列表（僅抓啟用中的職缺）
        job_map = {}
        if companies:
            company_ids = [c["id"] for c in companies]
            placeholders = ",".join(["%s"] * len(company_ids))
            cursor.execute(
                f"""
                SELECT company_id, title
                FROM internship_jobs
                WHERE company_id IN ({placeholders})
                  AND is_active = TRUE
                """,
                company_ids,
            )
            job_rows = cursor.fetchall()
            for job in job_rows:
                cid = job.get("company_id")
                title = job.get("title") or ""
                job_map.setdefault(cid, []).append(title)
        
        # 調試：記錄返回的公司狀態分布
        status_count = {}
        for company in companies:
            status = company.get('status', 'unknown')
            status_count[status] = status_count.get(status, 0) + 1
        print(f"📊 已審核公司查詢結果: 總數={len(companies)}, 狀態分布={status_count}")
        
        # 格式化時間
        from datetime import timezone, timedelta
        taiwan_tz = timezone(timedelta(hours=8))
        
        for company in companies:
            if company.get('upload_time') and isinstance(company['upload_time'], datetime):
                company['upload_time'] = company['upload_time'].astimezone(taiwan_tz).strftime("%Y-%m-%d %H:%M")
            else:
                company['upload_time'] = "-"
            
            if company.get('reviewed_at') and isinstance(company['reviewed_at'], datetime):
                company['reviewed_at'] = company['reviewed_at'].astimezone(taiwan_tz).strftime("%Y-%m-%d %H:%M")
            else:
                company['reviewed_at'] = "-"
            
            # 確保 is_open_current_semester 是布林值
            company['is_open_current_semester'] = bool(company.get('is_open_current_semester', False))
            # 附加職缺清單（前端顯示用）
            company['jobs'] = job_map.get(company['id'], [])
        
        return jsonify({"success": True, "companies": companies, "current_semester": current_semester_code})

    except Exception:
        print("❌ 取得已審核公司錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# 🔎 取得公司詳細資料 (包含職缺)
# =========================================================
@company_bp.route('/api/get_company_detail', methods=['GET'])
def get_company_detail():
    conn = None
    cursor = None
    try:
        if 'user_id' not in session:
            return jsonify({"success": False, "message": "請先登入"}), 403

        company_id = request.args.get('company_id', type=int)
        if not company_id:
            return jsonify({"success": False, "message": "缺少 company_id"}), 400

        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        # 查詢公司主資料
        cursor.execute("""
            SELECT 
                ic.id, ic.company_name, ic.status, ic.description AS company_intro, 
                ic.location AS company_address, ic.contact_person AS contact_name, 
                ic.contact_title, ic.contact_email, ic.contact_phone, 
                ic.reject_reason, ic.submitted_at, ic.reviewed_at, 
                u.name AS upload_teacher_name
            FROM internship_companies ic
            JOIN users u ON ic.uploaded_by_user_id = u.id
            WHERE ic.id = %s
        """, (company_id,))
        company = cursor.fetchone()

        if not company:
            return jsonify({"success": False, "message": "找不到公司資料"}), 404

        # 查詢職缺資料
        cursor.execute("""
            SELECT 
                title AS internship_unit, 
                description AS internship_content, 
                period AS internship_period, 
                work_time AS internship_time, 
                slots AS internship_quota, 
                remark, salary
            FROM internship_jobs
            WHERE company_id = %s
            AND is_active = TRUE
        """, (company_id,))
        jobs = cursor.fetchall()
        company['internship_jobs'] = jobs

        return jsonify({"success": True, "company": company})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"載入詳細資料失敗: {e}"}), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()


# =========================================================
# 📚 實習 QA - 取得所有問答
# =========================================================
@company_bp.route('/api/qa/list', methods=['GET'])
def qa_list():
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("""
            SELECT id, question, answer 
            FROM internship_qa
            ORDER BY sort_order ASC, id DESC
        """)
        data = cursor.fetchall()

        return jsonify({"success": True, "data": data})

    except Exception:
        import traceback
        print("❌ QA 列表錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500

    finally:
        if cursor: cursor.close()
        if conn: conn.close()


# =========================================================
# ➕ 實習 QA - 新增
# =========================================================
@company_bp.route('/api/qa/add', methods=['POST'])
def qa_add():
    data = request.json

    question = data.get("question", "").strip()
    answer   = data.get("answer", "").strip()
    sort     = data.get("sort_order", 0)

    if not question or not answer:
        return jsonify({"success": False, "message": "問題與答案不得為空"}), 400

    try:
        conn = get_db()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO internship_qa (question, answer, sort_order)
            VALUES (%s, %s, %s)
        """, (question, answer, sort))

        conn.commit()
        return jsonify({"success": True, "message": "新增成功"})

    except Exception:
        import traceback
        print("❌ QA 新增錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500

    finally:
        if cursor: cursor.close()
        if conn: conn.close()


# =========================================================
# ✏️ 實習 QA - 更新
# =========================================================
@company_bp.route('/api/qa/update/<int:qa_id>', methods=['PUT'])
def qa_update(qa_id):
    data = request.json

    question = data.get("question", "").strip()
    answer   = data.get("answer", "").strip()
    sort     = data.get("sort_order")

    if not question or not answer:
        return jsonify({"success": False, "message": "問題與答案不得為空"}), 400

    try:
        sort = int(sort) if str(sort).isdigit() else 0

        conn = get_db()
        cursor = conn.cursor()

        cursor.execute("""
            UPDATE internship_qa
            SET question=%s, answer=%s, sort_order=%s
            WHERE id=%s
        """, (question, answer, sort, qa_id))

        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"success": False, "message": "找不到該 QA"}), 404

        return jsonify({"success": True, "message": "更新成功"})

    except Exception:
        import traceback
        print("❌ QA 更新錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500

    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# 🗑️ 實習 QA - 刪除
# =========================================================
@company_bp.route('/api/qa/delete/<int:qa_id>', methods=['DELETE'])
def qa_delete(qa_id):
    try:
        conn = get_db()
        cursor = conn.cursor()

        cursor.execute("DELETE FROM internship_qa WHERE id=%s", (qa_id,))
        conn.commit()

        return jsonify({"success": True, "message": "刪除成功"})

    except Exception:
        import traceback
        print("❌ QA 刪除錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500

    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# API - 審核公司
# =========================================================
@company_bp.route("/api/approve_company", methods=["POST"])
def api_approve_company():
    data = request.get_json()
    company_id = data.get("company_id")
    status = data.get("status")

    if not company_id or status not in ['approved', 'rejected']:
        return jsonify({"success": False, "message": "參數錯誤"}), 400

    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute("SELECT company_name, status, advisor_user_id FROM internship_companies WHERE id = %s", (company_id,))
        company_row = cursor.fetchone()

        if not company_row:
            return jsonify({"success": False, "message": "查無此公司"}), 404

        company_name, current_status, advisor_user_id = company_row
        if current_status != 'pending':
            return jsonify({"success": False, "message": f"公司已被審核過（目前狀態為 {current_status}）"}), 400

        # 取得審核者的 user_id
        reviewer_id = session.get('user_id') if 'user_id' in session else None

        cursor.execute("""
            UPDATE internship_companies
            SET status = %s, reviewed_at = %s, reviewed_by_user_id = %s
            WHERE id = %s
        """, (status, datetime.now(), reviewer_id, company_id))
        
        # 如果核准且公司有指導老師，檢查是否需要更新廠商的 teacher_name
        updated_vendor_username = None
        if status == 'approved' and advisor_user_id:
            # 取得指導老師的名字
            cursor.execute("SELECT name FROM users WHERE id = %s", (advisor_user_id,))
            teacher_info = cursor.fetchone()
            teacher_name = teacher_info[0] if teacher_info and teacher_info[0] else None
            
            if teacher_name:
                # 檢查公司名稱是否匹配廠商對應的公司名稱
                vendor_company_map = {
                    'vendor': '人人人',
                    'vendora': '嘻嘻嘻'
                }
                
                # 檢查公司名稱是否在 vendor_company_map 的值中
                matched_vendor_username = None
                for vendor_username, mapped_company_name in vendor_company_map.items():
                    if company_name == mapped_company_name:
                        matched_vendor_username = vendor_username
                        break
                
                # 如果找到匹配的廠商，更新該廠商的 teacher_name 為該指導老師的名字
                if matched_vendor_username:
                    cursor.execute("""
                        UPDATE users 
                        SET teacher_name = %s 
                        WHERE username = %s AND role = 'vendor'
                    """, (teacher_name, matched_vendor_username))
                    updated_vendor_username = matched_vendor_username
        
        conn.commit()

        action_text = '核准' if status == 'approved' else '拒絕'
        message = f"公司「{company_name}」已{action_text}"
        # 如果更新了廠商的 teacher_name，在訊息中提示
        if updated_vendor_username:
            message += f" 已自動更新廠商 '{updated_vendor_username}' 的指導老師關聯。"
        return jsonify({"success": True, "message": message})

    except Exception:
        print("❌ 審核公司錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500

    finally:
        cursor.close()
        conn.close()

# =========================================================
# API - 設定公司本學期開放狀態
# =========================================================
@company_bp.route("/api/set_company_open_status", methods=["POST"])
def api_set_company_open_status():
    """設定公司在本學期是否開放"""
    data = request.get_json()
    company_id = data.get("company_id")
    is_open = data.get("is_open", False)

    if company_id is None:
        return jsonify({"success": False, "message": "缺少 company_id"}), 400

    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        
        # 取得當前學期代碼
        current_semester_code = get_current_semester_code(cursor)
        if not current_semester_code:
            return jsonify({"success": False, "message": "目前沒有設定當前學期"}), 400

        # 檢查公司是否存在且已審核通過
        cursor.execute("SELECT id, company_name, status FROM internship_companies WHERE id = %s", (company_id,))
        company = cursor.fetchone()
        
        if not company:
            return jsonify({"success": False, "message": "找不到該公司"}), 404
        
        if company['status'] != 'approved':
            return jsonify({"success": False, "message": "只有已審核通過的公司才能設定開放狀態"}), 400

        # 檢查是否已存在該公司該學期的記錄
        cursor.execute("""
            SELECT id FROM company_openings 
            WHERE company_id = %s AND semester = %s
        """, (company_id, current_semester_code))
        existing = cursor.fetchone()

        if existing:
            # 更新現有記錄
            cursor.execute("""
                UPDATE company_openings 
                SET is_open = %s, opened_at = %s
                WHERE company_id = %s AND semester = %s
            """, (is_open, datetime.now(), company_id, current_semester_code))
        else:
            # 建立新記錄
            cursor.execute("""
                INSERT INTO company_openings (company_id, semester, is_open, opened_at)
                VALUES (%s, %s, %s, %s)
            """, (company_id, current_semester_code, is_open, datetime.now()))

        conn.commit()
        
        status_text = '開放' if is_open else '關閉'
        return jsonify({
            "success": True, 
            "message": f"公司「{company['company_name']}」已{status_text}",
            "is_open": bool(is_open)
        })

    except Exception as e:
        print("❌ 設定公司開放狀態錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": f"伺服器錯誤: {str(e)}"}), 500

    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# 🖥️ 上傳公司頁面
# =========================================================
@company_bp.route('/upload_company', methods=['GET'])
def upload_company_form_page():
    # 傳遞使用者角色資訊給前端，用於顯示提示
    user_role = session.get('role', '')
    return render_template('company/upload_company.html', user_role=user_role)

# =========================================================
# API - 取得所有指導老師
# =========================================================
@company_bp.route("/api/get_all_teachers", methods=["GET"])
def api_get_all_teachers():
    """取得所有指導老師（teacher 和 director 角色）"""
    conn = None
    cursor = None
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("""
            SELECT id, name
            FROM users
            WHERE role IN ('teacher', 'director')
            ORDER BY name ASC
        """)
        teachers = cursor.fetchall()
        
        return jsonify({"success": True, "teachers": teachers})
    except Exception:
        print("❌ 取得指導老師列表錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# API - 更新公司指導老師
# =========================================================
@company_bp.route("/api/update_company_advisor", methods=["POST"])
def api_update_company_advisor():
    """更新公司的指導老師"""
    data = request.get_json()
    company_id = data.get("company_id")
    advisor_user_id = data.get("advisor_user_id")  # 可以是 None
    
    if not company_id:
        return jsonify({"success": False, "message": "缺少 company_id"}), 400
    
    conn = None
    cursor = None
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        
        # 檢查公司是否存在
        cursor.execute("SELECT id, company_name FROM internship_companies WHERE id = %s", (company_id,))
        company = cursor.fetchone()
        if not company:
            return jsonify({"success": False, "message": "找不到該公司"}), 404
        
        # 如果提供了 advisor_user_id，驗證該用戶是老師或主任
        if advisor_user_id:
            cursor.execute("SELECT id, name, role FROM users WHERE id = %s AND role IN ('teacher', 'director')", (advisor_user_id,))
            teacher = cursor.fetchone()
            if not teacher:
                return jsonify({"success": False, "message": "指定的用戶不是有效的指導老師"}), 400
        
        # 更新指導老師
        cursor.execute("""
            UPDATE internship_companies
            SET advisor_user_id = %s
            WHERE id = %s
        """, (advisor_user_id, company_id))
        
        # 取得更新後的指導老師名稱
        advisor_name = None
        if advisor_user_id:
            cursor.execute("SELECT name FROM users WHERE id = %s", (advisor_user_id,))
            advisor = cursor.fetchone()
            if advisor:
                advisor_name = advisor['name']
        
        # 更新所有相關廠商的 teacher_name
        updated_vendor_count = 0
        # 取得公司的 uploaded_by_user_id 和 contact_email
        cursor.execute("""
            SELECT uploaded_by_user_id, contact_email 
            FROM internship_companies 
            WHERE id = %s
        """, (company_id,))
        company_info = cursor.fetchone()
        
        vendor_ids_to_update = []
        
        # 1. 如果上傳者是廠商，更新該廠商的 teacher_name
        if company_info and company_info.get('uploaded_by_user_id'):
            cursor.execute("""
                SELECT id FROM users 
                WHERE id = %s AND role = 'vendor'
            """, (company_info['uploaded_by_user_id'],))
            vendor = cursor.fetchone()
            if vendor:
                vendor_ids_to_update.append(vendor['id'])
        
        # 2. 如果公司有 contact_email，查找所有匹配該 email 的廠商
        if company_info and company_info.get('contact_email'):
            cursor.execute("""
                SELECT id FROM users 
                WHERE email = %s AND role = 'vendor'
            """, (company_info['contact_email'],))
            vendors_by_email = cursor.fetchall()
            for vendor in vendors_by_email:
                if vendor['id'] not in vendor_ids_to_update:
                    vendor_ids_to_update.append(vendor['id'])
        
        # 3. 更新所有找到的廠商的 teacher_name
        # 如果 advisor_user_id 為 None，則清除 teacher_name（設為 NULL）
        # 如果 advisor_user_id 有值，則設定為指導老師的名稱
        teacher_name_value = advisor_name if advisor_user_id and advisor_name else None
        for vendor_id in vendor_ids_to_update:
            cursor.execute("""
                UPDATE users 
                SET teacher_name = %s 
                WHERE id = %s AND role = 'vendor'
            """, (teacher_name_value, vendor_id))
            updated_vendor_count += 1
        
        conn.commit()
        
        message = f"公司「{company['company_name']}」的指導老師已更新"
        # 如果更新了廠商的 teacher_name，在訊息中提示
        if updated_vendor_count > 0:
            message += f" 已自動更新 {updated_vendor_count} 個相關廠商的指導老師關聯。"
        
        return jsonify({
            "success": True,
            "message": message,
            "advisor_name": advisor_name,
            "updated_vendor_count": updated_vendor_count
        })
    except Exception:
        print("❌ 更新公司指導老師錯誤：", traceback.format_exc())
        conn.rollback()
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# 📥 導出公司審核數據
# =========================================================
@company_bp.route("/api/export_company_reviews", methods=["GET"])
def api_export_company_reviews():
    """導出公司審核數據為SQL文件"""
    conn = None
    cursor = None
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        
        # 查詢所有已審核的公司
        cursor.execute("""
            SELECT 
                ic.id,
                ic.company_name,
                ic.status,
                ic.reviewed_at,
                ic.reviewed_by_user_id,
                ic.advisor_user_id
            FROM internship_companies ic
            WHERE ic.status IN ('approved', 'rejected')
            ORDER BY ic.id
        """)
        companies = cursor.fetchall()
        
        # 查詢公司開放狀態
        cursor.execute("""
            SELECT 
                co.company_id,
                co.semester,
                co.is_open,
                co.opened_at
            FROM company_openings co
            ORDER BY co.company_id, co.semester
        """)
        openings = cursor.fetchall()
        openings_dict = {}
        for opening in openings:
            company_id = opening['company_id']
            if company_id not in openings_dict:
                openings_dict[company_id] = []
            openings_dict[company_id].append(opening)
        
        # 生成SQL內容
        sql_lines = []
        sql_lines.append("-- ============================================")
        sql_lines.append(f"-- 公司審核數據導出")
        sql_lines.append(f"-- 導出時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        sql_lines.append(f"-- 共 {len(companies)} 家公司")
        sql_lines.append("-- ============================================\n")
        sql_lines.append("START TRANSACTION;\n")
        
        # 更新審核狀態
        sql_lines.append("-- 更新公司審核狀態\n")
        for company in companies:
            company_id = company['id']
            company_name = company['company_name'].replace("'", "''")
            status = company['status']
            reviewed_at = company['reviewed_at']
            reviewed_by_user_id = company['reviewed_by_user_id']
            
            reviewed_at_str = f"'{reviewed_at.strftime('%Y-%m-%d %H:%M:%S')}'" if reviewed_at else "NULL"
            reviewed_by_str = str(reviewed_by_user_id) if reviewed_by_user_id else "NULL"
            
            sql_lines.append(f"-- 公司: {company_name} (ID: {company_id})")
            sql_lines.append(f"UPDATE internship_companies")
            sql_lines.append(f"SET status = '{status}',")
            sql_lines.append(f"    reviewed_at = {reviewed_at_str},")
            sql_lines.append(f"    reviewed_by_user_id = {reviewed_by_str}")
            sql_lines.append(f"WHERE id = {company_id};")
            sql_lines.append("")
        
        # 更新指導老師
        sql_lines.append("-- 更新公司指導老師\n")
        for company in companies:
            if company['advisor_user_id']:
                sql_lines.append(f"UPDATE internship_companies")
                sql_lines.append(f"SET advisor_user_id = {company['advisor_user_id']}")
                sql_lines.append(f"WHERE id = {company['id']};")
                sql_lines.append("")
        
        # 更新開放狀態
        sql_lines.append("-- 更新公司開放狀態\n")
        for company_id, opening_list in openings_dict.items():
            for opening in opening_list:
                semester = opening['semester']
                is_open = 1 if opening['is_open'] else 0
                opened_at = opening['opened_at']
                opened_at_str = f"'{opened_at.strftime('%Y-%m-%d %H:%M:%S')}'" if opened_at else "NOW()"
                
                sql_lines.append(f"INSERT INTO company_openings (company_id, semester, is_open, opened_at)")
                sql_lines.append(f"VALUES ({company_id}, '{semester}', {is_open}, {opened_at_str})")
                sql_lines.append(f"ON DUPLICATE KEY UPDATE")
                sql_lines.append(f"    is_open = {is_open},")
                sql_lines.append(f"    opened_at = {opened_at_str};")
                sql_lines.append("")
        
        sql_lines.append("COMMIT;")
        
        sql_content = '\n'.join(sql_lines)
        
        from flask import Response
        return Response(
            sql_content,
            mimetype='text/plain',
            headers={
                'Content-Disposition': f'attachment; filename=company_reviews_export_{datetime.now().strftime("%Y%m%d")}.sql'
            }
        )
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"導出失敗: {str(e)}"}), 500
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# =========================================================
# API - 取得待審核公司列表
# =========================================================
@company_bp.route("/api/get_pending_companies", methods=["GET"])
def api_get_pending_companies():
    """取得狀態為 pending 的待審核公司列表"""
    conn = None
    cursor = None
    try:
        # 權限檢查：只有主任、科助、管理員可以查看
        if 'user_id' not in session:
            return jsonify({"success": False, "message": "請先登入"}), 403
        
        user_role = session.get('role')
        if user_role not in ['director', 'ta', 'admin']:
            return jsonify({"success": False, "message": "無權限"}), 403

        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("""
            SELECT 
                ic.id,
                ic.company_name,
                u.name AS upload_teacher_name,
                ic.submitted_at
            FROM internship_companies ic
            LEFT JOIN users u ON ic.uploaded_by_user_id = u.id
            WHERE ic.status = 'pending'
            ORDER BY ic.submitted_at ASC
        """)

        companies = cursor.fetchall()

        # 格式化日期時間
        from datetime import datetime, timezone, timedelta
        taiwan_tz = timezone(timedelta(hours=8))
        
        for company in companies:
            if isinstance(company.get('submitted_at'), datetime):
                # 將 UTC 轉為台灣時間
                company['submitted_at'] = company['submitted_at'].astimezone(taiwan_tz).strftime("%Y-%m-%d %H:%M")
            else:
                company['submitted_at'] = company.get('submitted_at', '') or ''

        return jsonify({
            "success": True,
            "companies": companies
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"載入失敗: {str(e)}"
        }), 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# -------------------------
# 學生端：取得所有可投遞實習公司
# -------------------------
@company_bp.route('/api/student/companies', methods=['GET'])
def get_student_companies():
    """取得所有已審核通過的實習公司（不限制學期開放狀態）"""
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 取得當前學期代碼（用於標記哪些公司是當前學期開放的）
        from semester import get_current_semester_code
        current_semester_code = get_current_semester_code(cursor)
        
        # 查詢所有已審核通過的公司，並標記當前學期是否開放
        if current_semester_code:
            cursor.execute("""
                SELECT
                    ic.id,
                    ic.company_name,
                    ic.location,
                    NULL AS industry,
                    COUNT(DISTINCT ij.id) AS job_count,
                    COALESCE(co.is_open, FALSE) AS is_open_current_semester
                FROM internship_companies ic
                LEFT JOIN internship_jobs ij
                    ON ic.id = ij.company_id
                    AND ij.is_active = 1
                LEFT JOIN company_openings co
                    ON ic.id = co.company_id
                    AND co.semester = %s
                WHERE ic.status = 'approved'
                GROUP BY ic.id, co.is_open
                ORDER BY ic.company_name
            """, (current_semester_code,))
        else:
            # 如果沒有設定當前學期，只顯示公司基本信息
            cursor.execute("""
                SELECT
                    ic.id,
                    ic.company_name,
                    ic.location,
                    NULL AS industry,
                    COUNT(DISTINCT ij.id) AS job_count,
                    FALSE AS is_open_current_semester
                FROM internship_companies ic
                LEFT JOIN internship_jobs ij
                    ON ic.id = ij.company_id
                    AND ij.is_active = 1
                WHERE ic.status = 'approved'
                GROUP BY ic.id
                ORDER BY ic.company_name
            """)

        companies = cursor.fetchall()
        
        # 確保 is_open_current_semester 是布林值
        for company in companies:
            company['is_open_current_semester'] = bool(company.get('is_open_current_semester', False))
            company['job_count'] = int(company.get('job_count', 0))

        return jsonify({
            'success': True,
            'companies': companies
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'載入公司資料失敗: {str(e)}'
        }), 500
    finally:
        cursor.close()
        conn.close()

# =========================================================
# 🖥️ 審核公司頁面
# =========================================================
@company_bp.route('/approve_company', methods=['GET'])
def approve_company_form_page():
    return render_template('company/approve_company.html')

# =========================================================
# 🖥️ 查看公司頁面
# =========================================================
@company_bp.route("/look_company")
def look_company_page():
    return render_template("company/look_company.html")

# =========================================================
# 📤 學生投遞履歷
# =========================================================
@company_bp.route('/api/student/apply_company', methods=['POST'])
def apply_company():
    """學生投遞履歷到公司"""
    if 'user_id' not in session or session.get('role') != 'student':
        return jsonify({"success": False, "message": "未授權"}), 403
    
    data = request.get_json()
    company_id = data.get('company_id')
    job_id = data.get('job_id')
    resume_id = data.get('resume_id')
    
    if not company_id or not job_id or not resume_id:
        return jsonify({"success": False, "message": "缺少必要參數"}), 400
    
    user_id = session['user_id']
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    
    try:
        # 驗證履歷屬於該用戶，且為正式版本（可以投遞）
        cursor.execute("""
            SELECT id, status, category FROM resumes 
            WHERE id = %s AND user_id = %s
        """, (resume_id, user_id))
        resume = cursor.fetchone()
        
        if not resume:
            return jsonify({"success": False, "message": "履歷不存在或無權限"}), 403
        
        # 只允許投遞正式版本的履歷（category='ready'）
        resume_category = resume.get('category', 'draft')
        
        if resume_category != 'ready':
            return jsonify({
                "success": False, 
                "message": "請先在履歷管理頁面提交履歷為正式版本後再投遞"
            }), 400
        
        # 驗證公司和職缺存在
        cursor.execute("""
            SELECT ij.id FROM internship_jobs ij
            INNER JOIN internship_companies ic ON ij.company_id = ic.id
            WHERE ij.id = %s AND ic.id = %s AND ic.status = 'approved' AND ij.is_active = 1
        """, (job_id, company_id))
        job = cursor.fetchone()
        
        if not job:
            return jsonify({"success": False, "message": "職缺不存在或公司未審核通過"}), 400
        
        # 允許重複投遞：移除重複投遞檢查，允許同一版本的履歷可以重複投遞到同一職缺
        
        # 獲取當前學期
        from semester import get_current_semester_id
        current_semester_id = get_current_semester_id(cursor)
        
        # 獲取職缺名稱
        cursor.execute("SELECT title FROM internship_jobs WHERE id = %s", (job_id,))
        job_title_result = cursor.fetchone()
        job_title = job_title_result['title'] if job_title_result else ''
        
        # 獲取 preference_order：找到該學生在該學期中已使用的最大 preference_order，然後加 1
        # 如果最大 order < 100，則從 100 開始（避免與志願序的 1-5 衝突）
        if current_semester_id:
            cursor.execute("""
                SELECT COALESCE(MAX(preference_order), 0) AS max_order
                FROM student_preferences
                WHERE student_id = %s AND semester_id = %s
            """, (user_id, current_semester_id))
        else:
            cursor.execute("""
                SELECT COALESCE(MAX(preference_order), 0) AS max_order
                FROM student_preferences
                WHERE student_id = %s
            """, (user_id,))
        max_order_result = cursor.fetchone()
        max_order = max_order_result['max_order'] if max_order_result else 0
        # 如果最大 order < 100，則從 100 開始，避免與志願序（通常 1-5）衝突
        # 否則使用最大 order + 1
        preference_order = 100 if max_order < 100 else max_order + 1
        
        # 當投遞履歷時，確保 status='uploaded'（審核中）
        # category 保持 'ready'（正式版本），審核狀態用 status 表示
        cursor.execute("""
            UPDATE resumes 
            SET status = 'uploaded',
                updated_at = NOW()
            WHERE id = %s AND user_id = %s
            AND category = 'ready'
        """, (resume_id, user_id))
        
        # 插入投遞記錄到 student_preferences
        # 檢查 student_preferences 表是否有 resume_id 字段
        cursor.execute("""
            SELECT COLUMN_NAME 
            FROM INFORMATION_SCHEMA.COLUMNS 
            WHERE TABLE_SCHEMA = DATABASE() 
            AND TABLE_NAME = 'student_preferences' 
            AND COLUMN_NAME = 'resume_id'
        """)
        columns = [row['COLUMN_NAME'] for row in cursor.fetchall()]
        has_resume_id = 'resume_id' in columns
        
        # 根據表結構動態構建 INSERT 語句
        if has_resume_id:
            # 表有 resume_id 字段（不再使用 folder_id）
            if current_semester_id:
                cursor.execute("""
                    INSERT INTO student_preferences
                    (student_id, semester_id, preference_order, company_id, job_id, resume_id, job_title, status, submitted_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    user_id,
                    current_semester_id,
                    preference_order,
                    company_id,
                    job_id,
                    resume_id,
                    job_title,
                    'submitted',
                    datetime.now()
                ))
            else:
                cursor.execute("""
                    INSERT INTO student_preferences
                    (student_id, preference_order, company_id, job_id, resume_id, job_title, status, submitted_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    user_id,
                    preference_order,
                    company_id,
                    job_id,
                    resume_id,
                    job_title,
                    'submitted',
                    datetime.now()
                ))
        else:
            # 表沒有 folder_id 和 resume_id 字段（向後兼容）
            if current_semester_id:
                cursor.execute("""
                    INSERT INTO student_preferences
                    (student_id, semester_id, preference_order, company_id, job_id, job_title, status, submitted_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    user_id,
                    current_semester_id,
                    preference_order,
                    company_id,
                    job_id,
                    job_title,
                    'submitted',
                    datetime.now()
                ))
            else:
                cursor.execute("""
                    INSERT INTO student_preferences
                    (student_id, preference_order, company_id, job_id, job_title, status, submitted_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (
                    user_id,
                    preference_order,
                    company_id,
                    job_id,
                    job_title,
                    'submitted',
                    datetime.now()
                ))
        
        conn.commit()
        return jsonify({"success": True, "message": "投遞成功"})
        
    except Exception as e:
        conn.rollback()
        traceback.print_exc()
        return jsonify({"success": False, "message": f"投遞失敗: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()