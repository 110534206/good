from flask import Blueprint, render_template, session, redirect, url_for, request, jsonify, current_app, send_file
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from config import get_db
from semester import is_student_in_application_phase, should_show_intern_experience, should_show_image_recognize
import os
import re 
from docx import Document

users_bp = Blueprint("users_bp", __name__)


def _get_active_semester_year(cursor):
    """取得當前啟用學期學年（semesters 表 is_active=1 的 code 前三碼，如 1132->113），與 admin 邏輯一致。"""
    cursor.execute("SELECT code FROM semesters WHERE is_active = 1 LIMIT 1")
    row = cursor.fetchone()
    if not row or row.get('code') is None:
        cursor.execute("SELECT code FROM semesters WHERE code IS NOT NULL AND code != '' ORDER BY code DESC LIMIT 1")
        row = cursor.fetchone()
    if not row or row.get('code') is None:
        cursor.execute("SELECT code FROM semesters ORDER BY id DESC LIMIT 1")
        row = cursor.fetchone()
    raw = row.get('code') if row else None
    if raw is None:
        return None
    if isinstance(raw, int):
        return raw // 10 if raw >= 100 else None
    if isinstance(raw, bytes):
        raw = raw.decode('utf-8', errors='ignore')
    code = str(raw).strip()
    if len(code) >= 3:
        try:
            return int(code[:3])
        except (TypeError, ValueError):
            pass
    return None


# -------------------------
# API: 取得當前使用者基本資料（姓名、電子信箱、頭貼）供履歷表單帶入
# -------------------------
@users_bp.route('/api/current_user_profile', methods=['GET'])
def current_user_profile():
    if 'user_id' not in session:
        return jsonify({"success": False, "message": "未登入"}), 403
    user_id = session['user_id']
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("SELECT name, email, avatar_url FROM users WHERE id=%s", (user_id,))
        row = cursor.fetchone()
        if not row:
            return jsonify({"success": False, "message": "找不到使用者"}), 404
        avatar_url = (row.get("avatar_url") or "").replace("\\", "/").strip() or None
        return jsonify({
            "success": True,
            "data": {
                "name": row.get("name") or "",
                "email": row.get("email") or "",
                "avatar_url": avatar_url
            }
        })
    finally:
        cursor.close()
        conn.close()


role_map = {
    "student": "學生",
    "teacher": "指導老師",
    "director": "主任",
    "ta": "科助",
    "admin": "管理員",
    "vendor": "廠商",
    "class_teacher": "班導師"
}
role_map_reverse = {
    "學生": "student",
    "指導老師": "teacher",
    "班導師": "teacher",
    "主任": "director",
    "科助": "ta",
    "管理員": "admin",
    "廠商": "vendor",
    "班導師": "teacher" 
}

# -------------------------
# 指導老師首頁
# -------------------------
@users_bp.route('/teacher_home')
def teacher_home():
    # 允許 teacher、director、class_teacher 進入
    if 'username' not in session or session.get('role') not in ['teacher', 'director', 'class_teacher']:
        return redirect(url_for('auth_bp.login_page'))

    # 若目前是班導身分，切回指導老師身分
    if session.get('role') == 'class_teacher':
        # 不論原本是主任或老師，都暫時切回指導老師身份
        session['role'] = 'teacher'
        session['display_role'] = '指導老師'

    # 記得保留原始身份（供切回班導時使用）
    if 'original_role' not in session:
        # 若第一次登入，紀錄原始身份
        session['original_role'] = 'director' if session.get('role') == 'director' else 'teacher'

    return render_template('user_shared/teacher_home.html')

# -------------------------
# 指導老師：審核實習心得頁面
# -------------------------
@users_bp.route('/review_experience')
def review_experience_page():
    """指導老師檢視學生實習心得列表"""
    current_role = session.get("role")
    if 'username' not in session or current_role not in ['teacher', 'director', 'class_teacher']:
        return redirect(url_for('auth_bp.login_page'))

    return render_template('user_shared/review experience.html')

# -------------------------
# 班導首頁
# -------------------------
@users_bp.route("/class_teacher_home")
def class_teacher_home():
    # 確保只有老師或主任身份可以進入（包括 class_teacher）
    current_role = session.get("role")
    if "username" not in session or current_role not in ["teacher", "director", "class_teacher"]:
        return redirect(url_for("auth_bp.login_page"))

    # 如果當前是 class_teacher，需要恢復為原始角色進行檢查
    if current_role == "class_teacher":
        original_role = session.get("original_role")
        if original_role == "director":
            session["role"] = "director"
        else:
            session["role"] = "teacher"

    # 若沒有班導師身分，導回原本主頁
    if not session.get("is_homeroom"):
        current_role = session.get("role")
        if current_role == 'director':
            return redirect(url_for("users_bp.director_home")) # 主任導回主任主頁
        else:
            return redirect(url_for("users_bp.teacher_home")) # 老師導回指導老師主頁

    # 新增：進入班導頁時，暫時設定為 "class_teacher"
    session["role"] = "class_teacher"

    return render_template("user_shared/class_teacher_home.html",
                           username=session.get("username"),
                           original_role=session.get("original_role"))

# -------------------------
# 班導查看學生履歷
# -------------------------
@users_bp.route("/class_review_resume")
def class_review_resume():
    # 確保只有老師或主任身份可以進入（包括 class_teacher）
    current_role = session.get("role")
    if "username" not in session or current_role not in ["teacher", "director", "class_teacher"]:
        return redirect(url_for("auth_bp.login_page"))

    # 如果當前是 class_teacher，需要恢復為原始角色進行檢查
    if current_role == "class_teacher":
        original_role = session.get("original_role")
        if original_role == "director":
            session["role"] = "director"
        else:
            session["role"] = "teacher"

    # 若沒有班導師身分，導回原本主頁
    if not session.get("is_homeroom"):
        current_role = session.get("role")
        if current_role == 'director':
            return redirect(url_for("users_bp.director_home"))
        else:
            return redirect(url_for("users_bp.teacher_home"))

    # 進入班導頁時，暫時設定為 "class_teacher"
    session["role"] = "class_teacher"

    return render_template("resume/class_review_resume.html",
                           username=session.get("username"),
                           original_role=session.get("original_role"))

# -------------------------
# Helper - 取得所有學期代碼
# -------------------------
def get_all_semesters(cursor):
    """從 semesters 表格中獲取所有學期代碼和名稱。"""
    # 這裡假設 semesters 表中有 id 和 code (學期代碼，例如 1132) 欄位
    cursor.execute("SELECT code, code AS display_name FROM semesters ORDER BY code DESC")
    return cursor.fetchall()

# -------------------------
# API - 取得個人資料
# -------------------------
@users_bp.route("/api/profile", methods=["GET"])
def get_profile():
    user_id = session.get("user_id")
    if not user_id:
        return jsonify({"success": False, "message": "尚未登入"}), 401

    active_role = session.get("role", "")

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("""
            SELECT u.id, u.username, u.email, u.role AS original_role, u.name,
                   c.department, c.name AS class_name, u.class_id, u.avatar_url,
                   u.teacher_id, u.user_changed, u.admission_year AS db_admission_year
            FROM users u
            LEFT JOIN classes c ON u.class_id = c.id
            WHERE u.id = %s
        """, (user_id,))
        user = cursor.fetchone()

        if not user:
            return jsonify({"success": False, "message": "使用者不存在"}), 404

        # ------------------------------
        # 實習學期：僅學生顯示。由 internship_configs 取得該屆（或個人）設定，顯示學期代碼（如 1132、1142）。
        # 不限制「當前啟用學期」，111 入學等屆別才能看到自己的實習學期（如 1142）。
        # ------------------------------
        user['current_semester_display'] = ''
        user['current_semester_code'] = None
        is_student = (user.get('original_role') == 'student')
        db_ay = user.get('db_admission_year')
        admission_year_val = None
        if db_ay is not None and str(db_ay).strip() != '':
            try:
                admission_year_val = int(db_ay)
            except (TypeError, ValueError):
                pass
        if admission_year_val is None and is_student and user.get('username') and len(str(user.get('username', ''))) >= 3:
            try:
                admission_year_val = int(str(user['username'])[:3])
            except (TypeError, ValueError):
                pass
        if is_student and admission_year_val is not None:
            cursor.execute("""
                SELECT ic.semester_id, s.code
                FROM internship_configs ic
                JOIN semesters s ON s.id = ic.semester_id
                WHERE ic.user_id = %s OR (ic.user_id IS NULL AND ic.admission_year = %s)
                ORDER BY ic.user_id DESC
                LIMIT 1
            """, (user_id, admission_year_val))
            ic_row = cursor.fetchone()
            if ic_row:
                code_val = ic_row.get('code')
                user['current_semester_display'] = str(code_val).strip() if code_val is not None else ''
                user['current_semester_code'] = ic_row.get('semester_id')
        # ------------------------------

        original_role_from_db = user.pop("original_role")
        # 返回當前 session 的 active_role，供前端區分班導(class_teacher)與指導老師(teacher)以顯示對應標籤與資料
        user["role"] = active_role
        user["display_role"] = role_map.get(active_role, active_role)
        user["original_role"] = original_role_from_db  # 同時返回原始角色，供前端參考

        # 入學屆數：優先使用資料庫欄位，學生可從 username 前 3 碼推得
        db_ay = user.pop("db_admission_year", None)
        if db_ay is not None and str(db_ay).strip() != "":
            user["admission_year"] = str(db_ay)
        elif original_role_from_db == "student" and user.get("username") and len(user["username"]) >= 3:
            user["admission_year"] = user["username"][:3]
        else:
            user["admission_year"] = ""

        is_homeroom = session.get("is_homeroom", False)
        classes = []
        if original_role_from_db in ("teacher", "director"):
            cursor.execute("""
                SELECT c.id, c.name, c.department, c.admission_year, ct.role
                FROM classes c
                JOIN classes_teacher ct ON c.id = ct.class_id
                WHERE ct.teacher_id = %s
            """, (user["id"],))
            classes = cursor.fetchall()
            user["classes"] = classes

        user["is_homeroom"] = is_homeroom
        user["email"] = user["email"] or ""

        # 班導：帶班班級、年級（classes_teacher role=班導師）；指導老師：指導學生所屬班級（teacher_student_relations）
        # 年級 = (當前學年 - admission_year + 1)，班級顯示為「科系 年級+班序」如 資管 四孝
        active_year = _get_active_semester_year(cursor)
        grade_labels = ('一', '二', '三', '四', '五', '六')

        def _class_display_row(dept, cname, admission_yr):
            """由 admission_year 與純班序組合成「科系 年級班序」，如 資管科 四孝、資管科 三忠。"""
            dept_str = (dept or '').strip()
            name_str = (cname or '').strip()
            if not name_str:
                return dept_str or '-'
            if admission_yr is None or active_year is None:
                return f"{dept_str} {name_str}".strip() if dept_str else name_str
            try:
                ay_int = int(admission_yr)
            except (TypeError, ValueError):
                return f"{dept_str} {name_str}".strip() if dept_str else name_str
            grade_num = active_year - ay_int + 1
            if 1 <= grade_num <= 6:
                grade_char = grade_labels[grade_num - 1]
            elif grade_num > 0:
                grade_char = str(grade_num)
            else:
                grade_char = ''
            return f"{dept_str} {grade_char}{name_str}".strip() if grade_char else (f"{dept_str} {name_str}".strip() if dept_str else name_str)

        if active_role in ("teacher", "director", "class_teacher") and original_role_from_db in ("teacher", "director"):
            homeroom_class_names = []
            if is_homeroom and classes:
                homeroom_classes = [c for c in classes if c.get("role") == "classteacher"]
                homeroom_class_names = [_class_display_row(c.get('department'), c.get('name'), c.get('admission_year')) for c in homeroom_classes]
            user["homeroom_class_display"] = "、".join(homeroom_class_names) if homeroom_class_names else ""

            guided_class_names = []
            cursor.execute("""
                SELECT DISTINCT c.id, c.department, c.name, c.admission_year
                FROM teacher_student_relations tsr
                JOIN users u ON u.id = tsr.student_id AND u.role = 'student'
                JOIN classes c ON c.id = u.class_id
                WHERE tsr.teacher_id = %s
                ORDER BY c.department, c.name
            """, (user["id"],))
            guided_classes = cursor.fetchall()
            if guided_classes:
                guided_class_names = [_class_display_row(c.get('department'), c.get('name'), c.get('admission_year')) for c in guided_classes]
            user["guided_class_display"] = "、".join(guided_class_names) if guided_class_names else ""

            # 相容用：班導優先顯示帶班，否則顯示指導學生所屬
            user["class_display_name"] = user["homeroom_class_display"] or user["guided_class_display"]
        elif original_role_from_db == "student":
            dept = (user.get('department') or '').strip()
            cname = (user.get('class_name') or '').strip()
            # 班級格式：資管科 X孝，X 為動態年級（與 user_management 年級邏輯一致）
            ay = user.get('admission_year')
            if ay is not None and str(ay).strip() != '':
                try:
                    ay_int = int(ay)
                except (TypeError, ValueError):
                    ay_int = None
            else:
                ay_int = None
            if ay_int is None and user.get('username') and len(str(user.get('username', ''))) >= 3:
                try:
                    ay_int = int(str(user['username'])[:3])
                except (TypeError, ValueError):
                    pass
            grade_char = ''
            active_year = _get_active_semester_year(cursor)
            if active_year is not None and ay_int is not None:
                grade_num = active_year - ay_int + 1
                grade_labels = ('一', '二', '三', '四', '五', '六')
                if 1 <= grade_num <= 6:
                    grade_char = grade_labels[grade_num - 1]
                elif grade_num > 0:
                    grade_char = str(grade_num)
            # 年級顯示（供個人資料頁「年級」欄位）
            if grade_char:
                user["grade_display"] = {'一': '一年級', '二': '二年級', '三': '三年級', '四': '四年級', '五': '五年級', '六': '六年級'}.get(grade_char, grade_char + '年級')
            else:
                user["grade_display"] = ''
            if dept and cname:
                user["class_display_name"] = f"{dept} {grade_char}{cname}" if grade_char else f"{dept} {cname}"
            elif dept or cname:
                user["class_display_name"] = f"{dept} {grade_char}{cname}".strip() or dept or cname
            else:
                user["class_display_name"] = '-'
        else:
            user["class_display_name"] = ""

        # 如果是廠商，獲取對應的指導老師資訊
        # 優先使用 users 表中的 teacher_id 欄位（如果有的話）
        # 如果沒有，則從 internship_companies 表中查詢
        if original_role_from_db == "vendor":
            # 初始化 advisor_name
            user["advisor_name"] = ""
            
            # 首先檢查 users 表中是否有直接儲存的 teacher_id
            teacher_id_from_users = user.get("teacher_id")
            
            if teacher_id_from_users:
                # 如果有 teacher_id，查詢老師的名字
                cursor.execute("""
                    SELECT name FROM users 
                    WHERE id = %s AND role IN ('teacher', 'director')
                """, (teacher_id_from_users,))
                teacher_row = cursor.fetchone()
                if teacher_row:
                    teacher_name = teacher_row.get('name') if isinstance(teacher_row, dict) else teacher_row[0]
                    if teacher_name:
                        user["advisor_name"] = teacher_name
                        print(f"✅ 從 users 表讀取指導老師 ID: {teacher_id_from_users}, 名稱: {user['advisor_name']}")
            
            # 如果從 teacher_id 沒有找到，或者 teacher_id 為空，則從 internship_companies 表中查詢
            if not user.get("advisor_name"):
                vendor_email = user.get("email") or ""
                cursor.execute("""
                    SELECT DISTINCT u.id AS advisor_id, u.name AS advisor_name
                    FROM internship_companies ic
                    LEFT JOIN users u ON ic.advisor_user_id = u.id
                    WHERE (ic.uploaded_by_user_id = %s OR ic.contact_email = %s)
                      AND ic.advisor_user_id IS NOT NULL
                      AND u.name IS NOT NULL
                      AND u.name != ''
                """, (user_id, vendor_email))
                advisors = cursor.fetchall() or []
                
                # 調試信息：記錄查詢結果
                print(f"🔍 廠商 {user_id} (email: {vendor_email}) 從 internship_companies 查詢指導老師: {advisors}")
                
                # 收集所有指導老師名稱
                advisor_names = []
                if advisors:
                    for advisor in advisors:
                        advisor_name = advisor.get("advisor_name")
                        if advisor_name and advisor_name.strip():
                            if advisor_name not in advisor_names:  # 避免重複
                                advisor_names.append(advisor_name)
                
                # 如果有指導老師，顯示所有指導老師（用、分隔）
                if advisor_names:
                    user["advisor_name"] = "、".join(advisor_names)
                    print(f"✅ 從 internship_companies 表查詢到的指導老師名稱: {user['advisor_name']}")
                else:
                    print(f"⚠️ 廠商 {user_id} 沒有找到指導老師資訊")
        else:
            user["advisor_name"] = ""

        # 供前端計算班級年級顯示（學生／老師班級皆可用）
        user["active_semester_year"] = active_year

        return jsonify({"success": True, "user": user})
    except Exception as e:
        print("❌ 取得個人資料錯誤:", e)
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 更新個人資料
# -------------------------
@users_bp.route("/api/saveProfile", methods=["POST"])
def save_profile():
    if session.get('role') == 'guest':
        return jsonify({"success": False, "message": "訪客無權限操作此功能"}), 403

    data = request.get_json()
    username = data.get("username")
    role_display = data.get("role")
    name = data.get("name")
    class_id = data.get("class_id")

    if not username or not role_display or not name:
        return jsonify({"success": False, "message": "缺少必要欄位"}), 400

    role = role_map_reverse.get(role_display)
    if not role:
        return jsonify({"success": False, "message": "身分錯誤"}), 400

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    user_id = session.get("user_id")

    try:
        if not user_id:
            return jsonify({"success": False, "message": "請重新登入"}), 401

        # 檢查是否要修改帳號
        cursor.execute("SELECT username, role FROM users WHERE id = %s", (user_id,))
        current_user = cursor.fetchone()
        
        if not current_user:
            return jsonify({"success": False, "message": "使用者不存在"}), 404

        # 如果提供了新的 username 且與當前不同，則嘗試修改帳號
        if username and username.strip() and username.strip() != current_user.get("username", ""):
            # 檢查新帳號是否已被使用（同一角色下）
            cursor.execute("""
                SELECT id FROM users 
                WHERE username = %s AND role = %s AND id != %s
            """, (username.strip(), current_user["role"], user_id))
            existing = cursor.fetchone()
            if existing:
                return jsonify({"success": False, "message": "此帳號已被使用"}), 400
            
            # 更新帳號，並標記已修改過帳密（user_changed=1）
            cursor.execute("""
                UPDATE users 
                SET username = %s, user_changed = 1
                WHERE id = %s
            """, (username.strip(), user_id))
            
            # 更新 session 中的 username
            session['username'] = username.strip()

        cursor.execute("UPDATE users SET name=%s WHERE id=%s", (name, user_id))

        if role == "student":
            if class_id:
                cursor.execute("SELECT id FROM classes WHERE id=%s", (class_id,))
                if not cursor.fetchone():
                    return jsonify({"success": False, "message": "班級不存在"}), 404
                cursor.execute("UPDATE users SET class_id=%s WHERE id=%s", (class_id, user_id))
        else:
            cursor.execute("UPDATE users SET class_id=NULL WHERE id=%s", (user_id,))

        conn.commit()

        # 判斷是否班導師
        is_homeroom = False
        if role in ("teacher", "director"):
            cursor.execute("""
                SELECT COUNT(*) as count FROM classes_teacher 
                WHERE teacher_id = %s AND role = 'classteacher'
            """, (user_id,))
            result = cursor.fetchone()
            is_homeroom = result[0] > 0 if result else False

        return jsonify({
            "success": True,
            "message": "資料更新成功",
            "role": role,
            "is_homeroom": is_homeroom
        })
    except Exception as e:
        print("❌ 更新資料錯誤:", e)
        conn.rollback()
        return jsonify({"success": False, "message": "資料庫錯誤"}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 取得所有學期
# -------------------------
@users_bp.route("/api/semesters", methods=["GET"])
def get_semesters():
    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT id, code FROM semesters ORDER BY code DESC")
        semesters = cursor.fetchall()
        cursor.close()
        conn.close()
        return jsonify({"success": True, "semesters": semesters})
    except Exception as e:
        return jsonify({"success": False, "message": f"無法取得學期資料：{str(e)}"}), 500

# -------------------------
# API - 上傳頭像
# -------------------------
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@users_bp.route('/api/upload_avatar', methods=['POST'])
def upload_avatar():
    if "user_id" not in session or session.get('role') == 'guest':
        return jsonify({"success": False, "message": "未授權或訪客無權限"}), 401

    if 'avatar' not in request.files:
        return jsonify({"success": False, "message": "沒有檔案"}), 400

    file = request.files['avatar']
    if file and allowed_file(file.filename):
        filename = secure_filename(f"{session['user_id']}.png")
        
        avatars_folder = os.path.join(current_app.static_folder, "avatars")
        os.makedirs(avatars_folder, exist_ok=True)
        
        filepath = os.path.join(avatars_folder, filename)
        file.save(filepath)

        avatar_url = url_for('static', filename=f"avatars/{filename}")
        
        conn = get_db()
        cursor = conn.cursor()
        try:
            cursor.execute("UPDATE users SET avatar_url = %s WHERE id = %s", (avatar_url, session['user_id']))
            conn.commit()
        except Exception as e:
            print("❌ 更新頭像URL錯誤:", e)
            return jsonify({"success": False, "message": "更新頭像URL失敗"}), 500
        finally:
            cursor.close()
            conn.close()
        
        return jsonify({"success": True, "avatar_url": avatar_url})
    else:
        return jsonify({"success": False, "message": "檔案格式錯誤"}), 400

# -------------------------
# API - 變更帳號
# -------------------------
@users_bp.route('/api/change_username', methods=['POST'])
def change_username():
    if "user_id" not in session or session.get('role') == 'guest':
        return jsonify({"success": False, "message": "尚未登入或訪客無權限"}), 401

    data = request.get_json()
    new_username = data.get("new_username")

    if not new_username:
        return jsonify({"success": False, "message": "請輸入新帳號"}), 400

    user_id = session["user_id"]

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # 檢查使用者是否存在
        cursor.execute("SELECT username, role FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()

        if not user:
            return jsonify({"success": False, "message": "使用者不存在"}), 404

        # 檢查新帳號是否已被使用（同一角色下）
        cursor.execute("""
            SELECT id FROM users 
            WHERE username = %s AND role = %s AND id != %s
        """, (new_username, user["role"], user_id))
        existing = cursor.fetchone()
        if existing:
            return jsonify({"success": False, "message": "此帳號已被使用"}), 400

        # 更新帳號
        cursor.execute("""
            UPDATE users 
            SET username = %s
            WHERE id = %s
        """, (new_username, user_id))
        conn.commit()

        # 更新 session 中的 username
        session['username'] = new_username

        return jsonify({
            "success": True,
            "message": "帳號已更新"
        })
    except Exception as e:
        print("❌ 帳號變更錯誤:", e)
        conn.rollback()
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 變更密碼
# -------------------------
@users_bp.route('/api/change_password', methods=['POST'])
def change_password():
    if "user_id" not in session or session.get('role') == 'guest':
        return jsonify({"success": False, "message": "尚未登入或訪客無權限"}), 401

    data = request.get_json()
    old_password = data.get("old_password")
    new_password = data.get("new_password")

    if not old_password or not new_password:
        return jsonify({"success": False, "message": "請填寫所有欄位"}), 400

    user_id = session["user_id"]

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("SELECT password, role FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()

        if not user or not check_password_hash(user["password"], old_password):
            return jsonify({"success": False, "message": "舊密碼錯誤"}), 403

        is_homeroom = False
        if user["role"] in ("teacher", "director"):
            check_cursor = conn.cursor() 
            check_cursor.execute("""
                SELECT COUNT(*) as count FROM classes_teacher 
                WHERE teacher_id = %s AND role = 'classteacher'
            """, (user_id,))
            result = check_cursor.fetchone()
            is_homeroom = result[0] > 0 if result else False
            check_cursor.close()
            
        hashed_pw = generate_password_hash(new_password)
        cursor.execute("UPDATE users SET password = %s, user_changed = 1 WHERE id = %s", (hashed_pw, user_id))
        conn.commit()

        return jsonify({
            "success": True, 
            "message": "密碼已更新",
            "role": user["role"], 
            "is_homeroom": is_homeroom 
        })
    except Exception as e:
        print("❌ 密碼變更錯誤:", e)
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# 訪客主頁
# -------------------------
@users_bp.route('/visitor_page')
def visitor_page():
    # 確保用戶是以訪客身份進入
    if session.get('role') != 'visitor' and session.get('is_visitor') != True:
         # 如果用戶不是訪客身份，導回登入頁面
         return redirect(url_for('auth_bp.login_page'))
         
    return render_template('user_shared/visitor.html')

# -------------------------
# 廠商首頁
# ------------------------
@users_bp.route('/vendor_home')
def vendor_home():
    """
    實習廠商登入後進入的主頁。
    """
    # 權限檢查：必須是已登入的用戶，且角色為 'vendor'
    if 'username' not in session or session.get('role') != 'vendor':
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/vendor_home.html') 


@users_bp.route('/manage_positions')
def manage_positions_page():
    """
    廠商職位需求管理頁面。
    """
    if 'username' not in session or session.get('role') != 'vendor':
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/manage_positions.html')


@users_bp.route('/manage_positions/new')
def manage_positions_create_page():
    """
    廠商新增職缺頁面。
    """
    if 'username' not in session or session.get('role') != 'vendor':
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/create_position.html')

@users_bp.route('/manage_positions/edit/<int:job_id>')
def manage_positions_edit_page(job_id):
    """
    廠商編輯職缺頁面。
    """
    if 'username' not in session or session.get('role') != 'vendor':
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/create_position.html', job_id=job_id)


# -------------------------
# 廠商媒合結果頁面
# -------------------------
@users_bp.route('/confirm_matching')
def confirm_matching_page():
    """
    廠商查看媒合結果的頁面。
    """
    if 'username' not in session or session.get('role') != 'vendor':
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/confirm_matching.html')

# -------------------------
# # 頁面路由
# -------------------------

# 使用者首頁（學生前台）
@users_bp.route('/student_home')
def student_home():
    # 流程學期／實習學期：投遞、志願、行事曆、媒合結果在「實習學期或上一學期」顯示；實習心得學期 1132 時顯示；成績 AI 識別 1131 即顯示
    in_application_phase = True
    show_intern_experience = True
    show_image_recognize = True
    if session.get('role') == 'student' and session.get('user_id'):
        try:
            conn = get_db()
            cursor = conn.cursor(dictionary=True)
            in_application_phase = is_student_in_application_phase(cursor, session['user_id'])
            show_intern_experience = should_show_intern_experience(cursor, session['user_id'])
            show_image_recognize = should_show_image_recognize(cursor, session['user_id'])
            cursor.close()
            conn.close()
        except Exception:
            pass
    return render_template('user_shared/student_home.html',
                          in_application_phase=in_application_phase,
                          show_intern_experience=show_intern_experience,
                          show_image_recognize=show_image_recognize)


@users_bp.route('/image_recognize')
def image_recognize_page():
    """學生圖片識別頁面（從學生首頁卡片進入）"""
    if 'user_id' not in session or session.get('role') != 'student':
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/image_recognize.html')


@users_bp.route('/api/student/image_recognize', methods=['POST'])
def api_student_image_recognize():
    """
    成績 AI 識別：支援圖片（OCR）與 PDF（文字抽取）。
    """
    if 'user_id' not in session or session.get('role') != 'student':
        return jsonify({"success": False, "message": "請先以學生身分登入"}), 401

    image_file = request.files.get('image')
    if not image_file or image_file.filename == '':
        return jsonify({"success": False, "message": "請選擇一張圖片"}), 400

    # 不存檔，只回傳檔名與大小（KB）與辨識結果
    image_file.stream.seek(0, os.SEEK_END)
    size_bytes = image_file.stream.tell()
    image_file.stream.seek(0)

    size_kb = round(size_bytes / 1024, 1) if size_bytes is not None else 0

    filename = image_file.filename or ""
    ext = os.path.splitext(filename)[1].lower()

    # PDF：直接抽取文字，不走圖片 OCR
    if ext == ".pdf":
        try:
            import io
            from PyPDF2 import PdfReader
        except Exception:
            return jsonify({
                "success": True,
                "filename": filename,
                "size_kb": size_kb,
                "text": "",
                "confidence": None,
                "message": "伺服器尚未安裝 PyPDF2，無法解析 PDF 文字，目前僅回傳檔案資訊。"
            })

        try:
            image_file.stream.seek(0)
            pdf_bytes = image_file.stream.read()
            reader = PdfReader(io.BytesIO(pdf_bytes))
            texts = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    texts.append(page_text.strip())
            full_text = "\n\n".join(texts)
        except Exception:
            full_text = ""

        return jsonify({
            "success": True,
            "filename": filename,
            "size_kb": size_kb,
            "text": full_text,
            "confidence": None,
            "message": "PDF 已完成文字抽取（非影像 OCR）。"
        })

    # 圖片：使用 OCR（Pillow + pytesseract）
    try:
        from PIL import Image
        import pytesseract
        from pytesseract import Output
    except Exception:
        # 伺服器尚未安裝 OCR 套件時，仍回傳檔案資訊
        return jsonify({
            "success": True,
            "filename": filename,
            "size_kb": size_kb,
            "text": "",
            "confidence": None,
            "message": "伺服器尚未安裝 OCR 套件（Pillow / pytesseract），目前僅回傳檔案資訊。"
        })

    ocr_text = ""
    avg_conf = None
    try:
        image = Image.open(image_file.stream).convert("RGB")
        data = pytesseract.image_to_data(image, output_type=Output.DICT, config="--oem 3 --psm 6")
        texts = []
        confs = []
        for txt, conf in zip(data.get("text", []), data.get("conf", [])):
            txt = (txt or "").strip()
            try:
                c = float(conf)
            except Exception:
                c = -1
            if txt and c >= 0:
                texts.append(txt)
                confs.append(c)
        if texts:
            ocr_text = " ".join(texts)
        if confs:
            avg_conf = round(sum(confs) / len(confs), 1)
    except Exception:
        # OCR 失敗時，不中斷整個流程
        pass

    return jsonify({
        "success": True,
        "filename": filename,
        "size_kb": size_kb,
        "text": ocr_text,
        "confidence": avg_conf
    })


@users_bp.route('/api/student/image_to_docx', methods=['POST'])
def api_student_image_to_docx():
    """
    將圖片透過 OCR 文字識別後，產生一份 Word 檔並提供下載。
    """
    if 'user_id' not in session or session.get('role') != 'student':
        return jsonify({"success": False, "message": "請先以學生身分登入"}), 401

    image_file = request.files.get('image')
    if not image_file or image_file.filename == '':
        return jsonify({"success": False, "message": "請選擇一張圖片"}), 400

    # 先取得 OCR / 文字抽取結果（重複利用上面的邏輯）
    image_file.stream.seek(0, os.SEEK_END)
    size_bytes = image_file.stream.tell()
    image_file.stream.seek(0)

    size_kb = round(size_bytes / 1024, 1) if size_bytes is not None else 0

    filename = image_file.filename or ""
    ext = os.path.splitext(filename)[1].lower()

    ocr_text = ""
    avg_conf = None
    ocr_available = False

    # PDF：使用 PyPDF2 直接抽取文字
    if ext == ".pdf":
        try:
            import io
            from PyPDF2 import PdfReader
            ocr_available = True
            image_file.stream.seek(0)
            pdf_bytes = image_file.stream.read()
            reader = PdfReader(io.BytesIO(pdf_bytes))
            texts = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    texts.append(page_text.strip())
            ocr_text = "\n\n".join(texts)
        except Exception:
            ocr_available = False
    else:
        # 圖片：使用 OCR
        try:
            from PIL import Image
            import pytesseract
            from pytesseract import Output
            ocr_available = True
        except Exception:
            ocr_available = False

        if ocr_available:
            try:
                image = Image.open(image_file.stream).convert("RGB")
                data = pytesseract.image_to_data(image, output_type=Output.DICT, config="--oem 3 --psm 6")
                texts = []
                confs = []
                for txt, conf in zip(data.get("text", []), data.get("conf", [])):
                    txt = (txt or "").strip()
                    try:
                        c = float(conf)
                    except Exception:
                        c = -1
                    if txt and c >= 0:
                        texts.append(txt)
                        confs.append(c)
                if texts:
                    ocr_text = " ".join(texts)
                if confs:
                    avg_conf = round(sum(confs) / len(confs), 1)
            except Exception:
                ocr_available = False

    # 建立 Word 檔（成績 AI 識別報告）
    doc = Document()
    doc.add_heading('成績 AI 識別結果', level=1)
    doc.add_paragraph(f"檔名：{image_file.filename}")
    doc.add_paragraph(f"檔案大小：約 {size_kb} KB")

    if ocr_available and ocr_text:
        doc.add_paragraph("")
        doc.add_heading('辨識文字', level=2)
        doc.add_paragraph(ocr_text)
        if avg_conf is not None:
            doc.add_paragraph(f"平均信心度：約 {avg_conf}%")
    elif ocr_available and not ocr_text:
        doc.add_paragraph("")
        doc.add_paragraph("OCR 已執行，但未偵測到明顯文字或成績資訊，請確認成績單是否清晰。")
    else:
        doc.add_paragraph("")
        doc.add_paragraph("伺服器目前尚未安裝或無法使用 OCR 套件（Pillow / pytesseract），僅產生包含成績檔案資訊的報告。")

    # 儲存到暫存目錄
    tmp_dir = os.path.join(current_app.root_path, 'tmp')
    os.makedirs(tmp_dir, exist_ok=True)
    safe_name = secure_filename(os.path.splitext(image_file.filename)[0] or "score")
    docx_filename = f"{safe_name}_score_ocr.docx"
    docx_path = os.path.join(tmp_dir, docx_filename)
    doc.save(docx_path)

    return send_file(docx_path, as_attachment=True, download_name=docx_filename)

# 功能操作說明頁面
@users_bp.route('/operation_manual')
def operation_manual():
    return render_template('user_shared/operation_manual.html')

# 實習流程常見問題頁面
@users_bp.route('/intern_faq')
def intern_faq():
    return render_template('user_shared/intern_faq.html')

# 實習常見問題頁面（第二版）
@users_bp.route('/intern_faq_2')
@users_bp.route('/intern_faq_2/')
def intern_faq_2():
    return render_template('user_shared/intern_faq_2.html')

# 使用者首頁 (主任前台)
@users_bp.route("/director_home")
def director_home():
    if "username" not in session or session.get("role") != "director":
        return redirect(url_for("auth_bp.login_page"))

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id, company_name FROM internship_companies WHERE status='pending'")
    companies = cursor.fetchall()
    cursor.close()
    conn.close()
    
    # 🎯 確認：此處直接渲染主任主頁，不執行任何班導師跳轉邏輯。
    return render_template("user_shared/director_home.html", companies=companies)

# 科助
@users_bp.route('/ta_home')
def ta_home():
    return render_template('user_shared/ta_home.html')

# 科助工作台（媒合最終公告、未錄取名單／二面流程）
@users_bp.route('/ta_assistant_dashboard')
def ta_assistant_dashboard():
    if 'user_id' not in session or session.get('role') not in ['ta', 'admin']:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('ta/ta_assistant_dashboard.html')

# 面試排程頁面（科助）
@users_bp.route('/ta/interview_schedule')
def interview_schedule():
    # 權限檢查：允許 ta, admin 訪問
    if 'user_id' not in session or session.get('role') not in ['ta', 'admin']:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/interview_schedule.html')

# 面試排程頁面（主任、指導老師、班導、學生）
@users_bp.route('/interview_schedule')
def director_interview_schedule():
    # 權限檢查：允許 director, teacher, class_teacher, student, admin 訪問
    if 'user_id' not in session or session.get('role') not in ['director', 'teacher', 'class_teacher', 'student', 'admin']:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/interview_schedule.html')

# 廠商媒合結果頁面
@users_bp.route('/ta/match_results')
def ta_match_results():
    """TA查看廠商排序學生的結果頁面"""
    if 'user_id' not in session or session.get('role') not in ['ta', 'admin']:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/match_results.html')
    

# 志願序最終結果
@users_bp.route('/final_results')
def final_results():
    return render_template('user_shared/final_results.html')

# 指導老師／班導：未錄取名單查看頁面
@users_bp.route('/teacher/unadmitted_list')
def teacher_unadmitted_list():
    """
    指導老師／班導查看未錄取名單的頁面。
    允許 teacher、class_teacher、director、ta、admin 訪問。
    """
    if 'user_id' not in session or session.get('role') not in ['teacher', 'class_teacher', 'director', 'ta', 'admin']:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/teacher_unadmitted_list.html')

# 管理員首頁（後台）
@users_bp.route('/admin_home')
def admin_home():
    return render_template('admin/admin_home.html')

# 個人頁面
@users_bp.route('/profile')
def profile():
    return render_template('user_shared/profile.html')

# 主任統計資料頁面
@users_bp.route('/admin_statistics')
def admin_statistics():
    if "username" not in session or session.get("role") != "director":
        return redirect(url_for("auth_bp.login_page"))
    return render_template('user_shared/admin_statistics.html')

# 取得 session 資訊
@users_bp.route('/api/get-session')
def get_session():
    # 檢查是否為訪客
    if session.get("role") == "visitor" or session.get("user_id") == 0:
        return jsonify({
            "success": False,
            "role": "visitor"
        }), 401
    # 檢查是否為有效的登入用戶
    if "username" in session and "role" in session:
        return jsonify({
            "success": True,
            "username": session["username"],
            "role": session["role"]
        })
    return jsonify({"success": False}), 401
    
#實習成果
@users_bp.route('/intern_achievement')
def intern_achievement():
    if 'username' not in session or session.get('role') != 'student':
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/intern_achievement.html')

# -------------------------
# 廠商職缺瀏覽頁面（給指導老師、科助查看所有廠商職缺）
# -------------------------
@users_bp.route('/manage_vendor')
def manage_vendor_page():
    """
    指導老師、科助查看所有廠商職缺的頁面。
    """
    if 'username' not in session:
        return redirect(url_for('auth_bp.login_page'))
    # 只允許指導老師、科助查看
    allowed_roles = ['teacher', 'ta']
    if session.get('role') not in allowed_roles:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/manage_vendor.html')


@users_bp.route('/teacher/company/<int:company_id>')
def teacher_company_detail_page(company_id):
    """
    指導老師、科助查看單一廠商詳細資訊的頁面。
    """
    if 'username' not in session:
        return redirect(url_for('auth_bp.login_page'))
    allowed_roles = ['teacher', 'ta']
    if session.get('role') not in allowed_roles:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('user_shared/teacher_company_detail.html', company_id=company_id)


@users_bp.route('/teacher/company/<int:company_id>/resumes')
def teacher_company_resumes_page(company_id):
    """
    指導老師、科助查看特定公司的履歷審核頁面。
    """
    if 'username' not in session:
        return redirect(url_for('auth_bp.login_page'))
    allowed_roles = ['teacher', 'ta']
    if session.get('role') not in allowed_roles:
        return redirect(url_for('auth_bp.login_page'))
    return render_template('resume/company_resume_review.html', company_id=company_id)

# -------------------------
# API - 獲取所有公開的職缺（給指導老師、科助查看）
# -------------------------
@users_bp.route('/api/public/positions', methods=['GET'])
def get_public_positions():
    """
    獲取當前登入指導老師對接的公司和啟用的職缺。
    只允許指導老師、科助查看。
    指導老師只能看到 advisor_user_id 等於自己的公司。
    科助可以看到所有已審核通過的公司。
    """
    if 'username' not in session:
        return jsonify({"success": False, "message": "未登入"}), 401
    
    # 只允許指導老師、科助查看
    allowed_roles = ['teacher', 'ta']
    if session.get('role') not in allowed_roles:
        return jsonify({"success": False, "message": "無權限"}), 403
    
    user_id = session.get('user_id')
    user_role = session.get('role')
    company_filter = request.args.get("company_id", type=int)
    status_filter = (request.args.get("status") or "").strip().lower()
    keyword = (request.args.get("q") or "").strip()
    
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # 只查詢已審核通過的公司和啟用的職缺
        where_clauses = ["ic.status = 'approved'", "ij.is_active = 1"]
        params = []
        
        # 指導老師只能看到自己對接的公司，科助可以看到所有公司
        if user_role == 'teacher':
            where_clauses.append("ic.advisor_user_id = %s")
            params.append(user_id)

        # 科助 (ta) 可以看到所有已審核通過的公司，不需要額外過濾
        if company_filter:
            where_clauses.append("ij.company_id = %s")
            params.append(company_filter)
        
        if keyword:
            like = f"%{keyword}%"
            where_clauses.append("(ij.title LIKE %s OR ij.description LIKE %s OR ij.remark LIKE %s OR ic.company_name LIKE %s)")
            params.extend([like, like, like, like])
        
        query = f"""
            SELECT
                ij.id,
                ij.company_id,
                ic.company_name,
                ij.title,
                ij.slots,
                ij.description,
                ij.period,
                ij.work_time,
                ij.salary,
                ij.remark,
                ij.is_active
            FROM internship_jobs ij
            JOIN internship_companies ic ON ij.company_id = ic.id
            WHERE {' AND '.join(where_clauses)}
            ORDER BY ic.company_name, ij.title
        """
        cursor.execute(query, tuple(params))
        rows = cursor.fetchall() or []
        
        # 序列化職缺資料
        items = []
        for row in rows:
            items.append({
                "id": row["id"],
                "company_id": row["company_id"],
                "company_name": row["company_name"],
                "title": row["title"],
                "slots": row["slots"],
                "description": row["description"] or "",
                "period": row["period"] or "",
                "work_time": row["work_time"] or "",
                "salary": str(row["salary"]) if row["salary"] else "",
                "remark": row["remark"] or "",
                "is_active": bool(row["is_active"])
            })
        
        # 獲取公司列表（指導老師只能看到自己對接的公司，科助可以看到所有）
        company_where_clause = "ic.status = 'approved'"
        company_params = []
        if user_role == 'teacher':
            company_where_clause += " AND ic.advisor_user_id = %s"
            company_params.append(user_id)
        
        cursor.execute(f"""
            SELECT DISTINCT ic.id, ic.company_name, ic.advisor_user_id, u.name AS advisor_name
            FROM internship_companies ic
            LEFT JOIN users u ON ic.advisor_user_id = u.id
            WHERE {company_where_clause}
            ORDER BY ic.company_name
        """, tuple(company_params))
        companies = cursor.fetchall() or []
        companies_payload = [{"id": c["id"], "name": c["company_name"], "advisor_user_id": c["advisor_user_id"], "advisor_name": c.get("advisor_name")} for c in companies]
        
        # 統計資訊
        stats = {
            "total": len(items),
            "active": len(items),  # 這裡只顯示啟用的，所以全部都是 active
            "inactive": 0
        }
        
        return jsonify({
            "success": True,
            "companies": companies_payload,
            "items": items,
            "stats": stats
        })
    except Exception as exc:
        print(f"❌ 獲取公開職缺失敗：{exc}")
        return jsonify({"success": False, "message": f"載入失敗：{exc}"}), 500
    finally:
        cursor.close()
        conn.close()


@users_bp.route('/api/public/company/<int:company_id>', methods=['GET'])
def get_public_company(company_id):
    """
    取得指定公司的詳細資料與職缺，僅供指導老師、科助使用。
    """
    if 'username' not in session:
        return jsonify({"success": False, "message": "未登入"}), 401

    allowed_roles = ['teacher', 'ta']
    role = session.get('role')
    if role not in allowed_roles:
        return jsonify({"success": False, "message": "無權限"}), 403

    user_id = session.get('user_id')

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        company_conditions = ["ic.id = %s", "ic.status = 'approved'"]
        params = [company_id]
        if role == 'teacher':
            company_conditions.append("ic.advisor_user_id = %s")
            params.append(user_id)

        cursor.execute(f"""
            SELECT
                ic.id,
                ic.company_name,
                ic.description,
                ic.location,
                ic.contact_person,
                ic.contact_title,
                ic.contact_email,
                ic.contact_phone,
                ic.reviewed_at,
                ic.submitted_at
            FROM internship_companies ic
            WHERE {' AND '.join(company_conditions)}
            LIMIT 1
        """, tuple(params))
        company = cursor.fetchone()

        if not company:
            return jsonify({"success": False, "message": "找不到公司資料"}), 404

        cursor.execute("""
            SELECT
                ij.id,
                ij.title,
                ij.slots,
                ij.description,
                ij.period,
                ij.work_time,
                ij.salary,
                ij.remark,
                ij.is_active
            FROM internship_jobs ij
            WHERE ij.company_id = %s
            ORDER BY ij.is_active DESC, ij.title
        """, (company_id,))
        jobs = cursor.fetchall() or []

        company_payload = {
            "id": company["id"],
            "name": company["company_name"],
            "description": company.get("description") or "",
            "location": company.get("location") or "",
            "contact_person": company.get("contact_person") or "",
            "contact_title": company.get("contact_title") or "",
            "contact_email": company.get("contact_email") or "",
            "contact_phone": company.get("contact_phone") or "",
            "submitted_at": company.get("submitted_at"),
            "reviewed_at": company.get("reviewed_at")
        }

        job_items = []
        for job in jobs:
            job_items.append({
                "id": job["id"],
                "title": job["title"],
                "slots": job["slots"],
                "description": job.get("description") or "",
                "period": job.get("period") or "",
                "work_time": job.get("work_time") or "",
                "salary": str(job["salary"]) if job.get("salary") not in (None, "") else "",
                "remark": job.get("remark") or "",
                "is_active": bool(job.get("is_active"))
            })

        return jsonify({
            "success": True,
            "company": company_payload,
            "jobs": job_items
        })
    except Exception as exc:
        print(f"❌ 取得公司詳細資料失敗：{exc}")
        return jsonify({"success": False, "message": "載入失敗，請稍後再試"}), 500
    finally:
        cursor.close()
        conn.close()


@users_bp.route('/api/public/company/<int:company_id>/vendor-reviewed-students', methods=['GET'])
def get_vendor_reviewed_students(company_id):
    """
    取得指定公司中廠商已審核的學生履歷列表，僅供指導老師使用。
    """
    if 'username' not in session:
        return jsonify({"success": False, "message": "未登入"}), 401

    allowed_roles = ['teacher', 'ta']
    role = session.get('role')
    if role not in allowed_roles:
        return jsonify({"success": False, "message": "無權限"}), 403

    user_id = session.get('user_id')

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # 驗證公司是否屬於該指導老師
        company_conditions = ["ic.id = %s", "ic.status = 'approved'"]
        params = [company_id]
        if role == 'teacher':
            company_conditions.append("ic.advisor_user_id = %s")
            params.append(user_id)

        cursor.execute(f"""
            SELECT ic.id, ic.company_name
            FROM internship_companies ic
            WHERE {' AND '.join(company_conditions)}
            LIMIT 1
        """, tuple(params))
        company = cursor.fetchone()

        if not company:
            return jsonify({"success": False, "message": "找不到公司資料"}), 404

        # 查詢所有選擇該公司的學生履歷（包括已審核和未審核的）
        # 返回廠商審核狀態和留言（如果有的話）
        # 只顯示有廠商審核記錄的學生，每個學生只顯示一筆（最新的審核記錄）
        # 使用 resume_applications 表獲取廠商審核狀態（不再使用 vendor_preference_history 表）
        cursor.execute("""
            SELECT 
                u.id AS student_id,
                u.name AS student_name,
                u.username AS student_number,
                c.name AS class_name,
                r.id AS resume_id,
                r.original_filename,
                r.status AS resume_status,
                r.created_at AS resume_uploaded_at,
                sja.id AS application_id,
                ij.title AS job_title,
                -- 審核狀態從 resume_applications.apply_status 獲取
                COALESCE(ra.apply_status, 'uploaded') AS vendor_review_status,
                ra.company_comment AS vendor_comment,
                ra.updated_at AS vendor_reviewed_at
            FROM student_job_applications sja
            JOIN users u ON sja.student_id = u.id
            LEFT JOIN classes c ON u.class_id = c.id
            LEFT JOIN resumes r ON sja.resume_id = r.id AND r.user_id = u.id
            LEFT JOIN internship_jobs ij ON sja.job_id = ij.id
            LEFT JOIN resume_applications ra ON ra.application_id = sja.id AND ra.job_id = sja.job_id
            WHERE sja.company_id = %s
            AND ra.id IS NOT NULL
            AND ra.apply_status IN ('approved', 'rejected')
            ORDER BY u.id, ra.updated_at DESC
        """, (company_id,))
        
        students = cursor.fetchall() or []
        
        # 去重：每個學生只保留一筆記錄（最新的審核記錄）
        # 因為 ORDER BY u.id, vendor_reviewed_at DESC，所以每個學生的第一筆就是最新的
        seen_students = {}
        unique_students = []
        for s in students:
            student_id = s.get('student_id')
            if student_id not in seen_students:
                seen_students[student_id] = True
                unique_students.append(s)
        
        students = unique_students
        
        # 格式化日期
        from datetime import datetime
        for s in students:
            if isinstance(s.get('resume_uploaded_at'), datetime):
                s['resume_uploaded_at'] = s['resume_uploaded_at'].strftime("%Y-%m-%d %H:%M:%S")
            if isinstance(s.get('vendor_reviewed_at'), datetime):
                s['vendor_reviewed_at'] = s['vendor_reviewed_at'].strftime("%Y-%m-%d %H:%M:%S")

        return jsonify({"success": True, "students": students})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"查詢失敗: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


@users_bp.route('/api/public/company/<int:company_id>', methods=['PUT'])
def update_public_company(company_id):
    """
    更新指定公司的資訊，僅供指導老師、科助使用。
    """
    if 'username' not in session:
        return jsonify({"success": False, "message": "未登入"}), 401

    allowed_roles = ['teacher', 'ta']
    role = session.get('role')
    if role not in allowed_roles:
        return jsonify({"success": False, "message": "無權限"}), 403

    user_id = session.get('user_id')
    data = request.get_json() or {}

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # 檢查公司是否存在且用戶有權限
        company_conditions = ["ic.id = %s", "ic.status = 'approved'"]
        params = [company_id]
        if role == 'teacher':
            company_conditions.append("ic.advisor_user_id = %s")
            params.append(user_id)

        cursor.execute(f"""
            SELECT id
            FROM internship_companies ic
            WHERE {' AND '.join(company_conditions)}
            LIMIT 1
        """, tuple(params))
        company = cursor.fetchone()

        if not company:
            return jsonify({"success": False, "message": "找不到公司資料或無權限"}), 404

        # 準備更新欄位
        update_fields = []
        update_values = []

        if 'description' in data:
            update_fields.append("description = %s")
            update_values.append(data.get('description', '').strip())

        if 'contact_person' in data:
            update_fields.append("contact_person = %s")
            update_values.append(data.get('contact_person', '').strip())

        if 'contact_title' in data:
            update_fields.append("contact_title = %s")
            update_values.append(data.get('contact_title', '').strip())

        if 'contact_email' in data:
            update_fields.append("contact_email = %s")
            update_values.append(data.get('contact_email', '').strip())

        if 'contact_phone' in data:
            update_fields.append("contact_phone = %s")
            update_values.append(data.get('contact_phone', '').strip())

        if 'location' in data:
            update_fields.append("location = %s")
            update_values.append(data.get('location', '').strip())

        if not update_fields:
            return jsonify({"success": False, "message": "沒有要更新的欄位"}), 400

        # 執行更新
        update_values.append(company_id)
        cursor.execute(f"""
            UPDATE internship_companies
            SET {', '.join(update_fields)}
            WHERE id = %s
        """, tuple(update_values))

        conn.commit()

        return jsonify({
            "success": True,
            "message": "公司資訊更新成功"
        })
    except Exception as exc:
        conn.rollback()
        print(f"❌ 更新公司資訊失敗：{exc}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": "更新失敗，請稍後再試"}), 500
    finally:
        cursor.close()
        conn.close()


@users_bp.route('/api/public/company/<int:company_id>/jobs/<int:job_id>', methods=['PUT'])
def update_public_company_job(company_id, job_id):
    """
    更新指定公司的職缺資訊，僅供指導老師、科助使用。
    """
    if 'username' not in session:
        return jsonify({"success": False, "message": "未登入"}), 401

    allowed_roles = ['teacher', 'ta']
    role = session.get('role')
    if role not in allowed_roles:
        return jsonify({"success": False, "message": "無權限"}), 403

    user_id = session.get('user_id')
    data = request.get_json() or {}

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # 檢查公司是否存在且用戶有權限
        company_conditions = ["ic.id = %s", "ic.status = 'approved'"]
        params = [company_id]
        if role == 'teacher':
            company_conditions.append("ic.advisor_user_id = %s")
            params.append(user_id)

        cursor.execute(f"""
            SELECT id
            FROM internship_companies ic
            WHERE {' AND '.join(company_conditions)}
            LIMIT 1
        """, tuple(params))
        company = cursor.fetchone()

        if not company:
            return jsonify({"success": False, "message": "找不到公司資料或無權限"}), 404

        # 檢查職缺是否屬於該公司
        cursor.execute("""
            SELECT id, company_id
            FROM internship_jobs
            WHERE id = %s AND company_id = %s
            LIMIT 1
        """, (job_id, company_id))
        job = cursor.fetchone()

        if not job:
            return jsonify({"success": False, "message": "找不到職缺資料"}), 404

        # 準備更新欄位
        update_fields = []
        update_values = []

        if 'title' in data:
            title = data.get('title', '').strip()
            if title:
                update_fields.append("title = %s")
                update_values.append(title)

        if 'slots' in data:
            try:
                slots = int(data.get('slots', 0))
                if slots > 0:
                    update_fields.append("slots = %s")
                    update_values.append(slots)
            except (TypeError, ValueError):
                pass

        if 'period' in data:
            update_fields.append("period = %s")
            update_values.append(data.get('period', '').strip() or None)

        if 'work_time' in data:
            update_fields.append("work_time = %s")
            update_values.append(data.get('work_time', '').strip() or None)

        if 'salary' in data:
            update_fields.append("salary = %s")
            salary_value = data.get('salary', '').strip()
            update_values.append(salary_value if salary_value else None)

        if 'description' in data:
            update_fields.append("description = %s")
            update_values.append(data.get('description', '').strip() or None)

        if 'remark' in data:
            update_fields.append("remark = %s")
            update_values.append(data.get('remark', '').strip() or None)

        if not update_fields:
            return jsonify({"success": False, "message": "沒有要更新的欄位"}), 400

        # 執行更新
        update_values.append(job_id)
        cursor.execute(f"""
            UPDATE internship_jobs
            SET {', '.join(update_fields)}
            WHERE id = %s
        """, tuple(update_values))

        conn.commit()

        # 取得更新後的職缺資料
        cursor.execute("""
            SELECT
                id,
                title,
                slots,
                description,
                period,
                work_time,
                salary,
                remark,
                is_active
            FROM internship_jobs
            WHERE id = %s
        """, (job_id,))
        updated_job = cursor.fetchone()

        return jsonify({
            "success": True,
            "message": "職缺資訊更新成功",
            "item": {
                "id": updated_job["id"],
                "title": updated_job["title"],
                "slots": updated_job["slots"],
                "description": updated_job.get("description") or "",
                "period": updated_job.get("period") or "",
                "work_time": updated_job.get("work_time") or "",
                "salary": str(updated_job["salary"]) if updated_job.get("salary") not in (None, "") else "",
                "remark": updated_job.get("remark") or "",
                "is_active": bool(updated_job.get("is_active"))
            }
        })
    except Exception as exc:
        conn.rollback()
        print(f"❌ 更新職缺資訊失敗：{exc}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": "更新失敗，請稍後再試"}), 500
    finally:
        cursor.close()
        conn.close()

# =========================================================================
# ⭐ 新增功能：公司指導老師管理 API (為前端 manage_companies.html 服務)
# =========================================================================

# -------------------------
# API - 取得所有指導老師列表
# -------------------------
@users_bp.route('/api/get_all_teachers', methods=['GET'])
def get_all_teachers_api():
    """
    獲取所有可作為指導老師的使用者列表 (role: teacher, director)。
    用於前端下拉選單。
    """
    # 權限檢查：只允許科助 (ta)、主任 (director) 或管理員 (admin)
    allowed_roles = ['ta', 'director', 'admin']
    if 'username' not in session or session.get('role') not in allowed_roles:
        return jsonify({"success": False, "message": "無權限"}), 403

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        # 查詢所有指導老師和主任
        cursor.execute("""
            SELECT id, name, username
            FROM users
            WHERE role IN ('teacher', 'director')
            ORDER BY name
        """)
        teachers = cursor.fetchall()
        
        # 格式化輸出
        teachers_payload = [{
            "id": t["id"],
            "name": t["name"],
            "username": t["username"]
        } for t in teachers]

        return jsonify({"success": True, "teachers": teachers_payload})
    except Exception as e:
        print("❌ 取得指導老師列表錯誤:", e)
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 更新公司指導老師
# -------------------------
@users_bp.route('/api/update_company_advisor', methods=['POST'])
def update_company_advisor():
    """
    更新 internship_companies 表中特定公司的 advisor_user_id 欄位。
    """
    # 權限檢查：只允許科助 (ta) 或 主任 (director) 進行操作
    allowed_roles = ['ta', 'director']
    if 'username' not in session or session.get('role') not in allowed_roles:
        return jsonify({"success": False, "message": "無權限"}), 403

    data = request.get_json()
    company_id = data.get("company_id", type=int)
    # advisor_user_id 可以是 None (前端傳入空值代表移除指導老師)
    advisor_user_id = data.get("advisor_user_id") 

    if not company_id:
        return jsonify({"success": False, "message": "缺少公司 ID"}), 400
    
    # 將 None 或空字串轉換為 None (在 Python 中代表 SQL 的 NULL)
    if advisor_user_id == "":
        advisor_user_id = None
        
    conn = get_db()
    cursor = conn.cursor()
    try:
        # 執行更新操作：更新 internship_companies.advisor_user_id
        cursor.execute("""
            UPDATE internship_companies 
            SET advisor_user_id = %s
            WHERE id = %s
        """, (advisor_user_id, company_id))

        if cursor.rowcount == 0:
            conn.rollback()
            return jsonify({"success": False, "message": "找不到公司或更新失敗 (ID: {company_id})"}), 404
        
        conn.commit()

        return jsonify({"success": True, "message": "指導老師已更新"})
    except Exception as e:
        print("❌ 更新公司指導老師錯誤:", e)
        conn.rollback()
        return jsonify({"success": False, "message": "資料庫錯誤"}), 500
    finally:
        cursor.close()
        conn.close()        