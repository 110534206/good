import os
import re
import google.generativeai as genai
from flask import Blueprint, request, Response, jsonify, session, current_app, send_file
from config import get_db
import json
import traceback
from werkzeug.utils import secure_filename
from docx import Document
from PIL import Image, ImageEnhance
import io




# --- 初始化 AI Blueprint ---
ai_bp = Blueprint('ai_bp', __name__)

# --- 初始化 Google GenAI ---
api_key = os.getenv('GEMINI_API_KEY')

# 檢查 API Key 是否存在
if not api_key:
    print("AI 模組警告：在環境變數中找不到 GEMINI_API_KEY。")
    model = None # 將 model 設為 None
else:
    # 設定 Google Gen AI
    genai.configure(api_key=api_key)
    # 初始化模型
    model = genai.GenerativeModel('gemini-2.5-flash')

# ==========================================================
# 🧠 系統提示詞（System Prompt）
# ==========================================================
SYSTEM_PROMPT = """
你是一位專業的實習申請顧問，專長在協助學生撰寫要寄給實習廠商的自我介紹與申請訊息。
請在所有回覆中遵守以下原則：
1. 依據指定語氣設定（專業正式／親切隨和／謹慎的／學術的）維持一致語氣。
2. 將學生提供的履歷重點整理成可直接寄給廠商的訊息，強調技能、成果與申請動機。
3. 禁止加入道歉語、AI 身分或與申請無關的敘述。
4. 全文使用繁體中文，可搭配必要的英文專有名詞。
5. 以具體行動與可量化成果為核心，段落清晰，符合寄給廠商的禮節與期待。
6. 全程使用純文字，禁止產生星號、井字號、底線或其他 Markdown 標記符號。
7. **絕對禁止**在輸出中包含任何解釋性文字、前綴說明（如「這是為您改寫的...」、「以下是...」）、後綴註解或評論。
8. **只輸出修改後的文本內容**，直接從修改後的文本開始，不要有任何說明或介紹。
"""

# 履歷修改輸出格式規則（避免 AI 輸出 * ** # 等符號）
REVISE_OUTPUT_FORMAT_RULE = """
【輸出格式】必須為純文字：禁止使用星號(*)、雙星號(**)、井字號(#)、底線(_)等任何 Markdown 或項目符號。條列時請用換行與數字（如 1. 2.）或頓號、破折號，切勿使用 * 或 ** 標記。標題或重點請直接以文字呈現，不要加粗或裝飾符號。
"""

def _build_keyword_guidance(user_keywords: str) -> str:
    """
    Ensure user-specified keywords actually influence rewriting.
    This is appended into prompts across styles.
    """
    kw = (user_keywords or "").strip()
    if not kw:
        return ""
    return f"""
[目標類別／關鍵字] {kw}
[關鍵字規則]
1. 改寫時必須以「目標類別／關鍵字」為主軸，調整用字、技能描述、成果呈現與動機，使內容更貼近該方向。
2. 只能強化原文已有的經驗與能力，或做合理連結；禁止捏造不存在的經歷、專案、職稱、數字與成就。
3. 請至少在文中 2-4 次自然提到與目標類別相關的技能／工具／工作內容（用自然語句，不要輸出關鍵字清單）。
"""

# ==========================================================
# AI 處理的 API 端點
# ==========================================================
@ai_bp.route('/api/revise-resume', methods=['POST'])
def revise_resume():
    # 檢查 API Key 是否在啟動時成功載入
    if not api_key or not model:
        return jsonify({"error": "AI 服務未正確配置 API Key。"}), 500

    # 接收履歷文本、任務風格、語氣風格、關鍵字（選填）
    try:
        data = request.get_json()
        user_resume_text = data.get('resumeText')
        edit_style = data.get('style', 'polish')
        tone_style = data.get('tone', 'professional')
        user_keywords = (data.get('keywords') or '').strip()

        # 🌟 [新功能] 如果用戶沒有提供 resumeText，自動從資料庫讀取自傳
        if not user_resume_text or not user_resume_text.strip():
            # 檢查用戶是否已登入
            if 'user_id' not in session or session.get('role') != 'student':
                return jsonify({"error": "請先登入並提供履歷文本，或先上傳履歷。"}), 400
            
            user_id = session['user_id']
            conn = get_db()
            cursor = conn.cursor(dictionary=True)
            
            try:
                # 獲取學號
                cursor.execute("SELECT username FROM users WHERE id=%s", (user_id,))
                user_result = cursor.fetchone()
                if not user_result:
                    return jsonify({"error": "找不到使用者資訊。"}), 404
                
                student_id = user_result["username"]
                
                # 從資料庫讀取自傳
                cursor.execute("SELECT Autobiography FROM Student_Info WHERE StuID=%s", (student_id,))
                student_info = cursor.fetchone()
                
                if student_info and student_info.get('Autobiography'):
                    user_resume_text = str(student_info.get('Autobiography', '')).strip()
                    print(f"✅ 自動從資料庫讀取自傳內容，長度: {len(user_resume_text)}")
                else:
                    return jsonify({"error": "資料庫中沒有自傳內容，請先上傳履歷或手動輸入。"}), 400
                    
            except Exception as e:
                print(f"從資料庫讀取自傳失敗: {e}")
                return jsonify({"error": "無法從資料庫讀取自傳，請手動輸入。"}), 500
            finally:
                if cursor:
                    cursor.close()
                if conn:
                    conn.close()

        if not user_resume_text or not user_resume_text.strip():
            return jsonify({"error": "請提供履歷文本。"}), 400

    except Exception as e:
        print(f"請求解析錯誤: {e}")
        return jsonify({"error": "無效的請求格式。"}), 400

    try:
        final_prompt = ""
        
        # --- 步驟一：定義語氣風格 (Tone) ---
        if tone_style == 'friendly':
            tone_prompt = "語氣必須親切隨和。"
        elif tone_style == 'cautious':
            tone_prompt = "語氣必須專業、謹慎且精確。"
        elif tone_style == 'academic':
            tone_prompt = "語氣必須嚴謹、客觀且具學術性。"
        else:
            tone_prompt = "語氣必須專業正式且符合商業履歷標準。規則：1. 避免個人感悟、心態或哲學性描述。2. 強調具體行動和成就。"

        # --- 步驟二：定義主要任務 (Task) ---
        keyword_guidance = _build_keyword_guidance(user_keywords)

        if edit_style == 'light_polish':
            print(f"偵測任務: 輕度潤飾, 語氣: {tone_style}")
            final_prompt = f"""[任務] 在最大程度保留原始結構與內容的前提下，對以下 [原始文本] 進行輕度潤飾。
[嚴格規則]
1. 不得大幅改寫內容
2. 不得新增原文沒有的經驗或成就
3. 保留原本段落順序
4. 僅修正文法、語句不順與過於口語的表達
5. 可微幅優化語氣，但不得改變原意
6. 不得重組全文
7. 修改比例不得超過原文字數的30%
8. {tone_prompt}
9. 只輸出修改後文本，不要任何說明
{REVISE_OUTPUT_FORMAT_RULE}
{keyword_guidance}
[原始文本] {user_resume_text}
[修改後的文本]"""

        elif edit_style == 'keyword_focus':
            # --- 關鍵字導向：若使用者有提供關鍵字則直接使用，否則從原文提取 ---
            if user_keywords:
                keywords = user_keywords
                print(f"偵測任務: 關鍵字導向 (使用者關鍵字: {keywords}), 語氣: {tone_style}")
                final_prompt = (
                    "[任務] 你是一位頂尖的人力資源顧問。"
                    "請在保留原有段落架構和大部分句子的前提下，根據使用者指定的目標關鍵字／偏好類別，"
                    "對以下自傳做『輕度改寫＋關鍵字強化』，讓內容更貼合該方向（例如：行銷、硬體、軟體、客服、會計等）。"
                    "[關鍵規則]"
                    " 1. 必須自然融合並凸顯目標關鍵字相關的經驗、技能與動機，但禁止捏造新的經歷或專案。"
                    " 2. 優先保留原本的句意與順序，只在必要處調整用字與語氣，整體修改幅度不超過原文字數的 40%。"
                    f" 3. {tone_prompt}"
                    " 4. 使用強動詞開頭的行動句，盡可能量化成果。"
                    " 5. 絕對禁止包含任何解釋性文字、前綴說明、後綴註解或評論。"
                    " 6. 只輸出修改後的文本內容，直接從修改後的文本開始。"
                    f"{REVISE_OUTPUT_FORMAT_RULE}"
                    f"{keyword_guidance}"
                    f"[原始文本] {user_resume_text} [目標關鍵字／偏好類別] {keywords} [修改後的文本]"
                )
            else:
                keyword_prompt = f"[任務] 從以下履歷文本中提取 5-7 個最核心的技能和成就關鍵字。[規則] 以逗號 (,) 分隔所有關鍵字，並在**一行中**輸出。[原始文本] {user_resume_text} [關鍵字列表]"
                keyword_response = model.generate_content(keyword_prompt)
                keywords = keyword_response.text.strip()
                print(f"偵測任務: 關鍵字導向 (自動提取關鍵字: {keywords}), 語氣: {tone_style}")
                final_prompt = (
                    "[任務] 你是一位頂尖的人力資源顧問。"
                    "請在保留原有段落架構和大部分句子的前提下，根據下列核心關鍵字，"
                    "對自傳做『輕度改寫＋關鍵字強化』，讓內容更聚焦於學生的專長。"
                    "[關鍵規則]"
                    " 1. 必須自然凸顯核心關鍵字相關的技能和成就，但禁止捏造新的經歷或專案。"
                    " 2. 優先保留原本的句意與順序，只在必要處調整用字與語氣，整體修改幅度不超過原文字數的 40%。"
                    f" 3. {tone_prompt}"
                    " 4. 使用強動詞開頭的行動句，盡可能量化成果。"
                    " 5. 絕對禁止包含任何解釋性文字、前綴說明、後綴註解或評論。"
                    " 6. 只輸出修改後的文本內容，直接從修改後的文本開始。"
                    f"{REVISE_OUTPUT_FORMAT_RULE}"
                    f"{_build_keyword_guidance(keywords)}"
                    f"[原始文本] {user_resume_text} [核心關鍵字] {keywords} [修改後的文本]"
                )
        
        elif edit_style == 'concise':
            # --- 選項 2: 文案精簡 (保持段落、縮短篇幅) ---
            keyword_instruction = "此外，若以下有提供目標類別／關鍵字，請在精簡時特別保留或強化與之相關的表述。" if user_keywords else ""
            print(f"偵測任務: 文案精簡, 語氣: {tone_style}" + (f", 關鍵字: {user_keywords}" if user_keywords else ""))
            final_prompt = (
                "[任務] 在保留原本文本主要內容與段落結構的前提下，將以下 [原始文本] 改寫得更精簡、清楚且成就導向。"
                "[規則] "
                f"1. {tone_prompt} "
                "2. 盡量維持原本的段落分段方式，只需要適度合併或刪減句子，不要改成一條一條的條列句。 "
                "3. 刪除冗長重複、口語化與與主題無關的描述，保留關鍵經歷、成果與能力。 "
                "4. 可重新組織句子讓語意更清楚，但不要完全重寫成另一篇文章。 "
                "5. 目標是將總字數大約縮短到原來的 50%～70%，但不得犧牲重要資訊。 "
                "6. 絕對禁止包含任何解釋性文字、前綴說明、後綴註解或評論。 "
                "7. 只輸出修改後的文本內容，不要有「這是為您改寫的...」、「以下是...」等說明文字，直接從修改後文本開始。"
                f"{keyword_instruction} "
                f"{REVISE_OUTPUT_FORMAT_RULE}"
                f"[原始文本] {user_resume_text}"
            )
            if user_keywords:
                final_prompt += f"{keyword_guidance}"
            final_prompt += " [修改後的文本]"

        else: # 'polish' (預設)
            # --- 選項 3: 履歷美化 (預設) (一步驟) ---
            keyword_instruction = "此外，若以下有提供目標類別／關鍵字，請在美化時自然突出與之相關的經驗與動機。" if user_keywords else ""
            print(f"偵測任務: 履歷美化, 語氣: {tone_style}" + (f", 關鍵字: {user_keywords}" if user_keywords else ""))
            final_prompt = f"[任務] 專業地**美化並潤飾**以下 [原始文本]。[規則] 1. **{tone_prompt}** 2. 使用強動詞開頭的行動句。 3. 盡可能量化成果。 4. 修正文法。 5. **絕對禁止**包含任何解釋性文字、前綴說明、後綴註解或評論。 6. **只輸出修改後的文本內容**，不要有任何「這是為您改寫的...」、「以下是...」等說明文字。 7. 直接從修改後的文本開始輸出，不要有任何前綴。{keyword_instruction} {REVISE_OUTPUT_FORMAT_RULE}[原始文本] {user_resume_text}"
            if user_keywords:
                final_prompt += f"{keyword_guidance}"
            final_prompt += " [修改後的文本]"

        # --- 統一的串流輸出 ---
        def generate_stream():
            try:
                response_stream = model.generate_content(final_prompt, stream=True)
                for chunk in response_stream:
                    if chunk.text:
                        yield chunk.text
            except Exception as e:
                print(f"串流處理中發生錯誤: {e}")
                error_str = str(e)
                # 檢查是否為配額限制錯誤（429）
                if "429" in error_str or "quota" in error_str.lower() or "Quota exceeded" in error_str:
                    retry_seconds = None
                    if "retry in" in error_str.lower() or "retry_delay" in error_str.lower():
                        retry_match = re.search(r'retry in ([\d.]+)s', error_str, re.IGNORECASE)
                        if retry_match:
                            retry_seconds = int(float(retry_match.group(1)))
                    
                    error_message = "⚠️ AI 服務目前使用量已達上限（免費層級每日限制 20 次）"
                    if retry_seconds:
                        error_message += f"，請在 {retry_seconds} 秒後再試"
                    else:
                        error_message += "，請稍後再試或明天再使用此功能"
                    yield error_message
                else:
                    yield f"AI 服務處理失敗: {error_str}"

        headers = {
            'Content-Type': 'text/plain; charset=utf-8',
            'Transfer-Encoding': 'chunked',
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive'
        }
        return Response(generate_stream(), headers=headers)

    except Exception as e:
        print(f"Gemini API 呼叫失敗： {e}")
        error_str = str(e)
        if "429" in error_str or "quota" in error_str.lower() or "Quota exceeded" in error_str:
            retry_seconds = None
            if "retry in" in error_str.lower() or "retry_delay" in error_str.lower():
                retry_match = re.search(r'retry in ([\d.]+)s', error_str, re.IGNORECASE)
                if retry_match:
                    retry_seconds = int(float(retry_match.group(1)))
            
            error_message = "AI 服務目前使用量已達上限（每日限制 20 次）"
            if retry_seconds:
                error_message += f"，請在 {retry_seconds} 秒後再試"
            else:
                error_message += "，請稍後再試或明天再使用此功能"
            
            return jsonify({
                "error": error_message,
                "error_type": "quota_exceeded",
                "retry_after": retry_seconds
            }), 429

        return jsonify({"error": f"AI 服務處理失敗: {error_str}"}), 500

# ==========================================================
# AI 推薦志願序 API 端點
# ==========================================================
@ai_bp.route('/api/recommend-preferences', methods=['POST'])
def recommend_preferences():
    if not model:
        return jsonify({"success": False, "error": "AI 模型未正確初始化"}), 500
    
    # 檢查 API Key
    if not api_key or not model:
        return jsonify({"success": False, "error": "AI 服務未正確配置 API Key。"}), 500
    
    # 權限檢查
    if "user_id" not in session or session.get("role") != "student":
        return jsonify({"success": False, "error": "只有學生可以使用此功能。"}), 403
    
    user_id = session["user_id"]
    conn = None
    cursor = None
    
    try:
        data = request.get_json()
        target_student_id = data.get('student_id')
        if target_student_id:
            user_id = target_student_id
        else:
            user_id = session["user_id"]

        conn = get_db()
        if not conn:
            return jsonify({"error": "無法連接資料庫"}), 500

        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT username FROM users WHERE id=%s", (user_id,))
        user_result = cursor.fetchone()
        if not user_result:
            return jsonify({"success": False, "error": "找不到使用者資訊。"}), 404  
        else:
            student_id = user_result["username"]

        # 檢查是否有上傳履歷
        cursor.execute("SELECT id, status FROM resumes WHERE user_id = %s ORDER BY created_at DESC LIMIT 1", (user_id,))
        resume_record = cursor.fetchone()
        
        if not resume_record:
            return jsonify({
                "success": False,
                "error": "您尚未上傳履歷，請先完成履歷上傳後再使用推薦功能。"
            }), 400
        
        # 自動從資料庫獲取學生的履歷和成績資料
        cursor.execute("SELECT * FROM Student_Info WHERE StuID=%s",(student_id,))
        student_info = cursor.fetchone() or {}
        
        # 獲取課程成績
        try:
            cursor.execute("SHOW COLUMNS FROM course_grades LIKE 'ProofImage'")
            has_proof_image = cursor.fetchone() is not None
            cursor.execute("SHOW COLUMNS FROM course_grades LIKE 'transcript_path'")
            has_transcript_path = cursor.fetchone() is not None
        except:
            has_proof_image = False
            has_transcript_path = False
        
        if has_proof_image:
            cursor.execute("SELECT CourseName, Credits, Grade, ProofImage FROM course_grades WHERE StuID=%s", (student_id,))
        elif has_transcript_path:
            cursor.execute("SELECT CourseName, Credits, Grade, transcript_path FROM course_grades WHERE StuID=%s", (student_id,))
        else:
            cursor.execute("SELECT CourseName, Credits, Grade FROM course_grades WHERE StuID=%s", (student_id,))
        grades = cursor.fetchall() or []
        
        has_transcript_image = False
        for grade in grades:
            if grade.get('ProofImage') or grade.get('transcript_path'):
                has_transcript_image = True
                break
        
        # 獲取證照
        try:
            cursor.execute("SHOW COLUMNS FROM student_certifications")
            cert_columns = {row["Field"] for row in cursor.fetchall()}
            has_cert_path = 'CertPath' in cert_columns
            has_cert_photo_path = 'CertPhotoPath' in cert_columns
            cert_path_field = 'CertPath' if has_cert_path else ('CertPhotoPath' if has_cert_photo_path else None)
        except:
            cert_path_field = None
        
        try:
            cert_path_select = f"sc.{cert_path_field} AS CertPath" if cert_path_field else "NULL AS CertPath"
            cursor.execute(f"""
                SELECT
                    sc.CertName AS CertName,
                    sc.CertType AS CertType,
                    {cert_path_select},
                    sc.AcquisitionDate AS AcquisitionDate,
                    sc.cert_code AS cert_code,
                    CONCAT(COALESCE(cc.job_category, ''), COALESCE(cc.level, '')) AS cert_name_from_code,
                    cc.category AS cert_category,
                    ca.name AS authority_name
                FROM student_certifications sc
                LEFT JOIN certificate_codes cc ON sc.cert_code COLLATE utf8mb4_unicode_ci = cc.code COLLATE utf8mb4_unicode_ci
                LEFT JOIN cert_authorities ca ON cc.authority_id = ca.id
                WHERE sc.StuID = %s
                ORDER BY sc.AcquisitionDate DESC, sc.id ASC
            """, (student_id,))
            cert_rows = cursor.fetchall() or []
            
            certifications = []
            for row in cert_rows:
                cert_name = row.get('cert_name_from_code', '').strip() or row.get('CertName', '').strip()
                cert_type = row.get('cert_category', '').strip() or row.get('CertType', '').strip()
                cert_path = row.get('CertPath', '').strip() or ''
                acquisition_date = row.get('AcquisitionDate', '')
                
                if cert_name:
                    certifications.append({
                        'CertName': cert_name,
                        'CertType': cert_type if cert_type else '其他',
                        'CertPath': cert_path,
                        'AcquisitionDate': acquisition_date,
                        'AuthorityName': row.get('authority_name', '').strip()
                    })
        except Exception as e:
            print(f"⚠️ 證照完整查詢失敗，使用簡單查詢: {e}")
            try:
                if cert_path_field:
                    cursor.execute(f"SELECT CertName, CertType, {cert_path_field} AS CertPath, AcquisitionDate FROM student_certifications WHERE StuID=%s", (student_id,))
                else:
                    cursor.execute("SELECT CertName, CertType, AcquisitionDate FROM student_certifications WHERE StuID=%s", (student_id,))
                cert_rows = cursor.fetchall() or []
                certifications = []
                for row in cert_rows:
                    if row.get('CertName'):
                        certifications.append({
                            'CertName': row.get('CertName', '').strip(),
                            'CertType': row.get('CertType', '').strip() or '其他',
                            'CertPath': row.get('CertPath', '').strip() if cert_path_field else '',
                            'AcquisitionDate': row.get('AcquisitionDate', ''),
                            'AuthorityName': ''
                        })
            except Exception as e2:
                print(f"⚠️ 簡單查詢也失敗: {e2}")
                certifications = []
        
        # 獲取語言能力
        cursor.execute("SELECT Language, Level FROM student_languageskills WHERE StuID=%s", (student_id,))
        languages = cursor.fetchall() or []
        
        # 整理履歷重點文字
        resume_parts = []
        basic_info = []
        if student_info:
            if student_info.get('Major'):
                basic_info.append(f"主修領域：{student_info.get('Major')}")
            if student_info.get('Skills'):
                skills = str(student_info.get('Skills', '')).strip()
                if skills:
                    basic_info.append(f"技能專長：{skills}")
        
        if basic_info:
            resume_parts.append("【基本資訊（從資料庫 Student_Info 表讀取）】\n" + "\n".join(basic_info))
        
        if certifications:
            cert_list = []
            for c in certifications:
                cert_name = c.get('CertName', '').strip()
                cert_type = c.get('CertType', '').strip()
                acquisition_date = c.get('AcquisitionDate', '')
                authority_name = c.get('AuthorityName', '').strip()
                
                if cert_name:
                    cert_info = f"  - {cert_name}"
                    if cert_type:
                        cert_info += f" ({cert_type})"
                    if authority_name:
                        cert_info += f" - 發證單位：{authority_name}"
                    if acquisition_date:
                        cert_info += f" - 取得日期：{acquisition_date}"
                    cert_list.append(cert_info)
            if cert_list:
                resume_parts.append("【證照資格（從資料庫 student_certifications 表讀取，包含證照名稱、類別、發證單位、取得日期等完整資訊）】\n" + "\n".join(cert_list))
        
        if languages:
            lang_list = []
            for l in languages:
                lang = l.get('Language', '').strip()
                level = l.get('Level', '').strip()
                if lang:
                    if level:
                        lang_list.append(f"  - {lang}：{level}")
                    else:
                        lang_list.append(f"  - {lang}")
            if lang_list:
                resume_parts.append("【語言能力（從資料庫 student_languageskills 表讀取）】\n" + "\n".join(lang_list))
        
        if student_info and student_info.get('Autobiography'):
            autobiography = str(student_info.get('Autobiography', '')).strip()
            if autobiography:
                if len(autobiography) > 2000:
                    autobiography = autobiography[:2000] + "..."
                resume_parts.append("【自傳內容（從資料庫 Student_Info.Autobiography 欄位讀取 - AI 分析重點，請優先引用此內容）】\n" + autobiography)
        
        cert_images_count = sum(1 for c in certifications if c.get('CertPath'))
        if cert_images_count > 0:
            resume_parts.append(f"【證照圖片說明】\n學生已上傳 {cert_images_count} 張證照圖片至資料庫系統中，證照資料已完整記錄。")
        
        resume_text = "\n\n".join(resume_parts) if resume_parts else ""
        
        # 整理學業成績摘要
        grades_parts = []
        core_course_total_score = 0.0
        core_course_count = 0

        if grades:
            grade_points = {'A+': 4.3, 'A': 4.0, 'A-': 3.7, 'B+': 3.3, 'B': 3.0, 'B-': 2.7, 'C+': 2.3, 'C': 2.0, 'C-': 1.7, 'D': 1.0, 'F': 0.0}
            total_points = 0
            total_credits = 0
            excellent_courses = []
            good_courses = []
            all_courses_list = []

            for grade in grades:
                course_name = grade.get('CourseName', '').strip()
                if not course_name:
                    continue
                raw_credits = str(grade.get('Credits', '0'))
                if '/' in raw_credits:
                    raw_credits = raw_credits.split('/')[0]
                try:
                    credits = float(raw_credits or 0)
                except ValueError:
                    credits = 0.0

                grade_str = str(grade.get('Grade', '')).strip().upper()
                grade_value = grade.get('Grade', '').strip()
                if grade_value:
                    try:
                        numeric_grade = float(grade_value)
                        if 0 <= numeric_grade <= 100:
                            core_course_total_score += numeric_grade
                            core_course_count += 1
                    except (ValueError, TypeError):
                        pass
                
                if credits > 0 and grade_str in grade_points:
                    total_points += grade_points[grade_str] * credits
                    total_credits += credits
                
                if grade_str in ['A+', 'A', 'A-']:
                    excellent_courses.append(f"{course_name} ({grade_str})")
                elif grade_str in ['B+', 'B']:
                    good_courses.append(f"{course_name} ({grade_str})")
                
                if grade_str in grade_points:
                    all_courses_list.append(f"{course_name}: {grade_str}")
            
            if total_credits > 0:
                gpa = total_points / total_credits
                grades_parts.append(f"GPA: {gpa:.2f}/4.3")
            
            if excellent_courses:
                grades_parts.append(f"優秀課程成績（A以上）：{', '.join(excellent_courses[:8])}")
            if good_courses:
                grades_parts.append(f"良好課程成績（B以上）：{', '.join(good_courses[:5])}")
            if all_courses_list:
                grades_parts.append(f"\n完整課程列表（從資料庫 course_grades 資料表讀取，對應履歷中的「已修習專業核心科目」）：\n" + "\n".join(all_courses_list[:50]))
        
        if core_course_count > 0:
            core_course_avg_score = core_course_total_score / core_course_count
            grades_parts.append(f"專業核心科目平均成績（從資料庫 course_grades 資料表的 Grade 欄位讀取）：{core_course_avg_score:.2f} 分（共 {core_course_count} 門專業核心科目）")
        
        if has_transcript_image:
            grades_parts.append("成績單圖片：已上傳至資料庫（可於系統中查看）")
        
        grades_text = "\n".join(grades_parts) if grades_parts else ""
        
        # 取得所有公司和職缺
        cursor.execute("SELECT COUNT(*) as count FROM internship_companies")
        company_count = cursor.fetchone().get('count', 0)
        
        if company_count == 0:
                return jsonify({
                    "success": False,
                "error": "目前系統中沒有任何公司資料，請聯繫管理員新增公司。"
                }), 400
        
        from semester import get_current_semester_code
        current_semester_code = get_current_semester_code(cursor)
        
        if current_semester_code:
            cursor.execute("""
                SELECT 
                    ic.id AS company_id,
                    ic.company_name,
                    ic.description AS company_description,
                    ic.location AS company_address,
                    ij.id AS job_id,
                    ij.title AS job_title,
                    ij.description AS job_description,
                    ij.period AS job_period,
                    ij.work_time AS job_work_time,
                    ij.remark AS job_remark
                FROM internship_companies ic
                INNER JOIN company_openings co ON ic.id = co.company_id
                JOIN internship_jobs ij ON ic.id = ij.company_id
                WHERE ic.status = 'approved'
                  AND co.semester = %s
                  AND co.is_open = TRUE
                  AND ij.is_active = TRUE
                ORDER BY ic.company_name, ij.title
            """, (current_semester_code,))
        else:
            cursor.execute("""
                SELECT 
                    ic.id AS company_id,
                    ic.company_name,
                    ic.description AS company_description,
                    ic.location AS company_address,
                    ij.id AS job_id,
                    ij.title AS job_title,
                    ij.description AS job_description,
                    ij.period AS job_period,
                    ij.work_time AS job_work_time,
                    ij.remark AS job_remark
                FROM internship_companies ic
                JOIN internship_jobs ij ON ic.id = ij.company_id
                WHERE 1=0
            """)
        companies_jobs = cursor.fetchall()
        
        if not companies_jobs:
            cursor.execute("SELECT COUNT(*) as count FROM internship_jobs WHERE is_active = TRUE")
            job_count = cursor.fetchone().get('count', 0)
            if job_count == 0:
                return jsonify({
                    "success": False,
                    "error": "目前系統中沒有可用的職缺，請聯繫管理員新增職缺。"
                }), 400
            return jsonify({
                "success": False,
                "error": "目前沒有可選的公司和職缺組合。"
            }), 400
        
        companies_info = {}
        company_name_to_id = {}
        job_by_id = {}
        job_by_company_title = {}
        job_title_index = {}

        for item in companies_jobs:
            company_id = item['company_id']
            if company_id not in companies_info:
                companies_info[company_id] = {
                    'company_id': company_id,
                    'company_name': item['company_name'],
                    'company_description': item['company_description'] or '',
                    'company_address': item['company_address'] or '',
                    'jobs': []
                }
                company_name_to_id[item['company_name'].strip()] = company_id
            
            job_payload = {
                'job_id': item['job_id'],
                'job_title': item['job_title'],
                'job_description': item['job_description'] or '',
                'job_period': item['job_period'] or '',
                'job_work_time': item['job_work_time'] or '',
                'job_remark': item['job_remark'] or ''
            }
            combined_job = {**job_payload, 'company_id': company_id, 'company_name': item['company_name']}
            companies_info[company_id]['jobs'].append(job_payload)
            job_by_id[item['job_id']] = combined_job
            normalized_title = (item['job_title'] or '').strip().lower()
            job_by_company_title[(company_id, normalized_title)] = combined_job
            if normalized_title:
                job_title_index.setdefault(normalized_title, []).append(combined_job)
        
        companies_text = ""
        for company in companies_info.values():
            jobs_text = "\n".join([
                f"  - 職缺ID: {job['job_id']}, 職缺名稱: {job['job_title']}, "
                f"描述: {job['job_description']}, 實習期間: {job['job_period']}, "
                f"工作時間: {job['job_work_time']}, 備註: {job['job_remark']}"
                for job in company['jobs']
            ])
            companies_text += f"""
公司ID: {company['company_id']}
公司名稱: {company['company_name']}
公司描述: {company['company_description']}
公司地址: {company['company_address']}
職缺列表:
{jobs_text}
---
"""

        has_resume_data = bool(resume_text.strip())
        has_grades_data = bool(grades_text.strip())
        
        prompt = f"""
### 任務目標
你是專業實習顧問。請根據以下【真實資料庫數據】，為學生推薦最匹配的 5 個實習志願。

### 學生背景資料
1. 【自傳與技能】：{resume_text}
2. 【成績單摘要】：{grades_text}

### 可選公司與職缺
{companies_text}

### 輸出規範 (JSON)
每個 "reason" 必須簡潔地包含以下三點（禁止廢話）：
- 引用自傳中的 [興趣/經驗]。
- 引用成績單中的 [具體科目成績]。
- 引用 [證照名稱]。

請直接輸出 JSON，格式如下：
{{
  "recommendations": [
    {{
      "order": 1,
      "company_id": ID,
      "job_id": ID,
      "company_name": "名稱",
      "job_title": "職稱",
      "reason": "直接引述資料的推薦理由"
    }}
  ]
}}
"""
        response = model.generate_content(
            prompt,
            generation_config={
                "response_mime_type": "application/json",
                "temperature": 0.2
            }
        )

        ai_response_text = response.text.strip()
        if ai_response_text.startswith('```json'):
            ai_response_text = ai_response_text[7:]
        if ai_response_text.startswith('```'):
            ai_response_text = ai_response_text[3:]
        if ai_response_text.endswith('```'):
            ai_response_text = ai_response_text[:-3]
        ai_response_text = ai_response_text.strip()

        def try_parse_json(raw_text: str):
            try:
                return json.loads(raw_text)
            except json.JSONDecodeError:
                return None

        recommendations_data = try_parse_json(ai_response_text)
        if recommendations_data is None:
            first_brace = ai_response_text.find('{')
            last_brace = ai_response_text.rfind('}')
            parsed = None
            if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
                possible_json = ai_response_text[first_brace:last_brace+1]
                parsed = try_parse_json(possible_json)
            if parsed is None:
                recommendations = []
            else:
                recommendations_data = parsed
                recommendations = recommendations_data.get('recommendations', [])
        else:
            recommendations = recommendations_data.get('recommendations', [])

        valid = []
        for rec in recommendations:
            cid = rec.get('company_id')
            jid = rec.get('job_id')
            rec_company_name = (rec.get('company_name') or '').strip()
            rec_job_title = (rec.get('job_title') or '').strip()
            matched_job = None

            try:
                jid_int = int(str(jid)) if jid is not None and str(jid).strip().isdigit() else None
            except ValueError:
                jid_int = None

            try:
                cid_int = int(str(cid)) if cid is not None and str(cid).strip().isdigit() else None
            except ValueError:
                cid_int = None

            if jid_int and jid_int in job_by_id:
                job_info = job_by_id[jid_int]
                if not cid_int or cid_int == job_info['company_id']:
                    matched_job = job_info
                else:
                    matched_job = None

            if not matched_job and rec_job_title:
                normalized_title = rec_job_title.lower()
                if not cid_int and rec_company_name:
                    cid_int = company_name_to_id.get(rec_company_name)
                if cid_int:
                    key = (cid_int, normalized_title)
                    if key in job_by_company_title:
                        matched_job = job_by_company_title[key]
                if not matched_job and normalized_title in job_title_index:
                    possible_jobs = job_title_index[normalized_title]
                    if len(possible_jobs) == 1:
                        matched_job = possible_jobs[0]
                if not matched_job:
                    for job in job_by_id.values():
                        job_title_lower = (job['job_title'] or '').lower()
                        if normalized_title and (normalized_title in job_title_lower or job_title_lower in normalized_title):
                            if cid_int and job['company_id'] != cid_int:
                                continue
                            matched_job = job
                            break

            if matched_job:
                valid.append({
                    'order': rec.get('order'),
                    'company_id': matched_job['company_id'],
                    'job_id': matched_job['job_id'],
                    'company_name': matched_job['company_name'],
                    'job_title': matched_job['job_title'],
                    'reason': rec.get('reason', '')
                })

        if not valid:
            fallback_jobs = list(job_by_id.values())
            fallback_jobs.sort(key=lambda j: (j['company_name'], j['job_title']))
            fallback_limit = min(5, len(fallback_jobs))
            if fallback_limit == 0:
                return jsonify({"success": False, "error": "系統目前找不到可用職缺，請稍後再試。"}), 400
            for idx in range(fallback_limit):
                job = fallback_jobs[idx]
                valid.append({
                    'order': idx + 1,
                    'company_id': job['company_id'],
                    'job_id': job['job_id'],
                    'company_name': job['company_name'],
                    'job_title': job['job_title'],
                    'reason': "系統自動推薦：依照您目前的背景與熱門程度優先推薦此職缺。"
                })

        return jsonify({"success": True, "recommendations": valid})

    except Exception as e:
        traceback.print_exc()
        error_str = str(e)
        if "429" in error_str or "quota" in error_str.lower() or "Quota exceeded" in error_str:
            retry_seconds = None
            if "retry in" in error_str.lower() or "retry_delay" in error_str.lower():
                retry_match = re.search(r'retry in ([\d.]+)s', error_str, re.IGNORECASE)
                if retry_match:
                    retry_seconds = int(float(retry_match.group(1)))
            
            error_message = "AI 服務目前使用量已達上限（免費層級每日限制 20 次）"
            if retry_seconds:
                error_message += f"，請在 {retry_seconds} 秒後再試"
            else:
                error_message += "，請稍後再試或明天再使用此功能"
            
            return jsonify({
                "success": False, 
                "error": error_message,
                "error_type": "quota_exceeded",
                "retry_after": retry_seconds
            }), 429
        
        return jsonify({"success": False, "error": f"AI 服務處理失敗: {str(e)}"}), 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# ==========================================================
# API：更新自傳內容
# ==========================================================
@ai_bp.route('/api/update_autobiography', methods=['POST'])
def update_autobiography():
    """
    將 AI 美化後的自傳更新至資料庫
    """
    if 'user_id' not in session or session.get('role') != 'student':
        return jsonify({"success": False, "message": "只有學生可以使用此功能。"}), 403
    
    user_id = session['user_id']
    conn = None
    cursor = None
    
    try:
        data = request.get_json()
        autobiography = data.get('autobiography', '').strip()
        
        if not autobiography:
            return jsonify({"success": False, "message": "自傳內容不能為空。"}), 400
        
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT username FROM users WHERE id=%s", (user_id,))
        user_result = cursor.fetchone()
        if not user_result:
            return jsonify({"success": False, "message": "找不到使用者資訊。"}), 404
        
        student_id = user_result["username"]
        
        cursor.execute("""
            INSERT INTO Student_Info (StuID, Autobiography, UpdatedAt)
            VALUES (%s, %s, NOW())
            ON DUPLICATE KEY UPDATE
                Autobiography = VALUES(Autobiography),
                UpdatedAt = NOW()
        """, (student_id, autobiography))
        
        conn.commit()
        return jsonify({"success": True, "message": "自傳已成功更新。"})
        
    except Exception as e:
        traceback.print_exc()
        if conn:
            conn.rollback()
        return jsonify({"success": False, "message": f"更新失敗: {str(e)}"}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# ==========================================================
# 📷 OCR 與文件識別工具函數
# ==========================================================

# 本系「必修科目」完整參考清單（標準課名 + 應修學分，供 OCR 歸併；學分與 Excel 範本不一致時以 Excel 為準）
DEPARTMENT_CORE_REFERENCE_COURSES = [
    {"name": "系統分析與設計", "credits": "3"},
    {"name": "資訊科技", "credits": "2"},
    {"name": "資訊科技進階", "credits": "2"},
    {"name": "計算機網路", "credits": "2"},
    {"name": "JAVA程式語言", "credits": "3"},
    {"name": "會計概論", "credits": "3"},
    {"name": "會計學", "credits": "3"},
    {"name": "資料庫伺服器管理與實作", "credits": "2"},
    {"name": "資料庫管理實務(SQL)", "credits": "2"},
    {"name": "管理學", "credits": "2"},
    {"name": "作業系統", "credits": "2"},
    {"name": "行銷管理", "credits": "2"},
    {"name": "資料結構", "credits": "2"},
    {"name": "商品攝影與後製", "credits": "2"},
    {"name": "微積分", "credits": "3"},
    {"name": "微電影製作", "credits": "2"},
    {"name": "經濟學", "credits": "2"},
    {"name": "數位整合行銷", "credits": "2"},
    {"name": "中英文輸入", "credits": "2"},
    {"name": "行銷企劃書撰寫", "credits": "2"},
    {"name": "程式設計", "credits": "3"},
    {"name": "創意機器人", "credits": "2"},
    {"name": "行動網頁程式開發", "credits": "2"},
    {"name": "數位化資料處理", "credits": "2"},
    {"name": "辦公室自動化", "credits": "2"},
    {"name": "硬體裝修", "credits": "2"},
    {"name": "網頁設計", "credits": "2"},
    {"name": "商業套裝軟體", "credits": "2"},
    {"name": "電腦繪圖與動畫", "credits": "2"},
    {"name": "統計學", "credits": "3"},
]


def _reference_course_names_set():
    return {row["name"] for row in DEPARTMENT_CORE_REFERENCE_COURSES}


def _reference_credits_by_name():
    return {row["name"]: str(row["credits"]).strip() for row in DEPARTMENT_CORE_REFERENCE_COURSES}


def _format_core_reference_for_prompt():
    lines = []
    for row in DEPARTMENT_CORE_REFERENCE_COURSES:
        lines.append(f"- 「{row['name']}」（應修學分：{row['credits']}）")
    return "\n".join(lines)


def build_transcript_json_prompt():
    """組裝含固定課程清單的 JSON 模式提示詞。"""
    ref_block = _format_core_reference_for_prompt()
    allowed_names = "、".join([f'「{r["name"]}」' for r in DEPARTMENT_CORE_REFERENCE_COURSES])

    return f"""
你是成績單 OCR 助手。請只閱讀圖片內容，輸出**一段合法 JSON**（不要 Markdown、不要程式碼區塊、不要註解、不要多餘說明）。

【本系參考：專業核心科目與應修學分清單（僅用於「課名相似時」歸併成標準字串）】
以下課名若與圖上科目**明顯為同一門**（略稱、空格、全半形、括號寫法等），請將 `name` 設為清單中的**完整標準課名**（須與清單字元完全一致），以利與 Excel 比對。
**圖上所有其他必修科目**（不在清單內者）也必須逐列輸出：`name` 請填圖上可讀之**完整科目名稱**（與成績單一致或極接近），**絕對不可省略、不可因不在清單就丟棄**。
{ref_block}
可歸併時的標準課名僅能是：{allowed_names}。

JSON 物件鍵名必須完全一致：
{{
  "confidence": <0-100 的整數，代表本次整體辨識信心>,
  "transcript_has_credits_column": <boolean，圖片上是否「有」學分相關欄位或表頭>,
  "transcript_has_grade_column": <boolean，圖片上是否「有」成績／等第／分數欄位或表頭>,
  "courses": [
    {{
      "name": "<必填；與 Excel／學校成績單比對用。能對應清單則用清單全名；否則用圖上該列科目完整名稱>",
      "transcript_label": <字串；圖上該列實際顯示的科目文字，盡量與圖片一致；若與 name 相同可填相同字串或 null>,
      "credits": <字串；若圖上有學分列且該格可讀則填圖上數字（可含「2/2」等形式）。若圖上無學分列或該格空白，且 name 為清單中某一門，則可填該課**清單應修學分**。否則 null>,
      "grade": <字串；僅能從圖上成績欄讀取。若無成績欄或看不清則 null>
    }}
  ]
}}

規則（務必遵守）：
1. 只收錄「必修」科目；若表上有必修／選修標示，非必修一律不要輸出。
2. 只收錄第一學年～第三學年（依圖上學期／學年欄位判斷）；第四學年（含）以後不要輸出；無法判斷學年的列不要輸出。
3. **課名**：`name` 禁止自創與圖片無關的課程；除清單歸併外，應忠實反映圖上科目（含行銷、作業系統、攝影等任何出現在圖上的必修課）。
4. **成績**：`grade` 只允許來自圖片，禁止臆測。
5. 禁止補全圖上沒有的文字；看不清的欄位用 null。若同一門必修課分上下學期（或分段）各出現一列，可分行輸出；系統會自動取成績較高者合併為該科目一筆。
6. 若整張圖不是成績單或過於模糊：courses 輸出 []，confidence 請給 0～30。
7. 若成績單根本沒有「學分」欄，transcript_has_credits_column 必為 false；此時若 name 恰為清單課名，可把 `credits` 填清單應修學分；若 name 非清單課且圖上無學分，則 credits 為 null。
8. 若成績單根本沒有「成績／等第」欄，transcript_has_grade_column 必為 false，且每筆 grade 必為 null。
9. courses 內每筆都必須同時包含鍵 name、transcript_label、credits、grade（值可為 null），不要省略鍵。
""".strip()


def _transcript_courses_to_markdown(courses):
    """將結構化課程列表轉成前端既有 Markdown 表格。"""
    lines = [
        "| 科目名稱 | 學分 | 成績 |",
        "| --- | --- | --- |",
    ]
    for c in courses:
        name = str(c.get("name", "")).replace("|", "\\|").strip()
        cr = c.get("credits")
        gr = c.get("grade")
        cr_disp = "" if cr is None or cr == "" else str(cr).replace("|", "\\|")
        gr_disp = "" if gr is None or gr == "" else str(gr).replace("|", "\\|")
        lines.append(f"| {name} | {cr_disp} | {gr_disp} |")
    return "\n".join(lines)


def _compact_course_name_key(name):
    """與前端 normalizeCourseNameKeyForTable 一致：去空白、小寫。"""
    if not name:
        return ""
    return re.sub(r"[\s\u3000]+", "", str(name).strip().lower())


def _course_base_merge_key(name):
    """
    同一門課分上下學期／(一)(二) 等時的基底鍵，用於合併列並取較高成績（如會計概論、硬體裝修）。
    """
    s = _compact_course_name_key(name)
    if not s:
        return ""
    patterns = (
        r"\([上下]\)$",
        r"（[上下]）$",
        r"\([一二三四五六七八九十0-9]+\)$",
        r"（[一二三四五六七八九十0-9]+）$",
        r"\([IiⅠⅡⅢⅣ]+\)$",
        r"（[IiⅠⅡⅢⅣ]+）$",
        r"上學期$",
        r"下學期$",
        r"第[一二三四1-4]學期$",
    )
    for _ in range(24):
        hit = False
        for p in patterns:
            ns = re.sub(p, "", s, flags=re.UNICODE)
            if ns != s:
                s = ns
                hit = True
                break
        if not hit:
            break
    return s


def _transcript_grade_rank_for_merge(grade_str):
    """成績可比較數值（越大越好），供上下學期合併取高。"""
    if grade_str is None:
        return float("-inf")
    s = str(grade_str).strip().replace("分", "")
    if not s or s.lower() == "null":
        return float("-inf")
    letter_to_score = {"優": 95.0, "甲": 85.0, "乙": 75.0, "丙": 65.0, "丁": 55.0}
    if s in letter_to_score:
        return letter_to_score[s]
    if re.match(r"^(通過|合格)$", s, re.I):
        return 70.0
    cleaned = re.sub(r"[^\d.\-]", "", s)
    if cleaned:
        try:
            return float(cleaned)
        except ValueError:
            pass
    letter_map = {
        "a+": 98.0,
        "a": 95.0,
        "a-": 92.0,
        "b+": 88.0,
        "b": 85.0,
        "b-": 82.0,
        "c+": 78.0,
        "c": 75.0,
        "c-": 72.0,
        "d+": 68.0,
        "d": 65.0,
        "d-": 62.0,
        "e": 55.0,
        "f": 40.0,
    }
    k = s.lower()
    return letter_map[k] if k in letter_map else float("-inf")


def _merge_transcript_rows_same_course_semester(rows):
    """
    辨識結果中同一基底課名多列（上下學期等）合併為一筆，成績取較高者；維持原列順序（以首列代表）。
    """
    if not isinstance(rows, list) or len(rows) < 2:
        return rows

    buckets = {}
    bucket_order = []
    for idx, row in enumerate(rows):
        if not isinstance(row, dict):
            continue
        nm = row.get("name") or ""
        bkey = _course_base_merge_key(nm) or f"__single_{idx}__"
        if bkey not in buckets:
            buckets[bkey] = []
            bucket_order.append(bkey)
        buckets[bkey].append((idx, row))

    merged_by_key = {}
    for bkey in bucket_order:
        items = buckets[bkey]
        if len(items) < 2:
            merged_by_key[bkey] = None
            continue
        best_row = items[0][1]
        best_r = _transcript_grade_rank_for_merge(best_row.get("grade"))
        merged_grade = ""
        for _, r in items:
            g = r.get("grade")
            gs = "" if g is None else str(g).strip()
            if gs and gs.lower() != "null":
                if not merged_grade:
                    merged_grade = gs
                elif _transcript_grade_rank_for_merge(gs) > _transcript_grade_rank_for_merge(merged_grade):
                    merged_grade = gs
            tr = _transcript_grade_rank_for_merge(g)
            if tr > best_r:
                best_r = tr
                best_row = r
            elif tr == best_r and tr == float("-inf"):
                pass

        merged = dict(best_row)
        merged["name"] = best_row.get("name") or merged.get("name")
        merged["grade"] = merged_grade if merged_grade else ""
        thg_any = any(bool(it[1].get("compare_grade")) for it in items)
        merged["compare_grade"] = bool(thg_any) and bool(str(merged.get("grade") or "").strip())

        raw_parts = []
        for _, r in items:
            o = (r.get("ocr_raw_name") or "").strip()
            if o and o not in raw_parts:
                raw_parts.append(o)
        if raw_parts:
            merged["ocr_raw_name"] = "；".join(raw_parts)

        merged_by_key[bkey] = merged

    new_rows = []
    emitted_multi = set()
    for idx, row in enumerate(rows):
        if not isinstance(row, dict):
            new_rows.append(row)
            continue
        bkey = _course_base_merge_key(row.get("name") or "") or f"__single_{idx}__"
        m = merged_by_key.get(bkey)
        if m is None:
            new_rows.append(row)
            continue
        if bkey in emitted_multi:
            continue
        emitted_multi.add(bkey)
        new_rows.append(m)
    return new_rows


def _normalize_transcript_json_payload(parsed):
    """
    將模型 JSON 轉成前端比對用課程列表，並附帶「是否要比對學分／成績」旗標。
    僅在圖上確實有該欄位且該列有值時才 compare_* = True。
    支援：name + transcript_label；舊版 canonical_name（在清單內時優先作為標準名）。
    """
    if not isinstance(parsed, dict):
        return [], {}

    raw_list = parsed.get("courses")
    if not isinstance(raw_list, list):
        return [], {}

    allowed = _reference_course_names_set()
    ref_credits = _reference_credits_by_name()

    thc = parsed.get("transcript_has_credits_column")
    thg = parsed.get("transcript_has_grade_column")
    if not isinstance(thc, bool):
        thc = any(
            (c.get("credits") is not None and str(c.get("credits", "")).strip() != "")
            for c in raw_list
            if isinstance(c, dict)
        )
    if not isinstance(thg, bool):
        thg = any(
            (c.get("grade") is not None and str(c.get("grade", "")).strip() != "")
            for c in raw_list
            if isinstance(c, dict)
        )

    out = []
    for c in raw_list:
        if not isinstance(c, dict):
            continue

        legacy_canonical = (c.get("canonical_name") or "").strip()
        primary_name = (c.get("name") or "").strip()
        label_raw = c.get("transcript_label")
        label = "" if label_raw is None else str(label_raw).strip()

        # 舊版 JSON 可能只有 canonical_name：在清單內則採用，否則當作圖上課名
        if legacy_canonical in allowed:
            resolved_name = legacy_canonical
        elif primary_name:
            resolved_name = primary_name
        elif legacy_canonical:
            resolved_name = legacy_canonical
        elif label:
            resolved_name = label
        else:
            continue

        # 供前端顯示「圖上 → 標準」：圖上原文與比對用 name 不同時保留
        ocr_raw = (label if (label and label != resolved_name) else "") or ""

        cred_raw = c.get("credits", None)
        grade_raw = c.get("grade", None)
        cred_str = None if cred_raw is None else str(cred_raw).strip()
        grade_str = None if grade_raw is None else str(grade_raw).strip()
        if cred_str == "":
            cred_str = None
        if grade_str == "":
            grade_str = None

        # 圖上無學分列時：僅當課名為清單標準名時，補清單應修學分（供顯示；不比對學分）
        if not thc and cred_str is None and resolved_name in ref_credits:
            cred_str = ref_credits.get(resolved_name)

        compare_credits = bool(thc) and cred_str is not None
        compare_grade = bool(thg) and grade_str is not None

        out.append({
            "name": resolved_name,
            "credits": cred_str if cred_str is not None else "",
            "grade": grade_str if grade_str is not None else "",
            "ocr_raw_name": ocr_raw,
            "compare_credits": compare_credits,
            "compare_grade": compare_grade,
            "transcript_has_credits_column": bool(thc),
            "transcript_has_grade_column": bool(thg),
        })

    out = _merge_transcript_rows_same_course_semester(out)

    meta = {
        "transcript_has_credits_column": bool(thc),
        "transcript_has_grade_column": bool(thg),
        "reference_course_count": len(DEPARTMENT_CORE_REFERENCE_COURSES),
    }
    return out, meta


def _gemini_try_transcript_json(model_instance, img_data, mimetype):
    """請模型以 application/json 回傳；失敗回傳 None。"""
    try:
        generation_config = {"temperature": 0.05, "response_mime_type": "application/json"}
        response = model_instance.generate_content(
            [
                {"mime_type": mimetype or "image/jpeg", "data": img_data},
                build_transcript_json_prompt(),
            ],
            generation_config=generation_config,
        )
        if not response or not getattr(response, "text", None):
            return None
        raw = response.text.strip()
        parsed = json.loads(raw)
        if not isinstance(parsed, dict):
            return None
        return parsed
    except Exception as e:
        print(f"⚠️ JSON 模式辨識失敗: {e}")
        return None


def _gemini_try_transcript_markdown(model_instance, img_data, mimetype):
    """舊版 Markdown 表格輸出（備援）。"""
    ref_block = _format_core_reference_for_prompt()
    prompt = f"""
請嚴格辨識圖片中的文字，並將「課程名稱」、「學分」、「成績」整理成 Markdown 表格。

【本系參考：專業核心科目與應修學分清單（僅用於課名相似時歸併）】
圖上科目若與清單中某一門**明顯為同一門**（例如「JAVA 程式」→「JAVA程式語言」），表格「科目名稱」欄請用**清單完整標準課名**。
**其餘所有出現在圖上的第一～三學年必修科目**（即使不在清單內）也必須輸出，科目名稱欄請填**圖上完整課名**（與成績單一致），不可省略。
學分請優先使用圖上數字（可含「2/2」）；若圖上無學分列且該列已歸併為清單課名，可填清單應修學分。成績僅能來自圖片。
{ref_block}

【重要規則】
1. **只抓必修**：僅辨識並輸出「必修」科目。請忽略：選修、通識、體育、國防、勞作、軍訓、語文（若標示為選修）、以及其他非必修欄位。若成績單有「必修/選修」欄位，只保留標示為「必修」的列。
2. **只抓第一學年到第三學年**：僅辨識成績單上「第一學年」、「第二學年」、「第三學年」所對應學期的科目（例如 1101、1102、1111、1112、1121、1122 等，依成績單實際學期標示）。排除第四學年（含）以後、以及無法判斷學年的列。
3. **絕對不可無中生有**：只能輸出圖片上真正存在的文字，禁止預測或補全沒看到的課程。
4. **科目名稱**：每一列必修課都要輸出；能歸併清單者用標準名，否則用圖上課名全文。
5. **格式統一**：表格欄位請固定為 `| 科目名稱 | 學分 | 成績 |`。
6. **排除雜訊**：請忽略學校名稱、個人資料、蓋章、頁碼、排名、總學分、平均成績、操行成績、備註欄位等所有非課程本身的內容。
7. **清晰度與可信度判斷**：
   - 如果圖片模糊、無法辨識文字或根本不是成績單，請輸出：「[無法辨識] 請重新拍攝更清晰的圖片上傳。」
   - **請在輸出的最後一行（單獨一行）評估本次辨識的準確率（信心度），格式為：「信心度：[0-100]」**
"""
    response = model_instance.generate_content(
        [
            {"mime_type": mimetype or "image/jpeg", "data": img_data},
            prompt,
        ]
    )
    if not response or not response.text:
        return None, None

    gemini_text = response.text.strip()
    gemini_conf = 95.0
    confidence_match = re.search(r"信心度：\s*(\d+)", gemini_text.split("\n")[-1])
    if confidence_match:
        try:
            gemini_conf = float(confidence_match.group(1))
            gemini_text = "\n".join(gemini_text.split("\n")[:-1]).strip()
        except ValueError:
            pass
    return gemini_text, gemini_conf


def perform_ocr_on_file(file_storage):
    """
    對上傳的檔案 (Image/PDF) 進行 OCR 或文字識別。
    """
    filename = file_storage.filename or ""
    
    # 計算檔案大小
    file_storage.stream.seek(0, os.SEEK_END)
    size_bytes = file_storage.stream.tell()
    file_storage.stream.seek(0)
    size_kb = round(size_bytes / 1024, 1) if size_bytes is not None else 0
    
    ext = os.path.splitext(filename)[1].lower()
    
    # 1. PDF 處理：直接抽取文字
    if ext == ".pdf":
        return {
            "success": False, # 改為 False 以明確告知失敗
            "filename": filename,
            "size_kb": size_kb,
            "text": "",
            "confidence": None,
            "message": "伺服器未安裝 PDF 解析套件，無法處理 PDF 檔案。"
        }

    # 2. 圖片處理：優先使用 Google Gemini AI（先 JSON 結構化，失敗再 Markdown 備援）
    gemini_text = ""
    gemini_conf = None
    use_gemini = False
    structured_courses = None
    transcript_meta = None

    local_api_key = os.getenv('GEMINI_API_KEY')

    if local_api_key:
        try:
            file_storage.stream.seek(0)
            img_data = file_storage.read()
            file_storage.stream.seek(0)

            genai.configure(api_key=local_api_key)

            candidate_models = [
                'gemini-2.5-flash',
                'gemini-1.5-flash',
                'gemini-1.5-pro',
                'gemini-2.0-flash-exp',
            ]
            mimetype = file_storage.mimetype or 'image/jpeg'

            for model_name in candidate_models:
                try:
                    print(f"OCR 嘗試使用模型: {model_name}")
                    model_instance = genai.GenerativeModel(model_name)

                    parsed = _gemini_try_transcript_json(model_instance, img_data, mimetype)
                    if parsed is not None:
                        structured_courses, transcript_meta = _normalize_transcript_json_payload(parsed)
                        try:
                            conf_raw = parsed.get('confidence', 85)
                            gemini_conf = float(conf_raw) if conf_raw is not None else 85.0
                        except (TypeError, ValueError):
                            gemini_conf = 85.0
                        gemini_text = _transcript_courses_to_markdown(structured_courses)
                        use_gemini = True
                        print(f"✅ 使用 Gemini AI ({model_name}) JSON 模式完成圖片辨識，courses={len(structured_courses)}")
                        break

                    md_text, md_conf = _gemini_try_transcript_markdown(model_instance, img_data, mimetype)
                    if md_text:
                        gemini_text = md_text
                        gemini_conf = md_conf
                        structured_courses = None
                        transcript_meta = None
                        use_gemini = True
                        print(f"✅ 使用 Gemini AI ({model_name}) Markdown 備援完成圖片辨識")
                        break
                except Exception as inner_e:
                    print(f"⚠️ 模型 {model_name} 嘗試失敗: {inner_e}")
                    continue

            if not use_gemini:
                print("❌ 所有 Gemini 模型皆嘗試失敗")

        except Exception as e:
            print(f"⚠️ Gemini OCR 系統性錯誤: {e}")
            file_storage.stream.seek(0)

    if use_gemini and gemini_text:
        payload = {
            "success": True,
            "filename": filename,
            "size_kb": size_kb,
            "text": gemini_text,
            "confidence": gemini_conf,
        }
        if structured_courses is not None:
            payload["courses"] = structured_courses
        if transcript_meta:
            payload["transcript_meta"] = transcript_meta
        return payload

    # 3. 圖片處理：AI 辨識失敗
    return {
        "success": False, 
        "filename": filename,
        "size_kb": size_kb,
        "text": "",
        "confidence": None,
        "message": "AI 辨識失敗 (所有 Gemini 模型皆無法使用)，請稍後再試。"
    }

def create_ocr_docx(file_storage):
    """
    執行 OCR 並生成 Word 文件 (回傳文件路徑與檔名)。
    """
    result = perform_ocr_on_file(file_storage)
    ocr_text = result.get("text", "")
    avg_conf = result.get("confidence")
    filename = result.get("filename", "document")
    size_kb = result.get("size_kb", 0)
    
    doc = Document()
    doc.add_heading('成績 AI 識別結果', level=1)
    doc.add_paragraph(f"檔名：{filename}")
    doc.add_paragraph(f"檔案大小：約 {size_kb} KB")

    if ocr_text:
        doc.add_paragraph("")
        doc.add_heading('辨識文字', level=2)
        
        lines = ocr_text.strip().split('\n')
        has_table_structure = any(line.strip().startswith('|') and '---' in line for line in lines)
        
        if has_table_structure:
            try:
                table_data = []
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith('|') and stripped.endswith('|'):
                        if '---' in stripped: 
                            continue
                        cells = [c.strip() for c in stripped.strip('|').split('|')]
                        table_data.append(cells)
                    else:
                        if table_data:
                            _write_table_to_doc(doc, table_data)
                            table_data = []
                        if stripped:
                            doc.add_paragraph(stripped)
                if table_data:
                    _write_table_to_doc(doc, table_data)
            except Exception as e:
                print(f"Word 表格生成失敗: {e}")
                doc.add_paragraph(ocr_text)
        else:
            doc.add_paragraph(ocr_text)
            
        if avg_conf is not None:
            doc.add_paragraph(f"平均信心度：約 {avg_conf}%")
            
    else:
        doc.add_paragraph("")
        doc.add_paragraph(result.get("message", "未偵測到文字。"))

    tmp_dir = os.path.join(current_app.root_path, 'tmp')
    os.makedirs(tmp_dir, exist_ok=True)
    safe_name = secure_filename(os.path.splitext(filename)[0] or "score")
    docx_filename = f"{safe_name}_score_ocr.docx"
    docx_path = os.path.join(tmp_dir, docx_filename)
    doc.save(docx_path)
    
    return docx_path, docx_filename

def _write_table_to_doc(doc, table_data):
    if not table_data:
        return
    max_cols = max(len(row) for row in table_data)
    if max_cols > 0:
        table = doc.add_table(rows=0, cols=max_cols)
        table.style = 'Table Grid'
        for row_data in table_data:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = cell_text
