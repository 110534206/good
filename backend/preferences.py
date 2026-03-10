from flask import Blueprint, render_template, request, jsonify, session, send_file, redirect, url_for, flash
from config import get_db
from datetime import datetime
import traceback
from collections import defaultdict
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import io
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from semester import get_current_semester_code, get_current_semester_id, get_flow_semester_id, get_semester_code_for_company_openings, is_student_in_application_phase


def _get_active_semester_year(cursor):
    """取得當前學年（如 113），用於計算年級。"""
    cursor.execute("SELECT code FROM semesters WHERE is_active = 1 LIMIT 1")
    row = cursor.fetchone()
    if not row or row.get('code') is None:
        cursor.execute("SELECT code FROM semesters WHERE code IS NOT NULL AND code != '' ORDER BY code DESC LIMIT 1")
        row = cursor.fetchone()
    raw = row.get('code') if row else None
    if raw is None:
        return None
    if isinstance(raw, int):
        return raw // 10 if raw >= 100 else None
    code = str(raw).strip()
    if len(code) >= 3:
        try:
            return int(code[:3])
        except (TypeError, ValueError):
            pass
    return None


def _format_class_display(dept, cname, admission_yr, active_year):
    """班級顯示格式：科系 年級班序，如 資管科 四孝。"""
    dept_str = (dept or '').strip()
    name_str = (cname or '').strip()
    if not name_str:
        return dept_str or '—'
    if admission_yr is None or active_year is None:
        return f"{dept_str} {name_str}".strip() if dept_str else name_str
    try:
        ay_int = int(admission_yr)
    except (TypeError, ValueError):
        return f"{dept_str} {name_str}".strip() if dept_str else name_str
    grade_num = active_year - ay_int + 1
    grade_labels = ('一', '二', '三', '四', '五', '六')
    grade_char = grade_labels[grade_num - 1] if 1 <= grade_num <= 6 else (str(grade_num) if grade_num > 0 else '')
    return f"{dept_str} {grade_char}{name_str}".strip() if grade_char else (f"{dept_str} {name_str}".strip() if dept_str else name_str)
from notification import create_notification
preferences_bp = Blueprint("preferences_bp", __name__)

# -------------------------
# 輔助函數：處理志願序填寫截止時間後的狀態自動更新
# -------------------------
def update_preference_status_after_deadline(cursor, conn):
    """
    志願序填寫截止時間後，自動更新狀態：
    將所有 submitted 狀態的志願序自動改為 approved（班導審核通過）
    
    返回: (is_deadline_passed: bool, updated_count: int)
    """
    try:
        from semester import get_current_semester_deadline
        # 檢查志願序填寫截止時間：優先學期流程表 internship_flows，無則 fallback 公告
        now = datetime.now()
        preference_deadline = get_current_semester_deadline(cursor, 'preference')
        if preference_deadline is None:
            cursor.execute("""
                SELECT end_time 
                FROM announcement 
                WHERE title LIKE '[作業]%填寫志願序截止時間' AND is_published = 1
                ORDER BY created_at DESC 
                LIMIT 1
            """)
            deadline_result = cursor.fetchone()
            if deadline_result and deadline_result.get('end_time'):
                deadline = deadline_result['end_time']
                if isinstance(deadline, datetime):
                    preference_deadline = deadline
                else:
                    try:
                        preference_deadline = datetime.strptime(str(deadline), '%Y-%m-%d %H:%M:%S')
                    except Exception:
                        preference_deadline = datetime.strptime(str(deadline), '%Y-%m-%d %H:%M')
        is_preference_deadline_passed = preference_deadline is not None and now > preference_deadline
        
        # 如果已經過了截止時間，執行狀態更新
        if is_preference_deadline_passed:
            # 將所有 submitted 狀態的志願序自動改為 approved（班導審核通過）
            cursor.execute("""
                UPDATE student_preferences 
                SET status = 'approved', updated_at = NOW()
                WHERE status = 'submitted'
            """)
            updated_count = cursor.rowcount
            
            if updated_count > 0:
                conn.commit()
                print(f"✅ 志願序填寫截止時間已過，已將 {updated_count} 筆志願序狀態從 'submitted' 改為 'approved'（班導審核通過）")
            
            return is_preference_deadline_passed, updated_count
        
        return False, 0
    except Exception as e:
        print(f"❌ 更新志願序狀態錯誤: {e}")
        traceback.print_exc()
        return False, 0

# -------------------------
# 共用：取得班級學生志願（與欄位）
# -------------------------
def get_class_preferences(cursor, class_id):
    """
    依照你原本 schema 回傳類似的欄位。
    回傳 rows: student_id, student_name, student_number, preference_order, company_name, job_title, submitted_at,
                 company_address, contact_name, contact_phone, contact_email
    """
    cursor.execute("""
        SELECT 
            u.id AS student_id,
            u.name AS student_name,
            u.username AS student_number,
            sp.preference_order,
            sp.submitted_at,
            ic.id AS company_id,
            ic.company_name,
            ic.company_address,
            ic.contact_name,
            ic.contact_phone,
            ic.contact_email,
            ij.id AS job_id,
            ij.title AS job_title
        FROM users u
        LEFT JOIN student_preferences sp ON u.id = sp.student_id
        LEFT JOIN internship_companies ic ON sp.company_id = ic.id
        LEFT JOIN internship_jobs ij ON sp.job_id = ij.id
        WHERE u.class_id = %s AND u.role = 'student'
        ORDER BY u.name, sp.preference_order
    """, (class_id,))
    return cursor.fetchall()

# -------------------------
# 志願填寫頁面
# -------------------------
@preferences_bp.route("/fill_preferences", methods=["GET"])
def fill_preferences_page():
    # 允許未登入/非學生以預覽模式進入；若為學生則僅當前實習學期可填寫
    is_student = ("user_id" in session and session.get("role") == "student")
    student_id = session.get("user_id") if is_student else None
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        if is_student and student_id and not is_student_in_application_phase(cursor, student_id):
            flash("您尚未進入實習流程學期，無法使用填寫志願序功能。", "warning")
            return redirect(url_for("users_bp.student_home"))
        # 1) 取得開放學期開放的公司（下學期沿用上學期，故 1132 時仍以 1131 開放為準）
        openings_semester_code = get_semester_code_for_company_openings(cursor)
        
        if openings_semester_code:
            cursor.execute("""
                SELECT DISTINCT ic.id, ic.company_name AS name
                FROM internship_companies ic
                INNER JOIN company_openings co ON ic.id = co.company_id
                WHERE ic.status = 'approved'
                  AND co.semester = %s
                  AND co.is_open = TRUE
                ORDER BY ic.company_name
            """, (openings_semester_code,))
        else:
            # 如果沒有設定當前學期，返回空列表
            cursor.execute("SELECT id, company_name AS name FROM internship_companies WHERE 1=0")
        
        companies = cursor.fetchall() or []

        # 2) 簡化：不再計算名額，改為取得所有公司的 ID 列表
        # job_slots: { company_id(str): 1, ... } (1表示該公司可選)
        job_slots = {str(c['id']): 1 for c in companies} #

        # 3) 讀取學生已填寫的志願（若有，預覽模式則為空）
        prefs = []
        if is_student:
            cursor.execute("""
                SELECT 
                    sp.preference_order, 
                    sp.company_id, 
                    sp.job_id,
                    ij.title AS job_title 
                FROM student_preferences sp
                JOIN internship_jobs ij ON sp.job_id = ij.id
                WHERE sp.student_id=%s
                ORDER BY sp.preference_order
            """, (student_id,))
            prefs = cursor.fetchall() or []

        submitted = {
        int(p['preference_order']): {
        "company_id": p["company_id"],
        "job_id": p["job_id"],
        "job_title": p["job_title"],
        }
        for p in prefs
        }

        # **重要：填寫頁面使用 /preferences/fill_preferences.html**
        return render_template(
            "preferences/fill_preferences.html",
            companies=companies,
            submitted=submitted,
            job_slots=job_slots, # 僅用於前端 JS 判斷已選公司，不再代表名額
            company_remaining={}, 
            preview=(not is_student)
        )

    except Exception as e:
        traceback.print_exc()
        return "伺服器錯誤", 500

    finally:
        try:
            cursor.close()
            conn.close()
        except Exception:
            pass

# -------------------------
# 取得該公司所有職缺
# -------------------------
@preferences_bp.route("/api/get_jobs_by_company", methods=["GET"])
def get_jobs_by_company():
    company_id = request.args.get("company_id", type=int)
    if not company_id:
        return jsonify({"success": False, "message": "缺少公司 ID"}), 400

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("""
            SELECT id, title 
            FROM internship_jobs 
            WHERE company_id = %s AND is_active = TRUE
        """, (company_id,))
        jobs = cursor.fetchall() or []
        return jsonify({"success": True, "jobs": jobs})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"查詢失敗: {e}"}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# 取得學生自己的志願序
# -------------------------
@preferences_bp.route("/api/get_my_preferences", methods=["GET"])
def get_my_preferences():
    """學生查看自己的志願序（僅當前實習學期學生）"""
    if "user_id" not in session or session.get("role") != "student":
        return jsonify({"success": False, "message": "未授權"}), 403
    student_id = session.get("user_id")
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        if not is_student_in_application_phase(cursor, student_id):
            return jsonify({"success": False, "message": "您尚未進入實習流程學期，無法使用此功能"}), 403
        cursor.execute("""
            SELECT 
                sp.id,
                sp.preference_order, 
                sp.company_id, 
                sp.job_id,
                sp.status,
                sp.submitted_at,
                ic.company_name,
                ij.title AS job_title
            FROM student_preferences sp
            LEFT JOIN internship_companies ic ON sp.company_id = ic.id
            LEFT JOIN internship_jobs ij ON sp.job_id = ij.id
            WHERE sp.student_id = %s
            ORDER BY sp.preference_order
        """, (student_id,))
        preferences = cursor.fetchall() or []
        
        # 格式化日期
        for pref in preferences:
            if isinstance(pref.get('submitted_at'), datetime):
                pref['submitted_at'] = pref['submitted_at'].strftime("%Y-%m-%d %H:%M:%S")
        
        return jsonify({"success": True, "preferences": preferences})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"查詢失敗: {e}"}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# 儲存學生志願
# -------------------------
@preferences_bp.route("/api/save_preferences", methods=["POST"])
def save_preferences():
    # 權限檢查
    if "user_id" not in session or session.get("role") != "student":
        return jsonify({"success": False, "message": "未授權"}), 403
    student_id = session["user_id"]
    data = request.get_json(silent=True) or {}
    preferences = data.get("preferences", [])
    # 基本驗證
    if not preferences:
        return jsonify({"success": False, "message": "請至少選擇一個志願。"}), 400
    MAX_PREFS = 5
    if len(preferences) > MAX_PREFS:
        return jsonify({"success": False, "message": f"最多只能填寫 {MAX_PREFS} 個志願。"}), 400
    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    try:
        if not is_student_in_application_phase(cursor, student_id):
            return jsonify({"success": False, "message": "您尚未進入實習流程學期，無法填寫志願序"}), 403
        # 1) 檢查公司是否重複 - 移除此邏輯，以配合前端的「公司可重複選，職缺互斥」
        selected_job_ids = set() # 用來檢查職缺是否重複，以防萬一
        for p in preferences:
            cid = p.get("company_id")
            jid = p.get("job_id")
            if not cid or not jid:
                return jsonify({"success": False, "message": "每筆志願需包含 company_id 與 job_id。"}), 400
            
            # **重點：檢查職缺是否重複**
            if jid in selected_job_ids:
                return jsonify({"success": False, "message": f"職缺(ID: {jid}) 已在其他志願中選擇，同一職缺只能選擇一次。"}), 400
            selected_job_ids.add(jid)

        # 2) 取得當前學期ID
        current_semester_id = get_current_semester_id(cursor)
        print(f"💾 儲存志願序 - student_id: {student_id}, current_semester_id: {current_semester_id}, preferences數量: {len(preferences)}")
        
        # 3) 刪除學生舊紀錄並插入新志願
        if current_semester_id:
            cursor.execute("DELETE FROM student_preferences WHERE student_id=%s AND semester_id=%s", (student_id, current_semester_id))
            deleted_count = cursor.rowcount
            print(f"🗑️ 刪除舊志願序: {deleted_count} 筆")
        else:
            cursor.execute("DELETE FROM student_preferences WHERE student_id=%s", (student_id,))
            deleted_count = cursor.rowcount
            print(f"🗑️ 刪除舊志願序（無學期限制）: {deleted_count} 筆")

        inserted_count = 0
        for p in preferences:
            pref_order = int(p.get("order"))
            company_id = int(p.get("company_id"))
            job_id = int(p.get("job_id"))

            # 檢查 job_id 是否屬於該公司
            cursor.execute("""
                SELECT title FROM internship_jobs WHERE id=%s AND company_id=%s
            """, (job_id, company_id))
            job_row = cursor.fetchone()
            if not job_row:
                conn.rollback()
                return jsonify({"success": False, "message": f"職缺無效或不屬於該公司：job_id={job_id}, company_id={company_id}"}), 400

            # 確保 job_row 是 dict 結構，以便取出 title
            job_title = job_row.get("title") if isinstance(job_row, dict) else (job_row[0] if isinstance(job_row, tuple) else None)

            if not job_title:
                conn.rollback()
                return jsonify({"success": False, "message": f"無法取得職缺名稱：job_id={job_id}"}), 400

            # 插入志願序（包含 semester_id）
            if current_semester_id:
                cursor.execute("""
                    INSERT INTO student_preferences
                    (student_id, semester_id, preference_order, company_id, job_id, job_title, status, submitted_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    student_id,
                    current_semester_id,
                    pref_order,
                    company_id,
                    job_id,
                    job_title,
                    'submitted',  # 預設狀態為 'submitted'（已提交，待審核）
                    datetime.now()
                ))
                inserted_count += 1
                print(f"✅ 插入志願序 {pref_order}: company_id={company_id}, job_id={job_id}, semester_id={current_semester_id}")
            else:
                # 如果沒有設定當前學期，仍然可以儲存（不包含 semester_id）
                cursor.execute("""
                    INSERT INTO student_preferences
                    (student_id, preference_order, company_id, job_id, job_title, status, submitted_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (
                    student_id,
                    pref_order,
                    company_id,
                    job_id,
                    job_title,
                    'submitted',  # 預設狀態為 'submitted'（已提交，待審核）
                    datetime.now()
                ))
                inserted_count += 1
                print(f"✅ 插入志願序 {pref_order}: company_id={company_id}, job_id={job_id} (無學期)")

        # 4) 提交 transaction
        conn.commit()
        print(f"💾 志願序儲存完成: 共插入 {inserted_count} 筆")
        
        # 5) 驗證資料是否正確寫入
        try:
            if current_semester_id:
                cursor.execute("""
                    SELECT COUNT(*) as count
                    FROM student_preferences
                    WHERE student_id = %s AND semester_id = %s
                """, (student_id, current_semester_id))
            else:
                cursor.execute("""
                    SELECT COUNT(*) as count
                    FROM student_preferences
                    WHERE student_id = %s
                """, (student_id,))
            verify_result = cursor.fetchone()
            saved_count = verify_result.get('count', 0) if verify_result else 0
            print(f"✅ 驗證資料寫入: 資料庫中有 {saved_count} 筆志願序記錄")
            
            if saved_count != inserted_count:
                print(f"⚠️ 警告: 插入 {inserted_count} 筆，但資料庫中只有 {saved_count} 筆")
        except Exception as verify_error:
            print(f"⚠️ 驗證資料時發生錯誤: {verify_error}")
        
        # 6) 通知班導：學生已填寫志願序
        try:
            # 獲取學生資訊（姓名、班級）
            cursor.execute("""
                SELECT u.name AS student_name, u.class_id, c.name AS class_name
                FROM users u
                LEFT JOIN classes c ON u.class_id = c.id
                WHERE u.id = %s
            """, (student_id,))
            student_info = cursor.fetchone()
            
            if student_info:
                student_name = student_info.get('student_name', '學生')
                class_id = student_info.get('class_id')
                class_name = student_info.get('class_name', '')
                
                if class_id:
                    # 查找該班級的班導
                    cursor.execute("""
                        SELECT teacher_id
                        FROM classes_teacher
                        WHERE class_id = %s AND role = 'classteacher'
                        LIMIT 1
                    """, (class_id,))
                    teacher_row = cursor.fetchone()
                    
                    if teacher_row and teacher_row.get('teacher_id'):
                        teacher_id = teacher_row.get('teacher_id')
                        
                        # 創建通知給班導
                        title = "學生志願序已提交"
                        message = f"{student_name} 已提交實習志願序，請前往審核頁面查看。"
                        link_url = "/review_preferences"  # 志願序審核頁面
                        
                        notification_success = create_notification(
                            user_id=teacher_id,
                            title=title,
                            message=message,
                            category="ranking",  # 志願序分類
                            link_url=link_url
                        )
                        
                        if notification_success:
                            print(f"✅ 已通知班導（teacher_id: {teacher_id}）：學生 {student_name} 已提交志願序")
                        else:
                            print(f"⚠️ 通知班導失敗（teacher_id: {teacher_id}）")
                    else:
                        print(f"⚠️ 學生 {student_name} 的班級（class_id: {class_id}）沒有找到班導")
                else:
                    print(f"⚠️ 學生 {student_name} 沒有分配班級（class_id 為空）")
            else:
                print(f"⚠️ 無法找到學生資訊（student_id: {student_id}）")
                
        except Exception as notify_error:
            # 通知失敗不影響主流程，只記錄錯誤
            print(f"⚠️ 通知班導時發生錯誤: {notify_error}")
            traceback.print_exc()
        
        # 返回成功訊息，包含儲存的筆數
        return jsonify({
            "success": True, 
            "message": f"志願序已成功送出（共 {inserted_count} 筆志願）。",
            "inserted_count": inserted_count
        })

    except Exception as e:
        # rollback
        try:
            conn.rollback()
        except Exception:
            pass
        print(f"❌ 儲存志願序時發生錯誤: {e}")
        traceback.print_exc()
        return jsonify({
            "success": False, 
            "message": f"儲存失敗：{str(e)}。請稍後再試或聯絡管理員。"
        }), 500
    finally:
        try:
            cursor.close()
            conn.close()
        except Exception:
            pass

# -------------------------
# API - 選擇角色 (模擬登入)
# -------------------------
@preferences_bp.route('/api/select_role', methods=['POST'])
def select_role():
    data = request.json
    username = data.get("username")
    role = data.get("role")

    conn = get_db()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id FROM users WHERE username=%s AND role=%s", (username, role))
    user = cursor.fetchone()
    cursor.close()
    conn.close()

    if user:
        session["user_id"] = user["id"]
        session["role"] = role
        return jsonify({"success": True})
    else:
        return jsonify({"success": False, "message": "無此角色"}), 404

# -------------------------
# 班導查看志願序
# -------------------------
@preferences_bp.route('/review_preferences')
def review_preferences():
    if 'username' not in session or session.get('role') not in ['teacher', 'director', "class_teacher"]:
        return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導，取得該班導負責的所有班級
        cursor.execute("""
            SELECT ct.class_id
            FROM classes_teacher ct
            WHERE ct.teacher_id = %s AND ct.role = 'classteacher'
        """, (user_id,))
        class_rows = cursor.fetchall()
        if not class_rows:
            return "你不是班導，無法查看志願序", 403

        class_ids = [r['class_id'] for r in class_rows]
        placeholders = ','.join(['%s'] * len(class_ids))

        # 使用流程學期（1131），1132 時仍顯示 1131 的志願；班導僅看「四年級」志願（當屆實習生）
        flow_semester_id = get_flow_semester_id(cursor)
        current_semester_id = flow_semester_id or get_current_semester_id(cursor)
        # 流程學期對應四年級：1131→110 屆，學號 110xxx
        student_id_prefix = None
        if flow_semester_id:
            cursor.execute("SELECT code FROM semesters WHERE id = %s", (flow_semester_id,))
            row = cursor.fetchone()
            if row and row.get('code') and len(str(row['code'])) >= 3:
                try:
                    year_part = int(str(row['code'])[:3])
                    student_id_prefix = year_part - 3  # 113 → 110
                except (ValueError, TypeError):
                    pass
        
        # 檢查志願序填寫截止時間並自動更新狀態（依流程學期）
        is_preference_deadline_passed, update_count = update_preference_status_after_deadline(cursor, conn)
        
        print(f"🔍 班導審核志願序 - class_ids: {class_ids}, flow_semester_id: {flow_semester_id}, 僅四年級(學號前3碼={student_id_prefix})")

        # 1. 取得該班導班級學生（僅顯示四年級：1131 時只列 110 屆／學號 110xxx）
        cursor.execute(f"""
            SELECT u.id AS student_id, u.username AS student_number, u.name AS student_name,
                   c.department AS class_department, c.name AS class_name_raw,
                   COALESCE(c.admission_year, u.admission_year) AS class_admission_year
            FROM users u
            LEFT JOIN classes c ON u.class_id = c.id
            WHERE u.class_id IN ({placeholders}) AND u.role = 'student'
            ORDER BY u.username
        """, tuple(class_ids))
        raw_students = cursor.fetchall()
        active_year = _get_active_semester_year(cursor)
        all_class_students = []
        for s in raw_students:
            # 有設定「僅四年級」時，只保留當屆實習生（110 屆 / 學號 110xxx）
            if student_id_prefix is not None:
                adm = s.get('class_admission_year')
                try:
                    adm_int = int(adm) if adm is not None else None
                except (TypeError, ValueError):
                    adm_int = None
                username = str(s.get('student_number') or '')
                prefix_ok = (adm_int == student_id_prefix) or (len(username) >= 3 and username[:3] == str(student_id_prefix))
                if not prefix_ok:
                    continue
            class_display = _format_class_display(
                s.get('class_department'),
                s.get('class_name_raw'),
                s.get('class_admission_year'),
                active_year
            )
            # 若格式化結果為 —，且 classes.name 有值，直接使用（部分資料庫可能已存完整班級名）
            if class_display == '—' and (s.get('class_name_raw') or '').strip():
                class_display = (s.get('class_name_raw') or '').strip()
            all_class_students.append({
                'student_id': s['student_id'],
                'student_number': s.get('student_number') or '',
                'student_name': s.get('student_name') or '',
                'class_name': class_display
            })

        # 2. 查詢已填寫志願序的學生（僅流程學期、僅四年級：1131 學年班導只看 110 屆／學號 110xxx）
        if flow_semester_id:
            if student_id_prefix is not None:
                prefix_str = str(student_id_prefix)
                cursor.execute(f"""
                    SELECT 
                        u.id AS student_id,
                        u.name AS student_name,
                        u.username AS student_number,
                        sp.id AS preference_id,
                        sp.preference_order,
                        sp.company_id,
                        COALESCE(ic.company_name, '未知公司') AS company_name,
                        sp.job_id,
                        sp.job_title,
                        sp.status,
                        sp.submitted_at,
                        sp.semester_id
                    FROM student_preferences sp
                    JOIN users u ON sp.student_id = u.id
                    LEFT JOIN classes c ON u.class_id = c.id
                    LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                    WHERE u.class_id IN ({placeholders}) 
                      AND u.role = 'student'
                      AND sp.semester_id = %s
                      AND (COALESCE(c.admission_year, u.admission_year) = %s OR u.username LIKE %s)
                    ORDER BY u.name, sp.preference_order
                """, tuple(class_ids) + (flow_semester_id, student_id_prefix, prefix_str + '%'))
            else:
                cursor.execute(f"""
                    SELECT 
                        u.id AS student_id,
                        u.name AS student_name,
                        u.username AS student_number,
                        sp.id AS preference_id,
                        sp.preference_order,
                        sp.company_id,
                        COALESCE(ic.company_name, '未知公司') AS company_name,
                        sp.job_id,
                        sp.job_title,
                        sp.status,
                        sp.submitted_at,
                        sp.semester_id
                    FROM student_preferences sp
                    JOIN users u ON sp.student_id = u.id
                    LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                    WHERE u.class_id IN ({placeholders}) 
                      AND u.role = 'student'
                      AND sp.semester_id = %s
                    ORDER BY u.name, sp.preference_order
                """, tuple(class_ids) + (flow_semester_id,))
        else:
            print("⚠️ 沒有設定學期，查詢該班級所有志願序")
            cursor.execute(f"""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number,
                    sp.id AS preference_id,
                    sp.preference_order,
                    sp.company_id,
                    COALESCE(ic.company_name, '未知公司') AS company_name,
                    sp.job_id,
                    sp.job_title,
                    sp.status,
                    sp.submitted_at,
                    sp.semester_id
                FROM student_preferences sp
                JOIN users u ON sp.student_id = u.id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id IN ({placeholders}) 
                  AND u.role = 'student'
                ORDER BY u.name, sp.preference_order
            """, tuple(class_ids))
        
        results = cursor.fetchall()
        print(f"📋 查詢結果數量: {len(results)} 筆")
        if results:
            print(f"📋 第一筆資料: {results[0]}")
        
        # 格式化時間
        from datetime import timezone, timedelta
        taiwan_tz = timezone(timedelta(hours=8))
        
        for row in results:
            if row.get('submitted_at') and isinstance(row['submitted_at'], datetime):
                row['submitted_at'] = row['submitted_at'].astimezone(taiwan_tz).strftime("%Y-%m-%d %H:%M:%S")
            elif row.get('submitted_at'):
                row['submitted_at'] = str(row['submitted_at'])
            else:
                row['submitted_at'] = ""

        # 整理資料結構給前端使用
        student_data = defaultdict(lambda: {'student_id': None, 'student_number': '', 'preferences': []})
        processed_count = 0
        skipped_count = 0
        
        print(f"📋 開始處理 {len(results)} 筆查詢結果")
        
        for row in results:
            student_name = row.get('student_name')
            student_id = row.get('student_id')
            preference_id = row.get('preference_id')
            preference_order = row.get('preference_order')
            company_name = row.get('company_name') or '未知公司'  # 如果為 NULL，使用預設值
            status = row.get('status') or 'submitted'  # student_preferences.status 的預設值是 'submitted'
            submitted_at = row.get('submitted_at', '')
            
            # 調試：顯示每筆資料
            print(f"📝 處理資料: student={student_name}, student_id={student_id}, order={preference_order}, company={company_name}, company_id={row.get('company_id')}, status={status}")
            
            # 只添加有完整資料的志願序（至少要有學生名稱和順序）
            if student_name and preference_order:
                if student_data[student_name]['student_id'] is None:
                    student_data[student_name]['student_id'] = student_id
                    student_data[student_name]['student_number'] = row.get('student_number') or ''
                student_data[student_name]['preferences'].append({
                    'preference_id': preference_id,
                    'order': preference_order,
                    'company': company_name,
                    'job_title': row.get('job_title') or '',
                    'status': status,
                    'submitted_at': submitted_at or ''
                })
                processed_count += 1
            else:
                skipped_count += 1
                print(f"⚠️ 跳過資料: student_name={student_name}, order={preference_order}, company={company_name}")

        print(f"✅ 資料處理完成: 處理 {processed_count} 筆，跳過 {skipped_count} 筆")

        # 如果沒有資料，添加詳細調試信息
        if len(student_data) == 0:
            print(f"⚠️ 警告: 沒有找到任何學生的志願序資料")
            print(f"   查詢條件: class_ids={class_ids}, current_semester_id={current_semester_id}")
            
            # 檢查是否有該班級的學生
            cursor.execute(f"SELECT COUNT(*) as count FROM users WHERE class_id IN ({placeholders}) AND role = 'student'", tuple(class_ids))
            student_count = cursor.fetchone()
            print(f"   該班級共有學生: {student_count.get('count', 0) if student_count else 0} 人")
            
            # 檢查是否有志願序（不限班級）
            cursor.execute("SELECT COUNT(*) as count FROM student_preferences", ())
            all_prefs_count = cursor.fetchone()
            print(f"   系統中總共有志願序: {all_prefs_count.get('count', 0) if all_prefs_count else 0} 筆")
            
            # 檢查該班級學生的志願序（不限制學期）
            cursor.execute(f"""
                SELECT COUNT(*) as count
                FROM student_preferences sp
                JOIN users u ON sp.student_id = u.id
                WHERE u.class_id IN ({placeholders}) AND u.role = 'student'
            """, tuple(class_ids))
            class_prefs_count = cursor.fetchone()
            print(f"   該班級學生填寫的志願序（不限學期）: {class_prefs_count.get('count', 0) if class_prefs_count else 0} 筆")
            
            # 檢查當前學期的所有志願序（不限班級）
            if current_semester_id:
                cursor.execute("""
                    SELECT COUNT(*) as count
                    FROM student_preferences
                    WHERE semester_id = %s
                """, (current_semester_id,))
                semester_prefs_count = cursor.fetchone()
                print(f"   當前學期的所有志願序（不限班級）: {semester_prefs_count.get('count', 0) if semester_prefs_count else 0} 筆")
                
                # 檢查該班級學生在當前學期的志願序（詳細）
                cursor.execute(f"""
                    SELECT 
                        u.id AS student_id,
                        u.name AS student_name,
                        u.username AS student_number,
                        sp.preference_order,
                        sp.semester_id,
                        sp.company_id,
                        sp.status,
                        sp.submitted_at
                    FROM student_preferences sp
                    JOIN users u ON sp.student_id = u.id
                    WHERE u.class_id IN ({placeholders}) 
                      AND u.role = 'student'
                      AND sp.semester_id = %s
                    ORDER BY u.name, sp.preference_order
                    LIMIT 5
                """, tuple(class_ids) + (current_semester_id,))
                sample_data = cursor.fetchall()
                if sample_data:
                    print(f"   ✅ 找到了 {len(sample_data)} 筆樣本資料:")
                    for sample in sample_data:
                        print(f"      - 學生: {sample.get('student_name')} ({sample.get('student_number')}), 志願序: {sample.get('preference_order')}, 學期ID: {sample.get('semester_id')}")
                else:
                    print(f"   ❌ 查詢結果為空（即使使用相同的條件）")
                
                # 檢查該班級學生在當前學期的志願序（但查詢所有學期）
                cursor.execute(f"""
                    SELECT 
                        u.id AS student_id,
                        u.name AS student_name,
                        u.username AS student_number,
                        sp.preference_order,
                        sp.semester_id,
                        sp.company_id,
                        sp.status
                    FROM student_preferences sp
                    JOIN users u ON sp.student_id = u.id
                    WHERE u.class_id IN ({placeholders}) 
                      AND u.role = 'student'
                    ORDER BY u.name, sp.preference_order
                    LIMIT 5
                """, tuple(class_ids))
                all_semester_data = cursor.fetchall()
                if all_semester_data:
                    print(f"   📋 該班級學生在所有學期的志願序（樣本）:")
                    for sample in all_semester_data:
                        print(f"      - 學生: {sample.get('student_name')}, 志願序: {sample.get('preference_order')}, 學期ID: {sample.get('semester_id')}")
            
            # 檢查所有志願序的學期ID分佈
            cursor.execute("""
                SELECT semester_id, COUNT(*) as count
                FROM student_preferences
                GROUP BY semester_id
            """)
            semester_dist = cursor.fetchall()
            if semester_dist:
                print(f"   志願序的學期ID分佈:")
                for dist in semester_dist:
                    print(f"      - 學期ID {dist.get('semester_id')}: {dist.get('count')} 筆")
        
        # 如果沒有資料，並且有設定當前學期，嘗試查詢所有學期的資料作為備用
        if len(student_data) == 0 and current_semester_id:
            print(f"🔄 嘗試查詢該班級在所有學期的志願序（作為診斷）...")
            cursor.execute(f"""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number,
                    sp.id AS preference_id,
                    sp.preference_order,
                    sp.company_id,
                    COALESCE(ic.company_name, '未知公司') AS company_name,
                    sp.job_id,
                    sp.job_title,
                    sp.status,
                    sp.submitted_at,
                    sp.semester_id
                FROM student_preferences sp
                JOIN users u ON sp.student_id = u.id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id IN ({placeholders})
                  AND u.role = 'student'
                ORDER BY u.name, sp.preference_order
            """, tuple(class_ids))
            all_semester_results = cursor.fetchall()
            print(f"📋 該班級在所有學期的志願序數量: {len(all_semester_results)} 筆")
            
            if all_semester_results:
                print(f"💡 建議: 將查詢條件改為不限制學期，以顯示所有志願序")
                # 重新處理所有學期的資料
                for row in all_semester_results:
                    student_name = row.get('student_name')
                    student_id = row.get('student_id')
                    preference_id = row.get('preference_id')
                    preference_order = row.get('preference_order')
                    company_name = row.get('company_name') or '未知公司'
                    status = row.get('status') or 'submitted'  # student_preferences.status 的預設值是 'submitted'
                    submitted_at = row.get('submitted_at', '')
                    
                    if student_name and preference_order:
                        if student_data[student_name]['student_id'] is None:
                            student_data[student_name]['student_id'] = student_id
                            student_data[student_name]['student_number'] = row.get('student_number') or ''
                        student_data[student_name]['preferences'].append({
                            'preference_id': preference_id,
                            'order': preference_order,
                            'company': company_name,
                            'job_title': row.get('job_title') or '',
                            'status': status,
                            'submitted_at': submitted_at or ''
                        })
                print(f"✅ 已載入該班級在所有學期的志願序: {len(student_data)} 位學生")

        # 3. 區分已填寫與未填寫（在 fallback 之後執行）
        student_class_map = {s['student_id']: (s.get('class_name') or '—') for s in all_class_students if s.get('student_id')}
        filled_student_ids = {sinfo['student_id'] for sinfo in student_data.values() if sinfo.get('student_id')}
        unfilled_students = []
        for s in all_class_students:
            sid = s.get('student_id')
            sname = s.get('student_name') or ''
            sno = s.get('student_number') or ''
            if sid and sid not in filled_student_ids:
                unfilled_students.append({
                    'student_id': sid, 'student_number': sno, 'student_name': sname,
                    'class_name': s.get('class_name') or '—'
                })
        
        filled_student_list = []
        for sname, sinfo in student_data.items():
            sid = sinfo['student_id']
            filled_student_list.append({
                'student_id': sid,
                'student_number': sinfo.get('student_number', ''),
                'student_name': sname,
                'class_name': student_class_map.get(sid, '—'),
                'preferences': sinfo['preferences'],
                'status': sinfo['preferences'][0]['status'] if sinfo['preferences'] else 'submitted'
            })

        # 查詢所有實習公司與職缺（供志願結構分析表顯示完整選項）
        cursor.execute("""
            SELECT ic.id AS company_id, ic.company_name,
                   ij.id AS job_id, COALESCE(ij.title, '') AS job_title
            FROM internship_companies ic
            INNER JOIN internship_jobs ij ON ij.company_id = ic.id
            WHERE ic.is_active = 1 AND ic.status = 'approved' AND ij.is_active = 1
            ORDER BY ic.company_name, ij.title
        """)
        all_internship_jobs = cursor.fetchall()

        return render_template('preferences/review_preferences.html',
            student_data=student_data,
            all_class_students=all_class_students,
            filled_student_list=filled_student_list,
            unfilled_students=unfilled_students,
            all_internship_jobs=all_internship_jobs)

    except Exception as e:
        print("取得志願資料錯誤：", e)
        traceback.print_exc()
        return "伺服器錯誤", 500
    finally:
        try:
            cursor.close()
            conn.close()
        except Exception:
            pass

# -------------------------
# API - 審核學生志願序（通過/退件）
# -------------------------
@preferences_bp.route('/api/review_preferences', methods=['POST'])
def review_preferences_api():
    user_id = session.get('user_id')
    user_role = session.get('role')

    # 權限檢查：只允許班導、老師、主任審核
    ALLOWED_ROLES = ['teacher', 'admin', 'class_teacher', 'director']
    if not user_id or user_role not in ALLOWED_ROLES:
        return jsonify({"success": False, "message": "未授權或無權限"}), 403

    data = request.get_json()
    student_id = data.get('student_id')
    status = data.get('status')
    reason = data.get('reason', '')

    if not student_id:
        return jsonify({"success": False, "message": "缺少學生ID"}), 400

    if status not in ['approved', 'rejected']:
        return jsonify({"success": False, "message": "無效的狀態碼"}), 400

    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 檢查是否為班導（只能審核自己班級的學生）
        if user_role == 'class_teacher' or user_role == 'teacher':
            cursor.execute("""
                SELECT ct.teacher_id, u.class_id
                FROM classes_teacher ct
                JOIN users u ON ct.class_id = u.class_id
                WHERE ct.teacher_id = %s AND u.id = %s AND u.role = 'student'
            """, (user_id, student_id))
            relation = cursor.fetchone()
            if not relation:
                return jsonify({"success": False, "message": "您只能審核自己班級學生的志願序"}), 403

        # 取得當前學期ID
        current_semester_id = get_current_semester_id(cursor)

        if not current_semester_id:
            return jsonify({"success": False, "message": "尚未設定當前學期"}), 400

        # 取得學生資訊
        cursor.execute("SELECT name, email FROM users WHERE id = %s", (student_id,))
        student_info = cursor.fetchone()
        if not student_info:
            return jsonify({"success": False, "message": "找不到學生資料"}), 404

        student_name = student_info['name']
        student_email = student_info.get('email', '')

        # 取得審核者姓名
        cursor.execute("SELECT name FROM users WHERE id = %s", (user_id,))
        reviewer = cursor.fetchone()
        reviewer_name = reviewer['name'] if reviewer else "審核老師"

        # 更新該學生在當前學期的所有志願序狀態
        if status == 'approved':
            # 通過：更新所有志願序狀態
            cursor.execute("""
                UPDATE student_preferences 
                SET status = %s
                WHERE student_id = %s 
                  AND semester_id = %s
            """, (status, student_id, current_semester_id))
            
            # 發送通過通知
            notification_content = (
                f"恭喜您！您的實習志願序已由 {reviewer_name} 老師審核通過。\n"
                f"您可以繼續後續的實習申請流程。"
            )
            create_notification(
                user_id=student_id,
                title="志願序審核通過通知",
                message=notification_content,
                category="ranking"
            )

        elif status == 'rejected':
            # 退件：更新所有志願序狀態並記錄退件原因（可選）
            cursor.execute("""
                UPDATE student_preferences 
                SET status = %s
                WHERE student_id = %s 
                  AND semester_id = %s
            """, (status, student_id, current_semester_id))

            # 發送退件通知
            notification_content = (
                f"您的實習志願序已被 {reviewer_name} 老師退件。\n\n"
                f"請修改後重新提交。"
            )
            create_notification(
                user_id=student_id,
                title="志願序退件通知",
                message=notification_content,
                category="ranking",
                link_url="/fill_preferences"  # 連結到志願填寫頁面，方便學生修改
            )

        conn.commit()

        return jsonify({"success": True, "message": "志願序審核狀態更新成功"})

    except Exception as e:
        conn.rollback()
        traceback.print_exc()
        return jsonify({"success": False, "message": f"伺服器錯誤: {str(e)}"}), 500

    finally:
        cursor.close()
        conn.close()

# -------------------------
# Excel 導出功能
# -------------------------
@preferences_bp.route('/export_preferences_excel')
def export_preferences_excel():
    if 'username' not in session or session.get('role') not in ['teacher', 'director', 'class_teacher']:
        return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導（取完結果再重用 cursor，避免 Unread result found）
        cursor.execute("""
        SELECT c.id AS class_id, c.name AS class_name
        FROM classes c
        JOIN classes_teacher ct ON c.id = ct.class_id
        WHERE ct.teacher_id = %s AND ct.role = 'classteacher'
        """, (user_id,))
        class_rows = cursor.fetchall()
        class_info = class_rows[0] if class_rows else None
        if not class_info:
            return "你不是班導，無法導出志願序", 403

        class_id = class_info['class_id']
        class_name = class_info['class_name']

        # 取得當前學期ID
        current_semester_id = get_current_semester_id(cursor)

        # 查詢班上學生及其志願（只匯出已通過的志願序）
        if current_semester_id:
            cursor.execute("""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number, 
                    sp.preference_order,
                    ic.company_name,
                    sp.submitted_at
                FROM users u
                INNER JOIN student_preferences sp ON u.id = sp.student_id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id = %s 
                  AND u.role = 'student'
                  AND sp.status = 'approved'
                  AND sp.semester_id = %s
                ORDER BY u.name, sp.preference_order
            """, (class_id, current_semester_id))
        else:
            # 如果沒有設定當前學期，只匯出已通過的志願序
            cursor.execute("""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number, 
                    sp.preference_order,
                    ic.company_name,
                    sp.submitted_at
                FROM users u
                INNER JOIN student_preferences sp ON u.id = sp.student_id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id = %s 
                  AND u.role = 'student'
                  AND sp.status = 'approved'
                ORDER BY u.name, sp.preference_order
            """, (class_id,))
        results = cursor.fetchall()

        # 創建 Excel 工作簿
        wb = Workbook()
        ws = wb.active
        ws.title = f"{class_name}志願序"

        # 設定樣式
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="0066CC", end_color="0066CC", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # 寫入標題
        ws.merge_cells('A1:G1')
        title_cell = ws['A1']
        title_cell.value = f"{class_name} - 已通過學生實習志願序統計表"
        title_cell.font = Font(bold=True, size=16)
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 30

        # 寫入導出時間
        ws.merge_cells('A2:G2')
        time_cell = ws['A2']
        time_cell.value = f"導出時間：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
        time_cell.alignment = Alignment(horizontal="right")

        # 寫入欄位名稱
        headers = ['學生姓名', '學號', '第一志願', '第二志願', '第三志願', '第四志願', '第五志願']
        ws.row_dimensions[4].height = 25
        for col_num, header in enumerate(headers, 1):
            col_letter = get_column_letter(col_num)
            cell = ws[f'{col_letter}4']
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
            if col_num in [3, 4, 5, 6, 7]:
                ws.column_dimensions[col_letter].width = 25

        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 15

        # 整理學生資料
        student_data = defaultdict(lambda: {
            'name': '',
            'student_number': '',
            'preferences': [''] * 5,
            'submitted_times': [''] * 5
        })

        for row in results:
            student_name = row['student_name']
            if student_name:
                student_data[student_name]['name'] = student_name
                student_data[student_name]['student_number'] = row['student_number'] or ''
                
                if row['preference_order'] and row['company_name']:
                    order = row['preference_order'] - 1 # 轉為 0-based index
                    if 0 <= order < 5:
                        student_data[student_name]['preferences'][order] = row['company_name']
                        if row['submitted_at']:
                            student_data[student_name]['submitted_times'][order] = row['submitted_at'].strftime('%m/%d %H:%M')

        # 寫入學生資料
        row_num = 5
        for student_name in sorted(student_data.keys()):
            data = student_data[student_name]
            
            # 學生姓名
            name_cell = ws.cell(row=row_num, column=1, value=data['name'])
            name_cell.border = border
            name_cell.alignment = Alignment(horizontal="center", vertical="center")
            # 學號
            number_cell = ws.cell(row=row_num, column=2, value=data['student_number'])
            number_cell.border = border
            number_cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # 志願序
            for i in range(5):
                pref_text = data['preferences'][i] or ''
                
                cell = ws.cell(row=row_num, column=3+i, value=pref_text)
                cell.border = border
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                
            ws.row_dimensions[row_num].height = 30
            row_num += 1

        # 添加統計資訊
        ws.cell(row=row_num + 1, column=1, value="統計資訊：").font = Font(bold=True)
        
        # 統計各公司被選擇次數
        company_counts = defaultdict(int)
        for data in student_data.values():
            for pref in data['preferences']:
                if pref:
                    company_counts[pref] += 1

        stats_row = row_num + 2
        ws.cell(row=stats_row, column=1, value="公司名稱").font = Font(bold=True)
        ws.cell(row=stats_row, column=2, value="被選擇次數").font = Font(bold=True)
        stats_row += 1
        
        for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
            ws.cell(row=stats_row, column=1, value=company).border = border
            ws.cell(row=stats_row, column=2, value=count).border = border
            stats_row += 1


        # 建立 response
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        filename = f"{class_name}_已通過實習志願序_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        traceback.print_exc()
        return "導出 Excel 失敗", 500
    finally:
        try:
            cursor.close()
            conn.close()
        except Exception:
            pass

# -------------------------
# PDF 導出功能
# -------------------------
@preferences_bp.route('/export_preferences_pdf')
def export_preferences_pdf():
    if 'username' not in session or session.get('role') not in ['teacher', 'director', 'class_teacher']:
        return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導（取完結果再重用 cursor，避免 Unread result found）
        cursor.execute("""
        SELECT c.id AS class_id, c.name AS class_name
        FROM classes c
        JOIN classes_teacher ct ON c.id = ct.class_id
        WHERE ct.teacher_id = %s AND ct.role = 'classteacher'
        """, (user_id,))
        class_rows = cursor.fetchall()
        class_info = class_rows[0] if class_rows else None
        if not class_info:
            return "你不是班導，無法導出志願序", 403

        class_id = class_info['class_id']
        class_name = class_info['class_name']

        # 取得當前學期ID
        current_semester_id = get_current_semester_id(cursor)

        # 查詢班上學生及其志願（只匯出已通過的志願序）
        if current_semester_id:
            cursor.execute("""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number, 
                    sp.preference_order,
                    ic.company_name,
                    sp.submitted_at
                FROM users u
                INNER JOIN student_preferences sp ON u.id = sp.student_id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id = %s 
                  AND u.role = 'student'
                  AND sp.status = 'approved'
                  AND sp.semester_id = %s
                ORDER BY u.name, sp.preference_order
            """, (class_id, current_semester_id))
        else:
            # 如果沒有設定當前學期，只匯出已通過的志願序
            cursor.execute("""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number, 
                    sp.preference_order,
                    ic.company_name,
                    sp.submitted_at
                FROM users u
                INNER JOIN student_preferences sp ON u.id = sp.student_id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id = %s 
                  AND u.role = 'student'
                  AND sp.status = 'approved'
                ORDER BY u.name, sp.preference_order
            """, (class_id,))
        results = cursor.fetchall()

        # 創建 PDF 緩衝區
        pdf_buffer = io.BytesIO()
        
        # 註冊中文字體（嘗試使用系統字體）
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import platform
            
            # 根據作業系統選擇字體路徑
            if platform.system() == 'Windows':
                # Windows 系統字體路徑
                font_paths = [
                    'C:/Windows/Fonts/msjh.ttc',  # 微軟正黑體
                    'C:/Windows/Fonts/simsun.ttc',  # 新細明體
                    'C:/Windows/Fonts/kaiu.ttf',  # 標楷體
                ]
                font_name = None
                for font_path in font_paths:
                    try:
                        if 'msjh' in font_path.lower():
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                            font_name = 'ChineseFont'
                            break
                        elif 'simsun' in font_path.lower():
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                            font_name = 'ChineseFont'
                            break
                        elif 'kaiu' in font_path.lower():
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            font_name = 'ChineseFont'
                            break
                    except:
                        continue
                
                if not font_name:
                    # 如果找不到字體，使用 reportlab 的內建字體處理
                    font_name = 'Helvetica'
            else:
                # Linux/Mac 系統，嘗試使用常見字體
                font_name = 'Helvetica'
        except Exception as e:
            print(f"字體註冊失敗: {e}")
            font_name = 'Helvetica'
        
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, topMargin=1*inch, bottomMargin=1*inch)
        
        # 設定樣式
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=20,
            alignment=1, # 置中
            textColor=colors.HexColor('#0066CC'),
            fontName=font_name if font_name else 'Helvetica'
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=8,
            fontName=font_name if font_name else 'Helvetica'
        )

        # 建立內容
        story = []

        # 標題
        title = Paragraph(f"{class_name} - 已通過學生實習志願序統計表", title_style)
        story.append(title)
        
        # 日期
        date_text = f"導出時間：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
        date_para = Paragraph(date_text, normal_style)
        story.append(date_para)
        story.append(Spacer(1, 20))


        # 整理學生資料
        student_data = defaultdict(lambda: {
            'name': '',
            'student_number': '',
            'preferences': [''] * 5,
            'submitted_times': [''] * 5
        })

        for row in results:
            student_name = row['student_name']
            if student_name:
                student_data[student_name]['name'] = student_name
                student_data[student_name]['student_number'] = row['student_number'] or ''
                
                if row['preference_order'] and row['company_name']:
                    order = row['preference_order'] - 1 # 轉為 0-based index
                    if 0 <= order < 5:
                        company_name = row['company_name']
                        submitted_at = row['submitted_at'].strftime('%m/%d %H:%M') if row['submitted_at'] else ''
                        student_data[student_name]['preferences'][order] = company_name
                        student_data[student_name]['submitted_times'][order] = submitted_at


        # 學生表格
        table_data = [
            ['學生姓名', '學號', '第一志願', '第二志願', '第三志願', '第四志願', '第五志願']
        ]
        
        for student_name in sorted(student_data.keys()):
            data = student_data[student_name]
            row = [data['name'], data['student_number']]
            
            for i in range(5):
                pref_text = data['preferences'][i] or ''
                row.append(pref_text)
            
            table_data.append(row)

        table = Table(table_data, colWidths=[1*inch, 0.8*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
        
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), font_name if font_name else 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')])
        ])
        table.setStyle(table_style)
        story.append(table)
        story.append(Spacer(1, 20))

        # 統計資訊
        story.append(Paragraph("<b>統計資訊：</b>", normal_style))
        story.append(Spacer(1, 5))

        company_counts = defaultdict(int)
        for data in student_data.values():
            for pref in data['preferences']:
                if pref:
                    company_counts[pref] += 1

        stats_data = [
            ['公司名稱', '被選擇次數']
        ]
        
        for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
            stats_data.append([company, count])

        stats_table = Table(stats_data, colWidths=[3*inch, 1*inch])
        stats_table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), font_name if font_name else 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ])
        stats_table.setStyle(stats_table_style)
        story.append(stats_table)

        # 建立 PDF
        doc.build(story)
        pdf_buffer.seek(0)

        filename = f"{class_name}_已通過實習志願序_{datetime.now().strftime('%Y%m%d')}.pdf"

        return send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        traceback.print_exc()
        return "導出 PDF 失敗", 500
    finally:
        try:
            cursor.close()
            conn.close()
        except Exception:
            pass


# -------------------------
# Word 導出功能
# -------------------------
@preferences_bp.route('/export_preferences_word')
@preferences_bp.route('/export_preferences_docx')
def export_preferences_docx():
    if 'username' not in session or session.get('role') not in ['teacher', 'director', 'class_teacher']:
        return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 確認是否為班導（取完結果再重用 cursor，避免 Unread result found）
        cursor.execute("""
        SELECT c.id AS class_id, c.name AS class_name
        FROM classes c
        JOIN classes_teacher ct ON c.id = ct.class_id
        WHERE ct.teacher_id = %s AND ct.role = 'classteacher'
        """, (user_id,))
        class_rows = cursor.fetchall()
        class_info = class_rows[0] if class_rows else None
        if not class_info:
            return "你不是班導，無法導出志願序", 403

        class_id = class_info['class_id']
        class_name = class_info['class_name']

        # 取得當前學期ID
        current_semester_id = get_current_semester_id(cursor)

        # 查詢班上學生及其志願（只匯出已通過的志願序）
        if current_semester_id:
            cursor.execute("""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number, 
                    sp.preference_order,
                    ic.company_name,
                    sp.submitted_at
                FROM users u
                INNER JOIN student_preferences sp ON u.id = sp.student_id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id = %s 
                  AND u.role = 'student'
                  AND sp.status = 'approved'
                  AND sp.semester_id = %s
                ORDER BY u.name, sp.preference_order
            """, (class_id, current_semester_id))
        else:
            # 如果沒有設定當前學期，只匯出已通過的志願序
            cursor.execute("""
                SELECT 
                    u.id AS student_id,
                    u.name AS student_name,
                    u.username AS student_number, 
                    sp.preference_order,
                    ic.company_name,
                    sp.submitted_at
                FROM users u
                INNER JOIN student_preferences sp ON u.id = sp.student_id
                LEFT JOIN internship_companies ic ON sp.company_id = ic.id
                WHERE u.class_id = %s 
                  AND u.role = 'student'
                  AND sp.status = 'approved'
                ORDER BY u.name, sp.preference_order
            """, (class_id,))
        results = cursor.fetchall()

        # 整理學生資料
        student_data = defaultdict(lambda: {
            'name': '',
            'student_number': '',
            'preferences': [''] * 5,
            'submitted_times': [''] * 5
        })

        for row in results:
            student_name = row['student_name']
            if student_name:
                student_data[student_name]['name'] = student_name
                student_data[student_name]['student_number'] = row['student_number'] or ''
                
                if row['preference_order'] and row['company_name']:
                    order = row['preference_order'] - 1
                    if 0 <= order < 5:
                        student_data[student_name]['preferences'][order] = row['company_name']
                        if row['submitted_at']:
                            student_data[student_name]['submitted_times'][order] = row['submitted_at'].strftime('%m/%d %H:%M')

        # 建立 Word 文件
        doc = Document()
        title = doc.add_heading(f"{class_name} - 已通過學生實習志願序統計表", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"導出時間：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph("")

        # 學生表格
        table = doc.add_table(rows=1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        headers = ['學生姓名', '學號', '第一志願', '第二志願', '第三志願', '第四志願', '第五志願']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            # 設置標題欄位居中
            table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for student_name in sorted(student_data.keys()):
            data = student_data[student_name]
            row = table.add_row().cells
            row[0].text = data['name']
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[1].text = data['student_number']
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i in range(5):
                pref_text = data['preferences'][i] or ''
                row[2+i].text = pref_text
                # 設置內容置中對齊
                row[2+i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        doc.add_heading("統計資訊", level=1)

        # 統計資訊
        company_counts = defaultdict(int)
        for data in student_data.values():
            for pref in data['preferences']:
                if pref:
                    company_counts[pref] += 1
        
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = "Table Grid"
        stats_table.rows[0].cells[0].text = "公司名稱"
        stats_table.rows[0].cells[1].text = "被選擇次數"

        for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
            row = stats_table.add_row().cells
            row[0].text = company
            row[1].text = str(count)


        # 建立 response
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        filename = f"{class_name}_已通過實習志願序_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        traceback.print_exc()
        return "導出 Word 失敗", 500
    finally:
        try:
            cursor.close()
            conn.close()
        except Exception:
            pass

# -------------------------
# 公司端查看選擇學生的結果
# -------------------------
@preferences_bp.route('/review_company_choices')
def review_company_choices():
    if 'user_id' not in session or session.get('role') not in ['company_contact', 'director']:
        return redirect(url_for('auth_bp.login_page'))

    user_id = session.get('user_id')
    user_role = session.get('role')
    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        companies = []
        if user_role == 'director':
            # 主管身份可以看到所有公司（這裡簡化，假設 director 可見）
            cursor.execute("""
                SELECT id AS company_id, company_name FROM internship_companies 
                WHERE status = 'approved'
            """)
            companies = cursor.fetchall()
        elif user_role == 'company_contact':
            # 廠商聯絡人只能看到自己負責的公司
            cursor.execute("""
                SELECT ic.id AS company_id, ic.company_name
                FROM internship_companies ic
                JOIN company_contacts cc ON ic.id = cc.company_id
                WHERE cc.user_id = %s AND ic.status = 'approved'
            """, (user_id,))
            companies = cursor.fetchall()
        
        if not companies:
            return render_template(
                'preferences/admission_results.html',
                companies=[],
                student_data={},
                message="目前尚未綁定任何實習公司。"
            )

        # 找出選擇這些公司的學生
        company_ids = tuple([c['company_id'] for c in companies])
        # 使用 IN 查詢多個公司
        query = f"""
            SELECT 
                u.name AS student_name,
                u.username AS student_number,
                sp.preference_order,
                sp.submitted_at,
                ic.company_name,
                ij.title AS job_title
            FROM student_preferences sp
            JOIN users u ON sp.student_id = u.id
            JOIN internship_companies ic ON sp.company_id = ic.id
            JOIN internship_jobs ij ON sp.job_id = ij.id
            WHERE sp.company_id IN ({','.join(['%s'] * len(company_ids))}) 
            ORDER BY ic.company_name, sp.preference_order, u.name
        """
        cursor.execute(query, company_ids)
        rows = cursor.fetchall()

        # 整理成 {公司名稱: [學生資料...]} 結構
        student_data = defaultdict(list)
        for row in rows:
            student_data[row['company_name']].append({
                'student_name': row['student_name'],
                'student_number': row['student_number'],
                'preference_order': row['preference_order'],
                'job_title': row['job_title'],
                'submitted_at': row['submitted_at'].strftime('%Y-%m-%d %H:%M') if row['submitted_at'] else 'N/A'
            })


        return render_template(
            'preferences/admission_results.html',
            companies=companies,
            student_data=student_data,
            message=None
        )

    except Exception as e:
        traceback.print_exc()
        return "伺服器錯誤", 500
    finally:
        try:
            cursor.close()
            conn.close()
        except Exception:
            pass