from flask import Blueprint, request, jsonify, session, send_file, render_template, redirect, url_for
from werkzeug.utils import secure_filename
from config import get_db
from semester import get_current_semester_id 
from email_service import send_resume_rejection_email, send_resume_approval_email 
import os
import traceback
import json
from datetime import datetime

# 引入 docx 相關模組
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

resume_bp = Blueprint("resume_bp", __name__)

# 上傳資料夾設定
UPLOAD_FOLDER = "uploads/resumes"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# -------------------------
# Helper / 權限管理
# -------------------------
def get_user_by_username(cursor, username):
    cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
    return cursor.fetchone()

def get_user_by_id(cursor, user_id):
    cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
    return cursor.fetchone()

def get_director_department(cursor, user_id):
    """
    取得主任所屬 department（透過 classes_teacher -> classes.department）
    若管理多個班級，只回傳第一個有 department 的值（可擴充回傳 list）
    """
    cursor.execute("""
        SELECT DISTINCT c.department
        FROM classes c
        JOIN classes_teacher ct ON ct.class_id = c.id
        WHERE ct.teacher_id = %s
        LIMIT 1
    """, (user_id,))
    r = cursor.fetchone()
    return r['department'] if r and r.get('department') else None

def teacher_manages_class(cursor, teacher_id, class_id):
    cursor.execute("""
        SELECT 1 FROM classes_teacher
        WHERE teacher_id = %s AND class_id = %s
        LIMIT 1
    """, (teacher_id, class_id))
    return cursor.fetchone() is not None

def can_access_target_resume(cursor, session_user_id, session_role, target_user_id):
    """
    判斷 session 的使用者（session_user_id, session_role）是否可存取 target_user_id 的履歷
    - admin: 全部
    - ta: 只讀（此函式只處理存取權，呼叫端需再判斷是否為可寫操作）
    - student: 只能存取自己的履歷
    - teacher: 只能存取自己帶的班級學生
    - director: 只能存取自己科系的學生（由 classes_teacher -> classes.department 判斷）
    """
    # admin 可以
    if session_role == "admin":
        return True

    # student 只能自己
    if session_role == "student":
        return session_user_id == target_user_id

    # ta 可以讀所有（呼叫端若為寫動作需拒絕）
    if session_role == "ta":
        return True

    # teacher / director 需要查 student 的班級與科系
    cursor.execute("SELECT class_id FROM users WHERE id = %s", (target_user_id,))
    u = cursor.fetchone()
    if not u:
        return False
    target_class_id = u.get('class_id')

    if session_role == "class_teacher":
        return teacher_manages_class(cursor, session_user_id, target_class_id)

    if session_role == "director":
        # 取得 director 的 department（若沒有設定，則無法存取）
        director_dept = get_director_department(cursor, session_user_id)
        if not director_dept:
            return False
        # 取得 target student's department
        cursor.execute("SELECT c.department FROM classes c WHERE c.id = %s", (target_class_id,))
        cd = cursor.fetchone()
        if not cd:
            return False
        return cd.get('department') == director_dept

    # 預設拒絕
    return False

def require_login():
    return 'user_id' in session and 'role' in session

# -------------------------
# 新增：資料儲存與文件生成函式
# -------------------------

def save_structured_data(cursor, student_id, data):
    """將結構化資料寫入四個正規化表格。"""
    try:
        # 1. 寫入 Student_Info (學生基本資料)
        cursor.execute("""
            INSERT INTO Student_Info (StuID, StuName, BirthDate, Gender, Phone, Email, Address, ConductScore, Autobiography, PhotoPath, AbsencesPath)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE 
                StuName=VALUES(StuName), BirthDate=VALUES(BirthDate), Gender=VALUES(Gender), Phone=VALUES(Phone), 
                Email=VALUES(Email), Address=VALUES(Address), ConductScore=VALUES(ConductScore), 
                Autobiography=VALUES(Autobiography), PhotoPath=VALUES(PhotoPath), AbsencesPath=VALUES(AbsencesPath), UpdatedAt=NOW()
        """, (
            student_id, data.get('name'), data.get('birth_date'), data.get('gender'), data.get('phone'),
            data.get('email'), data.get('address'), data.get('conduct_score'), data.get('autobiography'),
            data.get('photo_path'), data.get('absences_path')
        ))

        # 2. 寫入 Course_Grades (修課成績)
        cursor.execute("DELETE FROM Course_Grades WHERE StuID = %s", (student_id,))
        courses = data.get('courses', [])
        for course in courses:
            if course.get('name'):
                cursor.execute("""
                    INSERT INTO Course_Grades (StuID, CourseName, Credits, Grade)
                    VALUES (%s, %s, %s, %s)
                """, (student_id, course['name'], course.get('credits'), course.get('grade')))
        
        # 3. 寫入 Certificate_Skills (證照與語文能力)
        cursor.execute("DELETE FROM Certificate_Skills WHERE StuID = %s", (student_id,))
        certs_languages = data.get('certs_languages', [])
        for item in certs_languages:
            if item.get('name'):
                cursor.execute("""
                    INSERT INTO Certificate_Skills (StuID, Type, Name, Proficiency, CertImagePath)
                    VALUES (%s, %s, %s, %s, %s)
                """, (student_id, item['type'], item['name'], item.get('proficiency'), item.get('image_path')))

        # 4. 寫入 Internship_Preferences (實習志願序)
        cursor.execute("DELETE FROM Internship_Preferences WHERE StuID = %s", (student_id,))
        preferences = data.get('preferences', [])
        for i, pref in enumerate(preferences):
            if pref.get('company_name'):
                cursor.execute("""
                    INSERT INTO Internship_Preferences (StuID, PreferenceRank, CompanyName, JobTitle)
                    VALUES (%s, %s, %s, %s)
                """, (student_id, i + 1, pref['company_name'], pref.get('job_title')))
        
        return True
    
    except Exception as e:
        print(f"寫入結構化資料錯誤: {e}")
        traceback.print_exc()
        return False

def get_student_info_for_doc(cursor, student_id):
    """從資料庫獲取所有生成文件所需資料。"""
    data = {}
    
    # 1. 基本資料
    cursor.execute("SELECT * FROM Student_Info WHERE StuID = %s", (student_id,))
    data['info'] = cursor.fetchone()
    
    # 2. 成績
    cursor.execute("SELECT CourseName, Credits, Grade FROM Course_Grades WHERE StuID = %s", (student_id,))
    data['grades'] = cursor.fetchall()
    
    # 3. 證照與語文
    cursor.execute("SELECT Type, Name, Proficiency, CertImagePath FROM Certificate_Skills WHERE StuID = %s", (student_id,))
    data['certs_languages'] = cursor.fetchall()
    
    # 4. 志願序
    cursor.execute("SELECT PreferenceRank, CompanyName, JobTitle FROM Internship_Preferences WHERE StuID = %s ORDER BY PreferenceRank", (student_id,))
    data['preferences'] = cursor.fetchall()
    
    return data

def generate_application_form_docx(student_data, output_path):
    """
    從結構化資料生成《實習申請表.docx》。
    **重要：請將此函式內容替換為針對您文件排版的實際填充邏輯。**
    """
    try:
        # 建議使用您提供的實習申請表作為範本，但為示範，此處從空白文件開始
        # document = Document('template/實習申請表_template.docx')
        document = Document() 
        
        info = student_data.get('info', {})
        grades = student_data.get('grades', [])
        prefs = student_data.get('preferences', [])
        certs_languages = student_data.get('certs_languages', [])

        document.add_heading('康寧大學 資訊管理科 校外企業實習申請表', 0)
        document.add_paragraph(f'學號：{info.get("StuID", "")}，姓名：{info.get("StuName", "")}').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f'出生日期：{info.get("BirthDate", "")}，性別：{info.get("Gender", "")}')
        document.add_paragraph(f'連絡電話：{info.get("Phone", "")}，電子信箱：{info.get("Email", "")}')
        document.add_paragraph(f'地址：{info.get("Address", "")}')

        # 填充修習專業核心科目 (複雜表格，僅示範結構)
        document.add_heading('已修習專業核心科目', level=1)
        grade_table = document.add_table(rows=len(grades) // 3 + 2, cols=9)
        grade_table.cell(0, 0).text = '科目名稱'
        grade_table.cell(0, 1).text = '學分'
        grade_table.cell(0, 2).text = '成績'
        # ... 此處需要複雜的表格填充邏輯來匹配文件結構 ...

        # 填充實習志願填寫
        document.add_heading('實習志願填寫', level=1)
        pref_table = document.add_table(rows=6, cols=3)
        pref_table.cell(0, 0).text = '志願序'
        pref_table.cell(0, 1).text = '公司名稱'
        pref_table.cell(0, 2).text = '工作項目'
        
        for i in range(5):
            pref = prefs[i] if i < len(prefs) else {'PreferenceRank': i+1, 'CompanyName': '', 'JobTitle': ''}
            pref_table.cell(i + 1, 0).text = str(pref['PreferenceRank'])
            pref_table.cell(i + 1, 1).text = pref['CompanyName']
            pref_table.cell(i + 1, 2).text = pref['JobTitle']

        # 填充其他資訊... (證照、自傳等)

        document.save(output_path)
        return True

    except Exception as e:
        print(f"生成 Word 文件錯誤: {e}")
        traceback.print_exc()
        return False


# ------------------------
# 核心 API 路由 (新增)
# ------------------------

@resume_bp.route('/api/submit_and_generate', methods=['POST'])
def submit_and_generate_api():
    """
    處理學生表單提交、資料正規化儲存、文件生成，並建立上傳紀錄。
    """
    try:
        # 權限檢查
        if session.get('role') != 'student' or not session.get('user_id'):
            return jsonify({"success": False, "message": "只有學生可以提交申請"}), 403

        user_id = session.get('user_id')
        data = request.get_json() 
        if not data:
             return jsonify({"success": False, "message": "缺少提交資料"}), 400
        
        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        # 取得學生學號 (StuID)
        cursor.execute("SELECT username FROM users WHERE id = %s", (user_id,))
        user_info = cursor.fetchone()
        student_id = user_info['username'] # 假設 StuID 存於 users.username 欄位
        
        # 1. 資料儲存與正規化
        if not save_structured_data(cursor, student_id, data):
             conn.close()
             return jsonify({"success": False, "message": "資料儲存失敗"}), 500
        
        # 2. 準備生成文件所需資料
        student_data_for_doc = get_student_info_for_doc(cursor, student_id)
        
        # 3. 自動生成檔案
        original_filename = f"{student_id}_實習申請表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        safe_filename = secure_filename(original_filename)
        save_path = os.path.join(UPLOAD_FOLDER, safe_filename)

        if not generate_application_form_docx(student_data_for_doc, save_path):
            conn.close()
            return jsonify({"success": False, "message": "文件生成失敗，請檢查後端日誌"}), 500

        # 4. 建立上傳紀錄
        semester_id = get_current_semester_id(cursor) # 假設此函式存在
        filesize = os.path.getsize(save_path)
        db_filepath = save_path.replace("\\", "/") # 處理路徑分隔符

        cursor.execute("""
            INSERT INTO resumes (user_id, semester_id, original_filename, filepath, filesize, status, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, NOW())
        """, (user_id, semester_id, original_filename, db_filepath, filesize, 'uploaded'))

        resume_id = cursor.lastrowid
        conn.commit()
        
        # 5. 回傳成功訊息和下載資訊
        return jsonify({
            "success": True,
            "resume_id": resume_id,
            "filename": original_filename,
            "status": "uploaded",
            "message": "申請表已生成並提交",
            "download_url": f"/api/download_resume/{resume_id}"
        })

    except Exception as e:
        traceback.print_exc()
        conn = get_db()
        conn.rollback()
        return jsonify({"success": False, "message": f"伺服器提交錯誤: {str(e)}"}), 500
    finally:
        if 'cursor' in locals() and cursor:
            cursor.close()
        if 'conn' in locals() and conn:
            conn.close()


# -------------------------
#  履歷上傳
# -------------------------
@resume_bp.route('/api/upload_resume', methods=['POST'])
def upload_resume_api():
    try:
        # 取得 session 角色
        role = session.get('role')
        if role != 'student':
            return jsonify({"success": False, "message": "只有學生可以上傳履歷"}), 403

        if 'resume' not in request.files:
            return jsonify({"success": False, "message": "未上傳檔案"}), 400

        file = request.files['resume']
        username = session.get('username')
        if not username:
            return jsonify({"success": False, "message": "未登入學生帳號"}), 403

        if file.filename == '':
            return jsonify({"success": False, "message": "檔案名稱為空"}), 400

        # 取得原始檔名、保護檔名
        original_filename = file.filename
        safe_filename = secure_filename(original_filename)
        # 取得副檔名
        ext = os.path.splitext(safe_filename)[1]  # 包含點，如 ".pdf"、".docx"

        # 用時間戳 + 副檔名做儲存檔名
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        stored_filename = f"{timestamp}{ext}"
        save_path = os.path.join(UPLOAD_FOLDER, stored_filename)

        file.save(save_path)

        # 儲存到資料庫時統一用斜線
        db_filepath = save_path.replace("\\", "/")

        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        if not user:
            cursor.close()
            conn.close()
            if os.path.exists(save_path):
                os.remove(save_path)
            return jsonify({"success": False, "message": "找不到使用者"}), 404

        user_id = user['id']
        filesize = os.path.getsize(save_path)

        # =========================================================
        # 自動標註：學期、班級、學號
        # =========================================================
        semester_id = get_current_semester_id(cursor)
        cursor.execute("SELECT class_id FROM users WHERE id = %s", (user_id,))
        user_info = cursor.fetchone()
        class_id = user_info['class_id'] if user_info else None

        # 插入履歷
        cursor.execute("""
            INSERT INTO resumes (user_id, semester_id, original_filename, filepath, filesize, status, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, NOW())
        """, (user_id, semester_id, original_filename, db_filepath, filesize, 'uploaded'))

        resume_id = cursor.lastrowid
        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({
            "success": True,
            "resume_id": resume_id,
            "filename": original_filename,
            "filesize": filesize,
            "status": "uploaded",
            "message": "履歷上傳成功"
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"上傳失敗: {str(e)}"}), 500

# -------------------------
# API - 下載履歷
# -------------------------
@resume_bp.route('/api/download_resume/<int:resume_id>', methods=['GET'])
def download_resume(resume_id):
    try:
        # 檢查登入（所有角色皆須登入）
        if not require_login():
            return jsonify({"success": False, "message": "未授權"}), 403

        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        # 取得 resume 與 owner
        cursor.execute("""
            SELECT r.filepath, r.original_filename, r.user_id
            FROM resumes r
            WHERE r.id = %s
        """, (resume_id,))
        resume = cursor.fetchone()
        if not resume:
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "找不到履歷"}), 404

        # 權限檢查（TA 和其他讀取角色會透過 can_access_target_resume）
        if not can_access_target_resume(cursor, session['user_id'], session['role'], resume['user_id']):
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "沒有權限下載該履歷"}), 403

        filepath = resume['filepath']
        cursor.close()
        conn.close()

        if not filepath or not os.path.exists(filepath):
            return jsonify({"success": False, "message": "檔案不存在"}), 404

        return send_file(filepath, as_attachment=True, download_name=resume["original_filename"])

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"下載失敗: {str(e)}"}), 500

# -------------------------
# API - 查詢使用者履歷列表（含權限檢查）
# -------------------------
@resume_bp.route('/api/list_resumes/<username>', methods=['GET'])
def list_resumes(username):
    try:
        role = session.get('role')
        user_id = session.get('user_id')

        if role is None:
            # 訪客無權查詢履歷
            return jsonify({"success": False, "message": "訪客無法查看履歷"}), 403

        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        user = get_user_by_username(cursor, username)
        if not user:
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "找不到使用者"}), 404

        target_user_id = user['id']

        # 權限檢查
        if not can_access_target_resume(cursor, user_id, role, target_user_id):
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "沒有權限查看該使用者的履歷"}), 403

        cursor.execute("""
            SELECT id, original_filename, status, comment, note, created_at
            FROM resumes
            WHERE user_id = %s
            ORDER BY created_at DESC
        """, (target_user_id,))
        resumes = cursor.fetchall()

        for r in resumes:
            if isinstance(r.get('created_at'), datetime):
                r['created_at'] = r['created_at'].strftime("%Y-%m-%d %H:%M:%S")

        cursor.close()
        conn.close()
        return jsonify({"success": True, "resumes": resumes})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"查詢失敗: {str(e)}"}), 500

# -------------------------
# API - 審核履歷（合併 approve/reject 的邏輯）
# -------------------------
@resume_bp.route('/api/review_resume/<int:resume_id>', methods=['POST'])
def review_resume(resume_id):
    if 'user_id' not in session:
        return jsonify({"success": False, "message": "未授權"}), 403

    user_id = session['user_id']
    role = session.get('role')
    data = request.get_json() or {}
    status = data.get("status")
    comment = data.get("comment", "")
    note = data.get("note", "")

    if status not in ["approved", "rejected"]:
        return jsonify({"success": False, "message": "無效的狀態"}), 400

    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        # 查履歷對應學生與班級
        cursor.execute("""
            SELECT r.id, r.user_id, u.class_id, c.department
            FROM resumes r
            JOIN users u ON r.user_id = u.id
            LEFT JOIN classes c ON u.class_id = c.id
            WHERE r.id = %s
        """, (resume_id,))
        resume = cursor.fetchone()

        if not resume:
            return jsonify({"success": False, "message": "找不到履歷"}), 404

        target_user_id = resume['user_id']

        # 權限檢查
        if role in ["class_teacher"]:
            if not teacher_manages_class(cursor, user_id, resume['class_id']):
                return jsonify({"success": False, "message": "沒有權限審核這份履歷"}), 403

        elif role == "director":
            director_dept = get_director_department(cursor, user_id)
            if not director_dept or director_dept != resume.get('department'):
                return jsonify({"success": False, "message": "主任無權限審核其他科系的履歷"}), 403

        elif role == "admin":
            pass  # admin 可以

        else:
            # ta, student, 其他角色不可審核
            return jsonify({"success": False, "message": "角色無權限審核"}), 403

        # 更新履歷狀態與備註
        cursor.execute("""
            UPDATE resumes
            SET status = %s, comment = %s, note = %s, updated_at = NOW()
            WHERE id = %s
        """, (status, comment, note, resume_id))
        
        
        # 處理通知邏輯 (退件或通過)
        
        # 獲取學生與審核者信息
        cursor.execute("""
            SELECT u.username, u.name, u.email
            FROM users u
            WHERE u.id = %s
        """, (target_user_id,))
        student = cursor.fetchone()

        cursor.execute("""
            SELECT u.name
            FROM users u
            WHERE u.id = %s
        """, (user_id,))
        reviewer = cursor.fetchone()
        reviewer_name = reviewer['name'] if reviewer else "老師"


        if student:
            # ==================================
            # 1. 處理退件通知 (Rejected)
            # ==================================
            if status == "rejected":
                
                # 創建退件通知（系統通知）
                try:
                    cursor.execute("""
                        INSERT INTO notifications (user_id, title, message, link_url, is_read, created_at)
                        VALUES (%s, %s, %s, %s, 0, NOW())
                    """, (
                        target_user_id,
                        "履歷退件通知",
                        f"您的履歷已被{reviewer_name}退件。\n\n退件原因：{comment if comment else '請查看老師留言'}\n\n請根據老師的建議修改履歷後重新上傳。",
                        '/upload_resume'
                    ))
                except Exception as e:
                    print(f"⚠️ 創建退件通知時發生錯誤: {e}")
                    pass
                
                # 發送郵件通知（如果學生有郵箱）
                if student.get('email'):
                    try:
                        send_resume_rejection_email(
                            student_email=student['email'],
                            student_name=student['name'],
                            reviewer_name=reviewer_name,
                            rejection_reason=comment if comment else ""
                        )
                    except Exception as e:
                        print(f"⚠️ 發送退件郵件時發生錯誤: {e}")
                        pass

            # ==================================
            # 2. 處理通過通知 (Approved) (新增邏輯)
            # ==================================
            elif status == "approved":
                
                # 創建通過通知（系統通知）
                try:
                    cursor.execute("""
                        INSERT INTO notifications (user_id, title, message, link_url, is_read, created_at)
                        VALUES (%s, %s, %s, %s, 0, NOW())
                    """, (
                        target_user_id,
                        "履歷審核通過通知",
                        f"您的履歷已由{reviewer_name}審核通過！您現在可以進行後續的實習步驟。",
                        '/upload_resume' # 連結到一個能查看履歷狀態的頁面
                    ))
                except Exception as e:
                    print(f"⚠️ 創建通過通知時發生錯誤: {e}")
                    pass
                
                # 發送郵件通知（如果學生有郵箱）
                if student.get('email'):
                    try:
                        # 呼叫新的郵件發送函式
                        send_resume_approval_email(
                            student_email=student['email'],
                            student_name=student['name'],
                            reviewer_name=reviewer_name
                        )
                    except Exception as e:
                        print(f"⚠️ 發送通過郵件時發生錯誤: {e}")
                        pass


        # 提交所有資料庫變更（包含 UPDATE resumes 和 INSERT notifications）
        conn.commit()

        return jsonify({"success": True, "message": "履歷審核成功"})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": str(e)}), 500

    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 查詢自己的履歷列表 (學生)
# -------------------------
@resume_bp.route('/api/get_my_resumes', methods=['GET'])
def get_my_resumes():
    if 'user_id' not in session or session.get('role') != 'student':
        return jsonify({"success": False, "message": "未授權"}), 403

    user_id = session['user_id']

    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        cursor.execute("""
            SELECT r.id, r.original_filename, r.status, r.comment, r.note, r.created_at AS upload_time
            FROM resumes r
            WHERE r.user_id = %s
            ORDER BY r.created_at DESC
        """, (user_id,))
        resumes = cursor.fetchall()

        for r in resumes:
            if isinstance(r.get('upload_time'), datetime):
                r['upload_time'] = r['upload_time'].strftime("%Y-%m-%d %H:%M:%S")

        return jsonify({"success": True, "resumes": resumes})
    
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 更新履歷欄位（comment, note）（含權限檢查）
# -------------------------
@resume_bp.route('/api/update_resume_field', methods=['POST'])
def update_resume_field():
    try:
        if not require_login():
            return jsonify({"success": False, "message": "未授權"}), 403

        data = request.get_json() or {}
        resume_id = data.get('resume_id')
        field = data.get('field')
        value = (data.get('value') or '').strip()

        allowed_fields = {
            "comment": "comment",
            "note": "note"
        }

        try:
            resume_id = int(resume_id)
        except (TypeError, ValueError):
            return jsonify({"success": False, "message": "resume_id 必須是數字"}), 400

        if field not in allowed_fields:
            return jsonify({"success": False, "message": "參數錯誤"}), 400

        conn = get_db()
        cursor = conn.cursor(dictionary=True)

        # 先找出 resume 的 owner
        cursor.execute("SELECT user_id FROM resumes WHERE id = %s", (resume_id,))
        r = cursor.fetchone()
        if not r:
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "找不到該履歷"}), 404

        owner_id = r['user_id']

        # 取得使用者角色與 id
        role = session.get('role')
        user_id = session['user_id']

        if role == "class_teacher":
            if not teacher_manages_class(cursor, user_id, get_user_by_id(cursor, owner_id)['class_id']):
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "沒有權限修改該履歷"}), 403

        elif role == "director":
            director_dept = get_director_department(cursor, user_id)
            cursor.execute("SELECT c.department FROM classes c JOIN users u ON u.class_id = c.id WHERE u.id = %s", (owner_id,))
            target_dept_row = cursor.fetchone()
            if not director_dept or not target_dept_row or director_dept != target_dept_row.get('department'):
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "沒有權限修改該履歷"}), 403

        elif role == "admin":
            pass  # admin 可以

        elif role == "student":
            # 學生只能修改自己的履歷，且只能修改 note 欄位
            if user_id != owner_id:
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "學生只能修改自己的履歷"}), 403
            if field != "note":
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "學生只能修改備註欄位"}), 403

        else:
            # ta 或其他角色不可修改
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "角色無權限修改"}), 403

        # 更新欄位
        sql = f"UPDATE resumes SET {allowed_fields[field]} = %s, updated_at = NOW() WHERE id = %s"
        cursor.execute(sql, (value, resume_id))
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({"success": True, "field": field, "resume_id": resume_id})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"伺服器錯誤: {str(e)}"}), 500
    
# -------------------------
# API - 查詢履歷狀態
# -------------------------
@resume_bp.route('/api/resume_status', methods=['GET'])
def resume_status():
    resume_id = request.args.get('resume_id')
    if not resume_id:
        return jsonify({"success": False, "message": "缺少 resume_id"}), 400

    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT status FROM resumes WHERE id = %s", (resume_id,))
        resume = cursor.fetchone()
        cursor.close()
        conn.close()

        if not resume:
            return jsonify({"success": False, "message": "找不到該履歷"}), 404

        return jsonify({"success": True, "status": resume['status']})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"伺服器錯誤: {str(e)}"}), 500

# -------------------------
# API - 查詢所有學生履歷（根據 username，含讀取權限檢查）
# -------------------------
@resume_bp.route('/api/get_student_resumes', methods=['GET'])
def get_student_resumes():
    if not require_login():
        return jsonify({"success": False, "message": "未授權"}), 403

    username = request.args.get('username')
    if not username:
        return jsonify({"success": False, "message": "缺少 username"}), 400

    user_id = session['user_id']
    role = session['role']

    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        cursor.execute("""
            SELECT u.id AS student_id, u.class_id, c.department
            FROM users u
            LEFT JOIN classes c ON u.class_id = c.id
            WHERE u.username = %s
        """, (username,))
        student = cursor.fetchone()
        if not student:
            return jsonify({"success": False, "message": "找不到學生"}), 404

        # 權限判斷（讀取）
        if role == "teacher":
            if not teacher_manages_class(cursor, user_id, student['class_id']):
                return jsonify({"success": False, "message": "沒有權限查看該學生履歷"}), 403

        elif role == "director":
            director_dept = get_director_department(cursor, user_id)
            if not director_dept or director_dept != student.get('department'):
                return jsonify({"success": False, "message": "沒有權限查看該學生履歷"}), 403

        elif role == "ta":
            pass  # TA 可讀全部（如需限制可在此修改）

        elif role == "admin":
            pass

        else:
            return jsonify({"success": False, "message": "角色無權限"}), 403

        # 取得該學生履歷
        cursor.execute("""
            SELECT r.id, r.original_filename, r.status, r.comment, r.note, r.created_at AS upload_time
            FROM resumes r
            WHERE r.user_id = %s
            ORDER BY r.created_at DESC
        """, (student['student_id'],))
        resumes = cursor.fetchall()

        for r in resumes:
            if isinstance(r.get('upload_time'), datetime):
                r['upload_time'] = r['upload_time'].strftime("%Y-%m-%d %H:%M:%S")

        return jsonify({"success": True, "resumes": resumes})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"伺服器錯誤: {str(e)}"}), 500

    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 取得班導 / 主任 履歷 (支援多班級 & 全系)（讀取）
# -------------------------
@resume_bp.route("/api/get_class_resumes", methods=["GET"])
def get_class_resumes():
    # 驗證登入
    if not require_login():
        return jsonify({"success": False, "message": "未授權"}), 403

    user_id = session['user_id']
    role = session['role']
    # mode: "homeroom" 僅看自己班；"director" 主任模式看全科；預設為 homeroom 對 teacher；director 預設依實際頁面傳入
    mode = request.args.get('mode', '').strip().lower()

    conn = get_db()
    cursor = conn.cursor(dictionary=True)

    try:
        resumes = []  # 初始化結果列表
        sql_query = ""
        sql_params = tuple()

        print(f"🔍 [DEBUG] get_class_resumes called - user_id: {user_id}, role: {role}")

        # ------------------------------------------------------------------
        # 1. 班導 (role == "class_teacher")
        # ------------------------------------------------------------------
        if role in ["class_teacher"]:
            sql_query = """
                SELECT 
                    r.id,
                    u.name AS student_name,
                    u.username AS student_number,
                    c.name AS class_name,
                    c.department,
                    r.original_filename,
                    r.filepath,
                    r.status,
                    r.comment,
                    r.note,
                    r.created_at
                FROM resumes r
                JOIN users u ON r.user_id = u.id
                LEFT JOIN classes c ON u.class_id = c.id
                JOIN classes_teacher ct ON ct.class_id = c.id
                WHERE ct.teacher_id = %s
                ORDER BY c.name, u.name
            """
            sql_params = (user_id,)

            cursor.execute(sql_query, sql_params)
            resumes = cursor.fetchall()

            if not resumes:
                print(f"⚠️ [DEBUG] Teacher/class_teacher user {user_id} has no assigned classes.")
                resumes = []

        # ------------------------------------------------------------------
        # 2. 主任 (role == "director")
        # ------------------------------------------------------------------
        elif role == "director":
            # director 根據 mode 控制可見範圍：
            # - mode=director → 同科系全部
            # - 其他/預設 → 僅自己帶的班級（班導模式）
            if mode == "director":
                # 取得主任所屬科系（使用 helper）
                department = get_director_department(cursor, user_id)

                if not department:
                    # 沒有設定科系 → 不顯示任何資料，以免越權
                    resumes = []
                    sql_query = ""
                    sql_params = tuple()
                else:
                    sql_query = """
                        SELECT 
                            r.id,
                            u.name AS student_name,
                            u.username AS student_number,
                            c.name AS class_name,
                            c.department,
                            r.original_filename,
                            r.filepath,
                            r.status,
                            r.comment,
                            r.note,
                            r.created_at
                        FROM resumes r
                        JOIN users u ON r.user_id = u.id
                        JOIN classes c ON u.class_id = c.id
                        WHERE c.department = %s
                        ORDER BY c.name, u.name
                    """
                    sql_params = (department,)
            else:
                # homeroom/預設：僅看自己帶的班級
                sql_query = """
                    SELECT 
                        r.id,
                        u.name AS student_name,
                        u.username AS student_number,
                        c.name AS class_name,
                        c.department,
                        r.original_filename,
                        r.filepath,
                        r.status,
                        r.comment,
                        r.note,
                        r.created_at
                    FROM resumes r
                    JOIN users u ON r.user_id = u.id
                    LEFT JOIN classes c ON u.class_id = c.id
                    JOIN classes_teacher ct ON ct.class_id = c.id
                    WHERE ct.teacher_id = %s
                    ORDER BY c.name, u.name
                """
                sql_params = (user_id,)

            # 執行 SQL 查詢 (主任邏輯在上面已完成查詢或準備好查詢字串)
            if sql_query:
                cursor.execute(sql_query, sql_params)
                resumes = cursor.fetchall()

        # ------------------------------------------------------------------
        # 3. TA 或 Admin (role == "ta" or "admin")
        # ------------------------------------------------------------------
        elif role in ["ta", "admin"]:
            sql_query = """
                SELECT 
                    r.id,
                    u.name AS student_name,
                    u.username AS student_number,
                    c.name AS class_name,
                    c.department,
                    r.original_filename,
                    r.filepath,
                    r.status,
                    r.comment,
                    r.note,
                    r.created_at
                FROM resumes r
                JOIN users u ON r.user_id = u.id
                LEFT JOIN classes c ON u.class_id = c.id
                ORDER BY c.name, u.name
            """
            cursor.execute(sql_query, tuple())
            resumes = cursor.fetchall()

        else:
            return jsonify({"success": False, "message": "無效的角色或權限"}), 403

        # 格式化日期時間並統一字段名稱
        for r in resumes:
            if isinstance(r.get('created_at'), datetime):
                r['created_at'] = r['created_at'].strftime("%Y-%m-%d %H:%M:%S")
            # 統一字段名稱，確保前端能正確訪問
            if 'student_name' in r:
                r['name'] = r['student_name']
            if 'student_number' in r:
                r['username'] = r['student_number']
            if 'class_name' in r:
                r['className'] = r['class_name']
            if 'created_at' in r:
                r['upload_time'] = r['created_at']

        print(f"✅ [DEBUG] Returning {len(resumes)} resumes for role {role}")
        return jsonify({"success": True, "resumes": resumes})

    except Exception:
        print("❌ 取得班級履歷資料錯誤：", traceback.format_exc())
        return jsonify({"success": False, "message": "伺服器錯誤"}), 500

    finally:
        cursor.close()
        conn.close()

# -------------------------
# API - 刪除履歷（需寫入權限）
# -------------------------
@resume_bp.route('/api/delete_resume', methods=['DELETE'])
def delete_resume():
    if not require_login():
        return jsonify({"success": False, "message": "未授權"}), 403

    resume_id = request.args.get('resume_id')
    if not resume_id:
        return jsonify({"success": False, "message": "缺少 resume_id"}), 400

    try:
        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT filepath, user_id FROM resumes WHERE id = %s", (resume_id,))
        result = cursor.fetchone()
        if not result:
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "找不到該履歷"}), 404

        owner_id = result['user_id']
        role = session['role']
        user_id = session['user_id']

        # 權限檢查開始
        if role == "class_teacher":
            # 取得 owner 的 class_id
            cursor.execute("SELECT class_id FROM users WHERE id = %s", (owner_id,))
            owner = cursor.fetchone()
            if not owner or not teacher_manages_class(cursor, user_id, owner.get('class_id')):
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "沒有權限刪除該履歷"}), 403

        elif role == "director":
            director_dept = get_director_department(cursor, user_id)
            cursor.execute("SELECT c.department FROM classes c JOIN users u ON u.class_id = c.id WHERE u.id = %s", (owner_id,))
            target_dept_row = cursor.fetchone()
            if not director_dept or not target_dept_row or director_dept != target_dept_row.get('department'):
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "沒有權限刪除該履歷"}), 403

        elif role == "admin":
            pass

        # 學生只能刪除自己的履歷
        elif role == "student": 
            if user_id != owner_id:
                # 嚴格確保學生只能刪除自己的履歷
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "學生只能刪除自己的履歷"}), 403
            pass # 自己的履歷，允許繼續執行刪除
            
        else:
            # ta, others 無刪除權限
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "角色無權限刪除"}), 403
        # 權限檢查結束

        # 刪除檔案與資料
        filepath = result['filepath']
        if filepath and os.path.exists(filepath):
            os.remove(filepath)

        cursor.execute("DELETE FROM resumes WHERE id = %s", (resume_id,))
        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({"success": True, "message": "履歷已刪除"})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"伺服器錯誤: {str(e)}"}), 500

# -------------------------
# API - submit_comment（寫入 note，整合 update_resume_field）
# -------------------------
@resume_bp.route('/api/submit_comment', methods=['POST'])
def submit_comment():
    try:
        # 直接呼叫 update_resume_field 的邏輯會比較乾淨，但為保持原 API 也支援，我用相同的權限檢查
        data = request.get_json() or {}
        resume_id = data.get('resume_id')
        comment = (data.get('comment') or '').strip()

        if not resume_id or not comment:
            return jsonify({"success": False, "message": "缺少必要參數"}), 400

        try:
            resume_id = int(resume_id)
        except (TypeError, ValueError):
            return jsonify({"success": False, "message": "resume_id 必須是數字"}), 400

        conn = get_db()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT id, user_id FROM resumes WHERE id=%s", (resume_id,))
        r = cursor.fetchone()
        if not r:
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "找不到該履歷"}), 404

        owner_id = r['user_id']

        # 權限檢查（寫入）
        role = session.get('role')
        user_id = session.get('user_id')
        if role == "class_teacher":
            cursor.execute("SELECT class_id FROM users WHERE id = %s", (owner_id,))
            owner = cursor.fetchone()
            if not owner or not teacher_manages_class(cursor, user_id, owner.get('class_id')):
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "沒有權限更新留言"}), 403
        elif role == "director":
            director_dept = get_director_department(cursor, user_id)
            cursor.execute("SELECT c.department FROM classes c JOIN users u ON u.class_id = c.id WHERE u.id = %s", (owner_id,))
            target_dept_row = cursor.fetchone()
            if not director_dept or not target_dept_row or director_dept != target_dept_row.get('department'):
                cursor.close()
                conn.close()
                return jsonify({"success": False, "message": "沒有權限更新留言"}), 403
        elif role == "admin":
            pass
        else:
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "角色無權限更新留言"}), 403

        cursor.execute("UPDATE resumes SET note=%s, updated_at=NOW() WHERE id=%s", (comment, resume_id))
        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({"success": True, "message": "留言更新成功"})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "message": f"伺服器錯誤: {str(e)}"}), 500

# ------------------------
# 頁面路由
# ------------------------

#上傳履歷頁面 (原本用於檔案上傳，現在可作為單純檔案管理頁面)
@resume_bp.route('/upload_resume')
def upload_resume_page():
    return render_template('resume/upload_resume.html')

#審核履歷頁面
@resume_bp.route('/review_resume')
def review_resume_page():
    return render_template('resume/review_resume.html')