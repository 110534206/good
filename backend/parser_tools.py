import docx
from config import get_db
import traceback

def parse_and_update_resume(file_path):
    """
    讀取 Word 檔案內容，提取學號與自傳，並更新至資料庫。
    """
    try:
        # 1. 讀取 Word 檔案
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # 將所有文字合併，方便搜尋
        content = "\n".join(full_text)
        
        # 2. 提取學號 (利用正則表達式或關鍵字定位)
        # 根據你的範例，格式通常是 "學號：110534244"
        import re
        stu_id_match = re.search(r"學號：(\d+)", content)
        stu_id = stu_id_match.group(1) if stu_id_match else None
        
        # 3. 提取自傳內容
        # 尋找「自傳」標題之後的所有文字
        autobiography = ""
        found_bio = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if "自傳" in text and len(text) < 5: # 找到「自傳」這兩個字作為標題
                found_bio = True
                continue
            if found_bio:
                autobiography += text + "\n"
        
        if not stu_id or not autobiography:
            print("❌ 解析失敗：找不到學號或自傳內容")
            return False

        # 4. 寫入資料庫 (Student_Info 表)
        conn = get_db()
        cursor = conn.cursor()
        
        sql = """
            INSERT INTO Student_Info (StuID, Autobiography, UpdatedAt)
            VALUES (%s, %s, NOW())
            ON DUPLICATE KEY UPDATE
                Autobiography = VALUES(Autobiography),
                UpdatedAt = NOW()
        """
        cursor.execute(sql, (stu_id, autobiography.strip()))
        conn.commit()
        
        print(f"✅ 成功解析並更新資料庫！學號: {stu_id}")
        return True

    except Exception as e:
        print(f"❌ 發生錯誤: {e}")
        traceback.print_exc()
        return False
    finally:
        if 'cursor' in locals(): cursor.close()
        if 'conn' in locals(): conn.close()

# 測試執行
# parse_and_update_resume("uploads/110534244_履歷.docx")